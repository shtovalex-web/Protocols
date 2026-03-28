# -*- coding: utf-8 -*-
"""Приложение: формирование протоколов проверки знаний (tkinter)."""

from __future__ import annotations

import json
import os
import re
import sqlite3
import tempfile
from copy import deepcopy
import tkinter as tk
import xml.etree.ElementTree as ET
import zipfile
from datetime import date
from pathlib import Path
from typing import Any
from tkinter import filedialog, messagebox
from tkinter import ttk

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run as DocxRun
from fpdf import FPDF

from commission_admin import (
    CommissionAdminPanel,
    CommissionState,
    apply_commission_insertions_to_line,
    build_commission_signature_suffix_payload,
    build_commission_template_payload,
    ensure_app_settings_table,
    refresh_commission_pool_from_excel,
)
from employees_io import (
    EMPLOYEES_EXCEL_FILENAME,
    EMPLOYEES_SHEET_NAME,
    EmployeeExcelError,
    EmployeeRecord,
    listbox_label_for_employee,
    load_employees_from_excel,
)


GRADE_OPTIONS = ("удовлетворительно", "неудовлетворительно")
CHECK_TYPE_OPTIONS = ("плановая", "внеплановая")

# Текст бланка протокола без организационной шапки (абзацы от «ПРОТОКОЛ №» и таблица ниже двух строк шапки).
PROTOCOL_BODY_FONT_PT = 11
# Строка «Программа (В)» + темы перед каждым работником в таблице результатов.
PROTOCOL_V_TABLE_HEADER_FONT_PT = 9
PROTOCOL_PREVIEW_HEADER_FONT = ("Segoe UI", 10)
PROTOCOL_PREVIEW_BODY_FONT = ("Segoe UI", PROTOCOL_BODY_FONT_PT)

PROTOCOL_TEMPLATE_FILENAME = "default_protocol.docx"


def _is_word_protocol_template(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() in (".docx", ".docm")
DATABASE_FILENAME = "protocols.db"
LAST_PROTOCOL_NO_STATE_FILENAME = "last_protocol_no.json"


def last_protocol_no_state_path() -> Path:
    return Path(__file__).resolve().parent / LAST_PROTOCOL_NO_STATE_FILENAME


def load_last_protocol_no() -> str:
    p = last_protocol_no_state_path()
    if not p.is_file():
        return ""
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError, TypeError):
        return ""
    if not isinstance(data, dict):
        return ""
    v = data.get("last_protocol_no", "")
    return v.strip() if isinstance(v, str) else ""


def save_last_protocol_no(value: str) -> None:
    p = last_protocol_no_state_path()
    try:
        p.write_text(
            json.dumps({"last_protocol_no": value}, ensure_ascii=False),
            encoding="utf-8",
        )
    except OSError:
        pass


V_PROF_SHEET_NAME = "V_PROF"
# Лист B: столбец 2 — полное название программы «Б» в шапке и в таблице протокола.
B_PROGRAM_SHEET_NAME = "B"
B_PROGRAM_TITLE_COL = 2
# Лист V_PROF: A — должность; 2 — текст для якорных абзацев после «…по программе(ам):» для «Б»;
# 3–4 — PP / СИЗ; с 5 — фрагменты программы «В».
V_PROF_TITLE_COL_B = 2
V_PROF_TITLE_COL_PP = 3
V_PROF_TITLE_COL_SIZ = 4
V_PROF_PARTS_FIRST_COL = 5
V_PROF_PARTS_LAST_COL = 22
# Версия схемы кэша фрагментов программы «В»; при смене столбцов — старые записи игнорируются.
V_PROF_PARTS_CACHE_SCHEMA = 2

# Ключ, лист-источник названия/якорей, текст по умолчанию если ячейка пуста.
PROTOCOL_PROGRAM_DEFS: tuple[tuple[str, str, str], ...] = (
    ("B", B_PROGRAM_SHEET_NAME, "Программа обучения «Б»"),
    (
        "PP",
        V_PROF_SHEET_NAME,
        "Программа обучения по оказанию первой помощи пострадавшим",
    ),
    (
        "SIZ",
        V_PROF_SHEET_NAME,
        "Программа обучения использование (применение) средств индивидуальной защиты",
    ),
    ("V", V_PROF_SHEET_NAME, "Программа обучения «В»"),
)
PROTOCOL_PROGRAM_UI_LABELS: dict[str, str] = {
    "B": (
        f"Программа «Б» (таблица — {B_PROGRAM_SHEET_NAME}, ст. {B_PROGRAM_TITLE_COL}; "
        f"абзац после проверки — {V_PROF_SHEET_NAME}, ст. {V_PROF_TITLE_COL_B})"
    ),
    "PP": f"Первая помощь (лист {V_PROF_SHEET_NAME}, столб. {V_PROF_TITLE_COL_PP})",
    "SIZ": f"СИЗ (лист {V_PROF_SHEET_NAME}, столб. {V_PROF_TITLE_COL_SIZ})",
    "V": f"Программа «В» (лист {V_PROF_SHEET_NAME}, столб. {V_PROF_PARTS_FIRST_COL}–{V_PROF_PARTS_LAST_COL})",
}
# Краткие подписи чекбоксов на главном экране (подробности — в окне администрирования).
PROTOCOL_PROGRAM_CHECKBOX_SHORT: dict[str, str] = {
    "B": "Программа «Б»",
    "PP": "Первая помощь (PP)",
    "SIZ": "СИЗ",
    "V": "Программа «В»",
}

# Полный перечень маркеров шаблона (см. кнопку «Переменные шаблона» в окне программы).
PROTOCOL_TEMPLATE_VARIABLES_DOC = """\
ПЕРЕМЕННЫЕ И МАРКЕРЫ ШАБЛОНА ПРОТОКОЛА
========================================

В шаблоне (.docx или .txt) не используются имена вроде {{ФИО}}. Программа ищет
описанные ниже фрагменты текста и подставляет данные из полей формы.

────────────────────────────────────────
1) НОМЕР ПРОТОКОЛА  →  поле «№ протокола»
────────────────────────────────────────
В шаблоне должна быть строка с фрагментом  ПРОТОКОЛ №  (любой регистр; допускаются # или знак №;
можно после названия организации). Подчёркивания _____ после знака не обязательны.

Если после «№» идут подчёркивания (от трёх символов «_» подряд) — они заменяются на номер;
если подчёркиваний нет — номер вставляется сразу после знака.

Если в форме указан номер — подставляется строка вида
  <номер>-<месяц>-<год>
где <номер> — то, что введено в поле «№ протокола», <месяц> — число 1–12 из поля «Дата»
(без ведущего нуля), <год> — две последние цифры года из «Даты» (например, при номере 12
и дате 25.03.2026 в бланк попадёт 12-3-26; строка шаблона остаётся «ПРОТОКОЛ № …»).
Если дата не в формате ДД.ММ.ГГГГ (или ГГ) — подставляется только введённый номер, без суффикса.
Если номер не введён — строка с «ПРОТОКОЛ №» не меняется (подчёркивания остаются).

────────────────────────────────────────
2) ДАТА  →  поле «Дата» (формат ДД.ММ.ГГГГ)
────────────────────────────────────────
Отдельная строка, у которой после удаления пробелов по краям:
  • начинается с  «__
  • заканчивается на  г.

Типичный вид (как в бланке):
  «__» ___________ 20__ г.

Всю такую строку программа заменяет на дату словами, например:
  «25» марта 2026 г.

Важно: строка «В соответствии с приказом … от «__» … 20__ г.» не подходит
под это правило (она не начинается с «__»).

────────────────────────────────────────
2а) ПРИКАЗ О КОМИССИИ  →  меню «Администрирование» → «Приказ и комиссия…» (база protocols.db)
────────────────────────────────────────
В абзаце, где есть слова о комиссии / приказе (например: комисс, председателя, членов, приказ и « от»),
при сохранённых в базе данных № и дате приказа, председателе и членах:
  • сразу после подстроки « от» (пробел и «от») вставляется дата приказа словами, как в п. 2
    (поле «Дата приказа» в формате ДД.ММ.ГГГГ);
  • после первого в абзаце « №», не входящего в сочетание «ПРОТОКОЛ №», подставляется номер приказа;
  • после слова «председателя» — с новой строки столбиком ФИО и должность выбранного председателя в родительном падеже;
  • после фразы «членов комиссии» (если есть) или слова «членов» — с новой строки по одному человеку на строку
    (ФИО и должность в родительном падеже).
Родительный падеж: при установленных пакетах pymorphy2 и pymorphy2-dicts-ru; иначе текст без склонения.
Рекомендуется: pip install pymorphy2 pymorphy2-dicts-ru

────────────────────────────────────────
3) ПРОГРАММА / ТЕМА  →  чекбоксы программ и строка подчёркиваний
────────────────────────────────────────
Для .docx: лист B — название программы «Б» в шапке и в таблице из первой непустой
ячейки столбца 2 (со 2-й строки). После абзаца «…по программе(ам):» для «Б» подставляется
текст из листа V_PROF, столбец 2, по совпадению должности (столбец A). На V_PROF: «PP» — столбец 3,
«СИЗ» — столбец 4. Программа «В»:
столбец A — должность; фрагменты программы — непустые ячейки столбцов 5–19 в строке
совпадения (сцепка в заголовке блока и в таблице). Кэш V_PROF — в protocols.db
(таблица v_prof_cache; поле parts_schema — сброс кэша фрагментов при смене логики столбцов),
сброс при изменении времени файла Excel.
Для ПП, СИЗ, «В» при совмещаемой второй профессии: в таблице — одна строка на работника после объединения
одинакового ФИО (в графе должности — только основная). Блок «Б» — все выбранные строки списка,
даже если ФИО повторяется. Фрагменты программы «В»
по двум должностям объединяются без повторов; строк «№ …» в ячейке результата столько же,
сколько фрагментов. Для «Б» при совмещении — две строки «№ …» (две записи в реестре);
для PP и СИЗ — одна строка «№ …» (обучение одно). Для «В» в таблице перед каждым работником
добавляется строка-шапка: «Программа (В)» и в скобках только темы этого работника (как в примере 123 прогВ.docx);
в первой колонке номер блока программы — только у шапки первого работника в блоке, у следующих шапок пусто; в колонке «ФИО» — только ФИО.
В шапке протокола названия программ через «; ».

Отдельная строка с подчёркиваниями должна:
  • в конце (после пробелов) иметь запятую:  ,
  • состоять в основном из символов подчёркивания _ (длина > 20),
  • кроме подчёркиваний, пробелов и запятой не содержать другого текста.

Для .txt или если программы не отмечены: используется поле «Доп. тема».

После абзаца, в котором есть слова: провела, проверку, программе, охраны труда
(типично: «…провела проверку знания требований охраны труда работников по программе(ам):»),
при отмеченных программах B / PP / SIZ / V подставляются абзацы с префиксом «- »:
лист V_PROF, столбец A — должность; для «Б» — ячейка столбца 2, «PP» — 3, «СИЗ» — 4;
для «В» — по одному абзацу на каждую непустую ячейку столбцов 5–19. Если совпадения
по должности или ячейки пусты — соответствующая строка не добавляется.

────────────────────────────────────────
4) ДОПОЛНИТЕЛЬНО ДЛЯ ФАЙЛА .docx (и для .txt со срезом)
────────────────────────────────────────
Чтобы из документа выделялся фрагмент бланка:

  • Нужна строка с маркером начала: фрагмент  ПРОТОКОЛ №  (или #); в этой строке не должно быть
    слова  ТЕХНИЧЕСКИХ  (так отличается пустой бланк от заполненного примера в одном файле).

  • Конец фрагмента: строка, начинающаяся с  Приложение
    (всё после неё в шаблон не попадает).

Если в .txt нет строки с  ПРОТОКОЛ № , обрабатывается весь
файл построчно — правила п. 1–3 всё равно применяются, где найдутся подходящие строки.

  • В самом шаблоне протокола (.docx) в конце бланка можно поставить плейсхолдеры подписей комиссии
    (лучше одним фрагментом в run Word): {{ПРЕДСЕДАТЕЛЬ}} — И.О. Фамилия, должность;
    {{ЧЛЕНЫ_КОМИССИИ}} — члены, по одной строке на человека; также {{CHAIR}}, {{MEMBERS}}.
    Данные берутся из «Приказ и комиссия»; отдельный файл подписей не используется.

────────────────────────────────────────
ИТОГ: какие «переменные» задавать в шаблоне явно
────────────────────────────────────────
  • ____  в строке «ПРОТОКОЛ № ____»     — номер в виде N-М-ГГ (см. п. 1)
  • строка вида «__» … 20__ г.           — дата
  • абзац о комиссии по приказу          — см. п. 2а ( от, №, председателя, членов)
  • длинная строка _____ … ,             — тема в шапке (см. п. 3)
  • в таблице .docx: см. п. 5

────────────────────────────────────────
5) ТАБЛИЦА «Результат проверки знаний»  →  поля формы
────────────────────────────────────────
В шаблоне .docx — таблица: две верхние строки-шапка (п/п, Фамилия, … и 1…7),
затем пример блока «строка программы» + «строка 1.1 с маркером ФИО».

При формировании все строки ниже двух шапок удаляются и строятся заново:
число блоков = числу отмеченных программ (B, PP, SIZ, V). В каждом блоке
первая строка — номер программы (1, 2, …) и название: «Б» — лист B (столб. 2), PP/СИЗ — V_PROF (3–4), «В» — по V_PROF;
далее строки сотрудников N.1…N.M. Для ПП, СИЗ, «В»: M — после объединения записей с одинаковым ФИО
(одна строка на человека; совмещение учитывается внутри записи). Для программы «Б»: M — все выбранные
строки списка, в том числе несколько с одним ФИО. Список работников — лист rabotnik.
  • результат проверки: оценка и регистрационные номера реестра Минтруда
    (поле в окне администрирования) — не путать с номером протокола.
    Несколько номеров — через запятую или с новой строки.
    Для «В»: по одной строке «№ …» на каждый непустой фрагмент (после объединения должностей);
    перед строкой каждого работника вставляется подзаголовок «Программа (В)» и в скобках — только его темы
    (как в шаблоне 123 прогВ.docx); в колонке «ФИО» — только ФИО.
    Для «Б»: при совмещении две строки «№ …», иначе одна; для PP и СИЗ — всегда одна строка «№ …».
  • тип проверки: плановая / внеплановая

Если маркеров «ФИО» в таблице нет, таблица не меняется.

Остальной текст шаблона — произвольный, его можно менять.
"""

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Подчёркивания-заполнители после «ПРОТОКОЛ №»: ASCII _, полноширинная ＿, double low line ‗
_US_PLACEHOLDER_RUN = re.compile(r"(?:_|＿|‗){3,}")

# Якорь бланка: слово «протокол» и знак номера (№ / # / U+2116), без требования подчёркиваний.
_PROTOCOL_NUMBER_ANCHOR_RE = re.compile(r"(?i)протокол\s*(?:№|#|\u2116)")


def _normalize_docx_plain_text(s: str) -> str:
    """Неразрывный пробел и невидимые символы — для надёжного поиска маркеров."""
    if not s:
        return s
    return (
        s.replace("\xa0", " ")
        .replace("\u200b", "")
        .replace("\ufeff", "")
        .replace("\u00ad", "")
    )


def _reject_protocol_line_as_false_positive(s: str) -> bool:
    """«по протоколу …» и т.п. — не бланк."""
    return bool(re.search(r"(?i)по\s+протокол", s))


def _line_has_protocol_form_marker(line: str) -> bool:
    """Строка содержит якорь «ПРОТОКОЛ №» (регистр и пробелы гибкие; допускаются # и знак U+2116)."""
    s = _normalize_docx_plain_text(line)
    if "технических" in s.lower():
        return False
    if _reject_protocol_line_as_false_positive(s):
        return False
    return bool(_PROTOCOL_NUMBER_ANCHOR_RE.search(s))


def _line_fills_protocol_number_slot(line: str) -> bool:
    """Абзац, в котором подставляем номер протокола (есть якорь «ПРОТОКОЛ №»)."""
    return _line_has_protocol_form_marker(line)


def _is_protocol_title_line_without_underscores(line: str) -> bool:
    """«… ПРОТОКОЛ №» в абзаце без символов подчёркивания (номер на следующей строке)."""
    s = _normalize_docx_plain_text(line)
    if "технических" in s.lower():
        return False
    if _reject_protocol_line_as_false_positive(s):
        return False
    if _US_PLACEHOLDER_RUN.search(s) is not None:
        return False
    return bool(_PROTOCOL_NUMBER_ANCHOR_RE.search(s))


def _is_placeholder_only_underscore_line(line: str) -> bool:
    """Абзац только из группы подчёркиваний (и пробелов) — подстановка номера сюда."""
    s = _normalize_docx_plain_text(line).strip()
    m = _US_PLACEHOLDER_RUN.search(s)
    if not m:
        return False
    return not (s[: m.start()].strip() or s[m.end() :].strip())


def _replace_first_underscore_run_with_protocol_number(line: str, pn: str) -> str:
    s = _normalize_docx_plain_text(line)
    s2, n = _US_PLACEHOLDER_RUN.subn(pn, s, count=1)
    return s2 if n else line


def _apply_protocol_number_after_anchor(line: str, pn: str) -> str:
    """
    После «ПРОТОКОЛ №» (/#/U+2116): если дальше идут подчёркивания — заменить их на номер;
    иначе вставить номер сразу после знака (если его ещё нет).
    """
    if not pn:
        return line
    s = _normalize_docx_plain_text(line)
    m = _PROTOCOL_NUMBER_ANCHOR_RE.search(s)
    if not m:
        return line
    pos = m.end()
    tail = s[pos:]
    um = re.match(r"^(\s*)(?:_|＿|‗){3,}", tail)
    if um:
        return s[:pos] + um.group(1) + pn + tail[um.end() :]
    rest = tail.lstrip()
    if rest.startswith(pn):
        return s
    if not tail.strip():
        return s[:pos] + " " + pn
    return s[:pos] + " " + pn + tail


def _ooxml_sym_char(elem: ET.Element) -> str | None:
    for k, v in elem.attrib.items():
        if k.endswith("}char") and v:
            try:
                return chr(int(v, 16))
            except ValueError:
                return None
    return None


def _w_p_element_plain_text(p_el: ET.Element) -> str:
    """Текст абзаца как в Word: w:t, символы w:sym (например №), табуляция, перенос строки."""
    parts: list[str] = []
    W = _W_NS
    for node in p_el.iter():
        tag = node.tag
        if tag == f"{W}t":
            parts.append(node.text or "")
        elif tag == f"{W}sym":
            ch = _ooxml_sym_char(node)
            if ch:
                parts.append(ch)
        elif tag == f"{W}tab":
            parts.append("\t")
        elif tag == f"{W}br":
            parts.append("\n")
    return "".join(parts)


def _sorted_word_hf_xml_names(zf: zipfile.ZipFile, kind: str) -> list[str]:
    prefix = f"word/{kind}"
    return sorted(n for n in zf.namelist() if n.startswith(prefix) and n.endswith(".xml"))


def _paragraph_texts_from_hf_or_fragment_xml(xml_bytes: bytes) -> list[str]:
    root = ET.fromstring(xml_bytes)
    return [_w_p_element_plain_text(p) for p in root.iter(f"{_W_NS}p")]


def _iter_document_paragraphs_in_order(doc: Document):
    """Колонтитулы, затем все w:p тела (таблицы, надписи и т.д.) — в порядке XML."""
    seen_hf: set[int] = set()
    for section in doc.sections:
        hel = section.header._element
        hid = id(hel)
        if hid not in seen_hf:
            seen_hf.add(hid)
            for p_el in hel.iter(qn("w:p")):
                yield DocxParagraph(p_el, doc)
    for p_el in doc.element.body.iter(qn("w:p")):
        yield DocxParagraph(p_el, doc)
    seen_hf.clear()
    for section in doc.sections:
        fel = section.footer._element
        fid = id(fel)
        if fid not in seen_hf:
            seen_hf.add(fid)
            for p_el in fel.iter(qn("w:p")):
                yield DocxParagraph(p_el, doc)


def _all_document_paragraphs_ordered(doc: Document) -> list[DocxParagraph]:
    return list(_iter_document_paragraphs_in_order(doc))


_MONTHS_GEN = (
    "",
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)


class ProtocolTemplateError(Exception):
    """Файл шаблона протокола отсутствует или не читается."""


def protocol_template_path() -> Path:
    return Path(__file__).resolve().parent / PROTOCOL_TEMPLATE_FILENAME


def database_path() -> Path:
    return Path(__file__).resolve().parent / DATABASE_FILENAME


def employees_excel_default_path() -> Path:
    return Path(__file__).resolve().parent / EMPLOYEES_EXCEL_FILENAME


def load_excel_first_nonempty_in_column(
    path: Path, sheet_name: str, column_one_based: int
) -> str:
    """Первая непустая ячейка в столбце листа, начиная со 2-й строки."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    if not path.is_file():
        return ""
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        names = {n.lower(): n for n in wb.sheetnames}
        sk = sheet_name.lower()
        if sk not in names:
            return ""
        ws = wb[names[sk]]
        max_r = min(ws.max_row or 200, 500)
        for r in range(2, max_r + 1):
            v = ws.cell(row=r, column=column_one_based).value
            if v is not None and str(v).strip():
                return str(v).strip()
        return ""
    finally:
        wb.close()


def load_v_prof_first_nonempty_in_column(path: Path, column_one_based: int) -> str:
    """Лист V_PROF: первая непустая ячейка в столбце, начиная со 2-й строки."""
    return load_excel_first_nonempty_in_column(path, V_PROF_SHEET_NAME, column_one_based)


def _collect_unique_professions_ordered(persons: list[EmployeeRecord]) -> list[str]:
    """Уникальные должности из выбранных записей: основная и совмещаемая, порядок первого появления."""
    seen: set[str] = set()
    out: list[str] = []
    for p in persons:
        for pr in (p.profession, p.profession2):
            t = (pr or "").strip()
            if not t:
                continue
            k = _norm_profession_key(t)
            if k in seen:
                continue
            seen.add(k)
            out.append(t)
    return out


def _select_best_row_by_profession_col_a(
    path: Path,
    profession: str,
    sheet_name: str,
    max_col: int,
    tie_break_cols_1based: tuple[int, int] | None,
) -> tuple[Any, ...] | None:
    """
    Строка Excel: лучшее совпадение должности в столбце A (1-based).
    tie_break_cols_1based — (первая, последняя) колонки для суммы длин при равном совпадении;
    None — без доп. критерия.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return None
    if not path.is_file() or not profession.strip():
        return None
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        names = {n.lower(): n for n in wb.sheetnames}
        sk = sheet_name.lower()
        if sk not in names:
            return None
        ws = wb[names[sk]]
        target = _norm_profession_key(profession)
        skip_headers = {
            "профессия",
            "должность",
            "специальность",
            "фио",
            "п/п",
            "наименование",
        }
        best: tuple[int, int, tuple[Any, ...]] | None = None
        for row in ws.iter_rows(
            min_row=1,
            max_row=min(ws.max_row or 500, 2000),
            min_col=1,
            max_col=max_col,
            values_only=True,
        ):
            if not row or row[0] is None:
                continue
            c0 = _norm_profession_key(str(row[0]))
            if not c0 or c0 in skip_headers:
                continue
            match_score = 0
            if c0 == target:
                match_score = 3
            elif target in c0 or c0 in target:
                match_score = 2
            if match_score == 0:
                continue
            row_tuple = tuple(row)
            merged_len = 0
            if tie_break_cols_1based is not None:
                t_lo, t_hi = tie_break_cols_1based
                lo_idx = t_lo - 1
                hi_excl = min(len(row_tuple), t_hi)
                for idx in range(lo_idx, hi_excl):
                    cell = row_tuple[idx]
                    if cell is not None and str(cell).strip():
                        merged_len += len(str(cell).strip())
            cand = (match_score, merged_len, row_tuple)
            if best is None or cand[0] > best[0] or (
                cand[0] == best[0] and cand[1] > best[1]
            ):
                best = cand
        return best[2] if best else None
    finally:
        wb.close()


def _v_prof_select_best_row(path: Path, profession: str) -> tuple[Any, ...] | None:
    """
    Строка листа V_PROF (ячейки 1–19), лучшее совпадение должности в столбце A.
    """
    return _select_best_row_by_profession_col_a(
        path,
        profession,
        V_PROF_SHEET_NAME,
        V_PROF_PARTS_LAST_COL,
        (V_PROF_PARTS_FIRST_COL, V_PROF_PARTS_LAST_COL),
    )


def _v_prof_anchor_line_from_row(row: tuple[Any, ...], column_one_based: int) -> str | None:
    """Одна ячейка заданного столбца строки V_PROF с префиксом «- »."""
    i = column_one_based - 1
    if i < 0 or i >= len(row):
        return None
    v = row[i]
    if v is None:
        return None
    s = str(v).strip()
    if not s:
        return None
    return "- " + s


def _v_prof_anchor_lines_program_v(row: tuple[Any, ...]) -> list[str]:
    """Не пустые ячейки столбцов 5–19, каждая с префиксом «- »."""
    out: list[str] = []
    lim = min(len(row), V_PROF_PARTS_LAST_COL)
    for idx in range(V_PROF_PARTS_FIRST_COL - 1, lim):
        cell = row[idx]
        if cell is None:
            continue
        t = str(cell).strip()
        if t:
            out.append("- " + t)
    return out


def build_fg_lines_for_selected_programs(
    path: Path,
    program_keys: list[str],
    persons_raw: list[EmployeeRecord],
) -> list[str]:
    """
    V_PROF, столбец A — должность; для «Б»/PP/СИЗ — ячейки столбцов 2–4 (якорь после проверки).
    Название блока «Б» в таблице по-прежнему с листа B (см. _collect_program_keys_and_titles).
    «В» — по одной строке на каждую непустую ячейку столбцов фрагментов V_PROF.
    """
    profs = _collect_unique_professions_ordered(persons_raw)
    if not profs or not program_keys:
        return []
    anchor_col_for_key: dict[str, int] = {
        "B": V_PROF_TITLE_COL_B,
        "PP": V_PROF_TITLE_COL_PP,
        "SIZ": V_PROF_TITLE_COL_SIZ,
    }
    row_cache: dict[str, tuple[Any, ...] | None] = {}

    def row_for(pr: str) -> tuple[Any, ...] | None:
        if pr not in row_cache:
            row_cache[pr] = _v_prof_select_best_row(path, pr)
        return row_cache[pr]

    lines: list[str] = []
    for pkey in program_keys:
        for pr in profs:
            row = row_for(pr)
            if not row:
                continue
            if pkey == "V":
                lines.extend(_v_prof_anchor_lines_program_v(row))
            else:
                col = anchor_col_for_key.get(pkey)
                if col is None:
                    continue
                s = _v_prof_anchor_line_from_row(row, col)
                if s:
                    lines.append(s)
    return lines


def _fg_line_comparison_key(s: str) -> str:
    """Ключ для сравнения строк программ (без маркера списка, с нормализацией пробелов и регистра)."""
    t = s.replace("\xa0", " ").strip()
    t = re.sub(r"^[-–—]\s*", "", t)
    return _norm_profession_key(t)


def _program_bullet_paragraph_indices_after_anchor(
    doc: Document, anchor: DocxParagraph
) -> tuple[int, int]:
    """
    Абзацы сразу после anchor, идущие подряд и начинающиеся с маркера списка (-, –, —).
    Возвращает (lo, hi) — полуинтервал индексов в doc.paragraphs; если пусто, lo == hi.
    """
    paras = list(doc.paragraphs)
    try:
        i = paras.index(anchor)
    except ValueError:
        return (0, 0)
    lo = i + 1
    hi = lo
    while hi < len(paras):
        t = paras[hi].text.replace("\xa0", " ").strip()
        if not t:
            break
        if re.match(r"^[-–—]", t):
            hi += 1
        else:
            break
    return (lo, hi)


def _is_anchor_paragraph_program_table(text: str) -> bool:
    """Абзац после которого вставляются строки из V_PROF (в т. ч. столб. 2 для «Б»)."""
    t = text.replace("\xa0", " ").lower()
    return (
        "провела" in t
        and "проверку" in t
        and "программе" in t
        and "охраны труда" in t
    )


def _insert_paragraph_after(
    paragraph: DocxParagraph, text: str = ""
) -> DocxParagraph:
    new_el = OxmlElement("w:p")
    paragraph._p.addnext(new_el)
    new_para = DocxParagraph(new_el, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_program_fg_lines_after_anchor(
    doc: Document,
    lines: list[str],
) -> int:
    """
    Вставляет абзацы после блока про проверку по программе(ам).
    Уже имеющиеся в шаблоне пункты (подряд после якоря, с «-»/«–»/«—») не дублируются;
    из списка на вставку убираются повторы. Новые строки добавляются после последнего
    такого пункта (или сразу после якоря, если пунктов ещё нет).
    """
    if not lines:
        return 0
    anchor: DocxParagraph | None = None
    for para in doc.paragraphs:
        if _is_anchor_paragraph_program_table(para.text):
            anchor = para
            break
    if anchor is None:
        return 0
    lo, hi = _program_bullet_paragraph_indices_after_anchor(doc, anchor)
    paras = list(doc.paragraphs)
    ref: DocxParagraph = paras[hi - 1] if hi > lo else anchor
    existing_keys: set[str] = set()
    for j in range(lo, hi):
        existing_keys.add(_fg_line_comparison_key(paras[j].text))
    to_insert: list[str] = []
    batch_seen: set[str] = set()
    for line in lines:
        k = _fg_line_comparison_key(line)
        if k in existing_keys or k in batch_seen:
            continue
        batch_seen.add(k)
        to_insert.append(line)
    if not to_insert:
        return 0
    for line in to_insert:
        ref = _insert_paragraph_after(ref, line)
    return len(to_insert)


def _norm_profession_key(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower().replace("ё", "е"))


def _table_employees_dedupe_by_fio(records: list[EmployeeRecord]) -> list[EmployeeRecord]:
    """
    Одна строка таблицы на одно ФИО: записи с одинаковым ФИО (несколько строк Excel и т. п.)
    объединяются; уникальные должности идут в profession и profession2 (остальные отбрасываются).
    """
    if not records:
        return []
    buckets: dict[str, list[EmployeeRecord]] = {}
    order_keys: list[str] = []
    for i, p in enumerate(records):
        key = _norm_profession_key(p.fio or "")
        if not key:
            key = f"__noid_{i}__"
        if key not in buckets:
            buckets[key] = []
            order_keys.append(key)
        buckets[key].append(p)
    out: list[EmployeeRecord] = []
    for key in order_keys:
        grp = buckets[key]
        base = grp[0]
        profs_ordered: list[str] = []

        def add_prof(s: str) -> None:
            t = (s or "").strip()
            if not t:
                return
            nk = _norm_profession_key(t)
            if not any(_norm_profession_key(x) == nk for x in profs_ordered):
                profs_ordered.append(t)

        add_prof(base.profession)
        add_prof(base.profession2)
        for extra in grp[1:]:
            add_prof(extra.profession)
            add_prof(extra.profession2)
        main_p = profs_ordered[0] if profs_ordered else (base.profession or "").strip()
        p2 = profs_ordered[1] if len(profs_ordered) > 1 else ""
        sub = (base.subdivision or "").strip()
        for extra in grp[1:]:
            if not sub:
                sub = (extra.subdivision or "").strip()
        out.append(
            EmployeeRecord(
                fio=base.fio,
                profession=main_p,
                subdivision=sub,
                profession2=p2,
            )
        )
    return out


def _ensure_v_prof_cache_table(conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS v_prof_cache (
            file_path TEXT NOT NULL,
            profession_norm TEXT NOT NULL,
            file_mtime REAL NOT NULL,
            row_text TEXT NOT NULL,
            parts_json TEXT,
            PRIMARY KEY (file_path, profession_norm)
        )
        """
    )
    try:
        conn.execute("ALTER TABLE v_prof_cache ADD COLUMN parts_json TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        conn.execute("ALTER TABLE v_prof_cache ADD COLUMN parts_schema INTEGER DEFAULT 1")
    except sqlite3.OperationalError:
        pass


def _v_cache_invalidate_if_stale(conn: sqlite3.Connection, path_resolved: str, mtime: float) -> None:
    row = conn.execute(
        "SELECT file_mtime FROM v_prof_cache WHERE file_path = ? LIMIT 1",
        (path_resolved,),
    ).fetchone()
    if row is not None and row[0] != mtime:
        conn.execute("DELETE FROM v_prof_cache WHERE file_path = ?", (path_resolved,))


def v_cache_get_row_text(path: Path, profession: str) -> str | None:
    """Кэш строки V_PROF по должности (обновляется при смене времени файла Excel)."""
    if not path.is_file() or not profession.strip():
        return None
    p = str(path.resolve())
    mtime = path.stat().st_mtime
    key = _norm_profession_key(profession)
    with sqlite3.connect(database_path()) as conn:
        _ensure_v_prof_cache_table(conn)
        _v_cache_invalidate_if_stale(conn, p, mtime)
        conn.commit()
        row = conn.execute(
            """
            SELECT row_text FROM v_prof_cache
            WHERE file_path = ? AND profession_norm = ? AND file_mtime = ?
            """,
            (p, key, mtime),
        ).fetchone()
        if row:
            return row[0]
    return None


def v_cache_put_row_text(path: Path, profession: str, text: str, parts: list[str] | None = None) -> None:
    p = str(path.resolve())
    mtime = path.stat().st_mtime
    key = _norm_profession_key(profession)
    pj = json.dumps(parts, ensure_ascii=False) if parts is not None else None
    with sqlite3.connect(database_path()) as conn:
        _ensure_v_prof_cache_table(conn)
        conn.execute(
            """
            INSERT OR REPLACE INTO v_prof_cache
            (file_path, profession_norm, file_mtime, row_text, parts_json, parts_schema)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (p, key, mtime, text, pj, V_PROF_PARTS_CACHE_SCHEMA),
        )
        conn.commit()


def v_cache_get_parts_json(path: Path, profession: str) -> list[str] | None:
    if not path.is_file() or not profession.strip():
        return None
    p = str(path.resolve())
    mtime = path.stat().st_mtime
    key = _norm_profession_key(profession)
    with sqlite3.connect(database_path()) as conn:
        _ensure_v_prof_cache_table(conn)
        _v_cache_invalidate_if_stale(conn, p, mtime)
        conn.commit()
        row = conn.execute(
            """
            SELECT parts_json, COALESCE(parts_schema, 0) FROM v_prof_cache
            WHERE file_path = ? AND profession_norm = ? AND file_mtime = ?
            """,
            (p, key, mtime),
        ).fetchone()
        if row and row[0] and int(row[1]) == V_PROF_PARTS_CACHE_SCHEMA:
            try:
                data = json.loads(row[0])
                if isinstance(data, list):
                    return [str(x) for x in data]
            except json.JSONDecodeError:
                return None
    return None


def _v_prof_find_best_row_parts(path: Path, profession: str) -> list[str]:
    """Не пустые ячейки столбцов 5–19 строки V_PROF для должности (столбец A)."""
    row = _v_prof_select_best_row(path, profession)
    if not row:
        return []
    parts: list[str] = []
    lim = min(len(row), V_PROF_PARTS_LAST_COL)
    for idx in range(V_PROF_PARTS_FIRST_COL - 1, lim):
        cell = row[idx]
        if cell is None:
            continue
        t = str(cell).strip()
        if t:
            parts.append(t)
    return parts


def read_v_prof_row_parts_list(path: Path, profession: str) -> list[str]:
    """Упорядоченный список непустых ячеек столбцов 5–19 для должности (A); с кэшем."""
    cached = v_cache_get_parts_json(path, profession)
    if cached is not None:
        return cached
    parts = _v_prof_find_best_row_parts(path, profession)
    joined = ", ".join(parts)
    v_cache_put_row_text(path, profession, joined, parts)
    return parts


def read_v_prof_row_joined(path: Path, profession: str) -> str:
    """Сцепка ячеек V_PROF через запятую (для совместимости)."""
    return ", ".join(read_v_prof_row_parts_list(path, profession))


def expand_persons_for_separate_profession_rows(persons: list[EmployeeRecord]) -> list[EmployeeRecord]:
    """
    Устаревший разворот: две строки на работника с совмещением.
    Таблица протокола в .docx теперь строится по исходным записям (одна строка на ФИО);
    функция оставлена для совместимости и вспомогательного использования.
    """
    out: list[EmployeeRecord] = []
    for p in persons:
        out.append(
            EmployeeRecord(
                fio=p.fio,
                profession=p.profession,
                subdivision=p.subdivision,
                profession2="",
            )
        )
        p2 = (p.profession2 or "").strip()
        if not p2:
            continue
        out.append(
            EmployeeRecord(
                fio=p.fio,
                profession=p2,
                subdivision=p.subdivision,
                profession2="",
            )
        )
    return out


def v_program_merged_parts_for_raw_employee(path: Path, emp: EmployeeRecord) -> list[str]:
    """
    Все фрагменты V_PROF для сотрудника из исходной записи Excel: основная + совмещаемая
    должность без повторов по тексту (порядок: сначала основная строка V_PROF, затем уникальное из второй).
    """
    primary = read_v_prof_row_parts_list(path, emp.profession)
    seen = {_norm_profession_key(x) for x in primary}
    out = list(primary)
    p2 = (emp.profession2 or "").strip()
    if not p2:
        return out
    for x in read_v_prof_row_parts_list(path, p2):
        nk = _norm_profession_key(x)
        if nk not in seen:
            seen.add(nk)
            out.append(x)
    return out


def resolve_v_program_inner_text_global(
    path: Path, persons_raw: list[EmployeeRecord], fallback: str
) -> str:
    """Один объединённый текст фрагментов В по всем выбранным сотрудникам (без повторов)."""
    seen: set[str] = set()
    ordered: list[str] = []
    for emp in persons_raw:
        for part in v_program_merged_parts_for_raw_employee(path, emp):
            nk = _norm_profession_key(part)
            if nk not in seen:
                seen.add(nk)
                ordered.append(part)
    return ", ".join(ordered) if ordered else fallback


def format_v_program_table_block_title(inner: str, fallback: str) -> str:
    """Строка заголовка блока В в таблице: «Программа (В)» и объединённый текст в скобках."""
    core = inner.strip() if inner.strip() else fallback
    return f"Программа (В)\n({core})"


def _format_v_result_cell(grade: str, registry_no: str, v_parts: list[str]) -> str:
    """
    Результат для блока В: оценка, затем по строке «№ …» на каждый фрагмент программы
    (основная + совмещаемая должность); номера реестра по порядку. Без фрагментов — как обычная ячейка.
    """
    if not v_parts:
        return _format_table_result_grade(grade, registry_no)
    return _format_n_registry_lines(grade, registry_no, len(v_parts))


def _protocol_program_fallback_title(program_key: str) -> str:
    for k, _, fb in PROTOCOL_PROGRAM_DEFS:
        if k == program_key:
            return fb
    return ""


def _profession_cell_primary_only(emp: EmployeeRecord) -> str:
    """Только основная должность (совмещение в таблице не дублируется второй строкой)."""
    return (emp.profession or "").strip()


def _rebuild_registry_rows_for_program(pkey: str, emp: EmployeeRecord) -> int:
    """Сколько строк «№» в ячейке результата для блока программы."""
    pk = (pkey or "").strip().upper()
    has2 = bool((emp.profession2 or "").strip())
    if pk == "B":
        return 2 if has2 else 1
    if pk in ("PP", "SIZ"):
        return 1
    return 1


def init_db() -> None:
    with sqlite3.connect(database_path()) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS protocols (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                fio TEXT NOT NULL,
                topic TEXT,
                date TEXT,
                grade TEXT,
                content TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        _ensure_v_prof_cache_table(conn)
        ensure_app_settings_table(conn)
        conn.commit()


def save_protocol(fio: str, topic: str, date: str, grade: str, content: str) -> int:
    with sqlite3.connect(database_path()) as conn:
        cur = conn.execute(
            """
            INSERT INTO protocols (fio, topic, date, grade, content)
            VALUES (?, ?, ?, ?, ?)
            """,
            (fio, topic, date, grade, content),
        )
        conn.commit()
        return int(cur.lastrowid)


def get_all_protocols() -> list[dict[str, Any]]:
    with sqlite3.connect(database_path()) as conn:
        conn.row_factory = sqlite3.Row
        cur = conn.execute(
            """
            SELECT id, fio, topic, date, grade, content, created_at
            FROM protocols
            ORDER BY id DESC
            """
        )
        return [dict(row) for row in cur.fetchall()]


def clear_protocol_journal() -> int:
    """Удаляет все строки из таблицы журнала protocols. Кэш v_prof_cache не трогает."""
    with sqlite3.connect(database_path()) as conn:
        conn.execute("DELETE FROM protocols")
        ch = conn.execute("SELECT changes()").fetchone()
        conn.commit()
        return int(ch[0]) if ch and ch[0] is not None else 0


def _load_form_lines_from_txt(path: Path) -> list[str]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as e:
        raise ProtocolTemplateError(
            f"Текстовый шаблон должен быть в кодировке UTF-8:\n{path}"
        ) from e
    lines = text.splitlines()
    try:
        return _slice_form_template(lines)
    except ValueError:
        return lines


def load_protocol_form_lines(path: Path | None = None) -> list[str]:
    p = protocol_template_path() if path is None else Path(path).expanduser().resolve()
    if not p.is_file():
        hint = (
            f"Положите файл «{PROTOCOL_TEMPLATE_FILENAME}» в папку с программой или выберите шаблон."
            if path is None
            else ""
        )
        raise ProtocolTemplateError(
            f"Файл шаблона не найден:\n{p}" + (f"\n\n{hint}" if hint else "")
        )
    suf = p.suffix.lower()
    if suf == ".txt":
        try:
            return _load_form_lines_from_txt(p)
        except OSError as e:
            raise ProtocolTemplateError(f"Не удалось прочитать шаблон:\n{e}") from e
    if suf in (".docx", ".docm"):
        try:
            return _load_form_lines_from_docx(p)
        except zipfile.BadZipFile as e:
            raise ProtocolTemplateError(
                f"Файл не является корректным документом Word (.docx):\n{p}"
            ) from e
        except ET.ParseError as e:
            raise ProtocolTemplateError(f"Не удалось разобрать XML в шаблоне:\n{p}") from e
        except ValueError as e:
            raise ProtocolTemplateError(str(e)) from e
        except OSError as e:
            raise ProtocolTemplateError(f"Не удалось прочитать шаблон:\n{e}") from e
    raise ProtocolTemplateError(
        f"Поддерживаются шаблоны .docx / .docm и .txt; указан файл: {p.suffix or '(без расширения)'}"
    )


def _paragraphs_from_docx(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        lines: list[str] = []
        for name in _sorted_word_hf_xml_names(zf, "header"):
            lines.extend(_paragraph_texts_from_hf_or_fragment_xml(zf.read(name)))
        root = ET.fromstring(zf.read("word/document.xml"))
        body = root.find(f".//{_W_NS}body")
        if body is None:
            raise ValueError("В документе Word не найден раздел body.")
        for p in body.iter(f"{_W_NS}p"):
            lines.append(_w_p_element_plain_text(p))
        for name in _sorted_word_hf_xml_names(zf, "footer"):
            lines.extend(_paragraph_texts_from_hf_or_fragment_xml(zf.read(name)))
    return lines


def _find_form_template_bounds(lines: list[str]) -> tuple[int, int]:
    start = None
    for i, line in enumerate(lines):
        if _line_has_protocol_form_marker(line):
            start = i
            break
    if start is None:
        for i in range(len(lines) - 1):
            for glue in ("", " ", "\n"):
                combo = glue.join((lines[i], lines[i + 1]))
                if _line_has_protocol_form_marker(combo):
                    start = i
                    break
            if start is not None:
                break
    if start is None:
        raise ValueError(
            "В файле шаблона не найден бланк: в тексте, таблице или колонтитуле должна быть строка "
            "с фрагментом «ПРОТОКОЛ №» (допускаются латинский # или знак № U+2116; любой регистр). "
            "Номер подставляется после знака или вместо подчёркиваний сразу после него. "
            "Пример шаблона: default_protocol.docx в папке программы."
        )
    end = len(lines)
    for j in range(start, len(lines)):
        if lines[j].startswith("Приложение"):
            end = j
            break
    return start, end


def _slice_form_template(lines: list[str]) -> list[str]:
    start, end = _find_form_template_bounds(lines)
    return lines[start:end]


def _load_form_lines_from_docx(path: Path) -> list[str]:
    raw = _paragraphs_from_docx(path)
    try:
        return _slice_form_template(raw)
    except ValueError as e:
        raise ValueError(f"{e}\n\nПроверяемый файл:\n{path.resolve()}") from e


def _split_dmy_triplet(date_str: str) -> tuple[str, str, str] | None:
    """День, месяц, год как строки; разделитель . / или - (порядок ДД.ММ.ГГГГ)."""
    s = date_str.strip()
    for sep in (".", "/", "-"):
        if sep not in s:
            continue
        parts = [p.strip() for p in s.split(sep)]
        if len(parts) == 3:
            return (parts[0], parts[1], parts[2])
    return None


def _format_date_protocol_line(date_str: str) -> str:
    sp = _split_dmy_triplet(date_str)
    if sp is None:
        return date_str
    d_s, m_s, y = sp
    try:
        d, m = int(d_s), int(m_s)
        if len(y) == 2:
            y = f"20{y}"
        if not (1 <= m <= 12):
            return date_str
        return f"«{d}» {_MONTHS_GEN[m]} {y} г."
    except ValueError:
        return date_str


def _parse_dmy_month_year(date_str: str) -> tuple[int, int] | None:
    """ДД.ММ.ГГГГ или ДД.ММ.ГГ (разделитель . / -) → (месяц 1–12, полный год). Некорректная дата — None."""
    sp = _split_dmy_triplet(date_str)
    if sp is None:
        return None
    d_s, m_s, y = sp
    try:
        d, m = int(d_s), int(m_s)
        if len(y) == 2:
            y_full = int(f"20{y}")
        else:
            y_full = int(y)
        if not (1 <= m <= 12 and 1 <= d <= 31):
            return None
        return m, y_full
    except ValueError:
        return None


def format_protocol_number_for_template(protocol_no: str, date_str: str) -> str:
    """
    Подстановка вместо ____ в строке «ПРОТОКОЛ №»: <номер>-<месяц>-<год_2_цифры>.
    Месяц без ведущего нуля. Пустой номер — пустая строка (не подставлять).
    """
    raw = (protocol_no or "").strip()
    if not raw:
        return ""
    my = _parse_dmy_month_year(date_str)
    if my is None:
        return raw
    month, year_full = my
    yy = year_full % 100
    return f"{raw}-{month}-{yy:02d}"


def _is_program_underscore_line(line: str) -> bool:
    s = line.strip()
    if not s.endswith(","):
        return False
    core = s[:-1].replace("_", "").strip()
    return len(s) > 20 and not core


def _fill_protocol_form(
    form: list[str],
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
) -> list[str]:
    """Подстановка полей по маркерам шаблона. Список маркеров — PROTOCOL_TEMPLATE_VARIABLES_DOC."""
    payload = build_commission_template_payload(_format_date_protocol_line)
    result: list[str] = []
    idx = 0
    while idx < len(form):
        line = form[idx]
        nxt = form[idx + 1] if idx + 1 < len(form) else None
        chunks: list[str]

        if (
            nxt is not None
            and _is_protocol_title_line_without_underscores(line)
            and _is_placeholder_only_underscore_line(nxt)
        ):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            chunks = [
                line,
                _replace_first_underscore_run_with_protocol_number(nxt, pn) if pn else nxt,
            ]
            idx += 2
        elif _line_fills_protocol_number_slot(line):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            chunks = [_apply_protocol_number_after_anchor(line, pn) if pn else line]
            idx += 1
        elif line.strip().startswith("«__»") and line.strip().endswith("г."):
            chunks = [_format_date_protocol_line(date_str)]
            idx += 1
        elif _is_program_underscore_line(line):
            chunks = [(theme.strip() + ",") if theme.strip() else line]
            idx += 1
        else:
            chunks = [line]
            idx += 1

        for out_line in chunks:
            result.append(
                apply_commission_insertions_to_line(
                    out_line,
                    date_words=payload["date_words"],
                    order_no=payload["order_no"],
                    chair_gen=payload["chair"],
                    members_gen=payload["members"],
                )
            )
    return result


def build_protocol_text(
    theme: str,
    date_str: str,
    protocol_no: str = "",
    template_path: Path | None = None,
) -> str:
    form = load_protocol_form_lines(template_path)
    filled = _fill_protocol_form(
        form,
        protocol_no=protocol_no,
        date_str=date_str,
        theme=theme,
    )
    return "\n".join(filled)


def _configure_fpdf_font(pdf: FPDF, content: str) -> None:
    """Стандартный Helvetica для Latin-1; для кириллицы — системный Arial (TTF, uni)."""
    try:
        content.encode("latin-1")
        pdf.set_font("Helvetica", "", 12)
    except UnicodeEncodeError:
        ttf = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "arial.ttf"
        if not ttf.is_file():
            raise RuntimeError(
                "Стандартные шрифты PDF не поддерживают кириллицу; "
                "не найден arial.ttf в папке Fonts Windows."
            ) from None
        pdf.add_font("ArialUnicode", "", str(ttf), uni=True)
        pdf.set_font("ArialUnicode", "", 12)


def write_protocol_pdf(path: str, content: str) -> None:
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _configure_fpdf_font(pdf, content)
    normalized = content.replace("\r\n", "\n").replace("\r", "\n")
    for line in normalized.split("\n"):
        pdf.multi_cell(0, 7, line)
    pdf.output(path, "F")


def write_protocol_pdf_from_docx_template(
    template_path: Path,
    pdf_path: str,
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
    table_persons: list[EmployeeRecord],
    program_titles: list[str] | None = None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons_v_raw: list[EmployeeRecord] | None = None,
    grade: str = "",
    registry_no: str = "",
    check_type: str = "плановая",
) -> None:
    """Собирает DOCX из шаблона и конвертирует в PDF через Word (сохраняет оформление)."""
    try:
        from docx2pdf import convert
    except ImportError as e:
        raise RuntimeError(
            "Для PDF с форматированием Word установите: pip install docx2pdf\n"
            "и приложение Microsoft Word."
        ) from e

    fd, tmp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc, _ = build_filled_protocol_document(
            template_path,
            protocol_no=protocol_no,
            date_str=date_str,
            theme=theme,
            table_persons=table_persons,
            program_titles=program_titles,
            program_keys=program_keys,
            excel_path=excel_path,
            persons_v_raw=persons_v_raw,
            grade=grade,
            registry_no=registry_no,
            check_type=check_type,
        )
        doc.save(tmp_docx)
        convert(
            str(Path(tmp_docx).resolve()),
            str(Path(pdf_path).resolve()),
        )
    finally:
        try:
            os.unlink(tmp_docx)
        except OSError:
            pass


def write_protocol_docx(path: str, content: str) -> None:
    doc = Document()
    normalized = content.replace("\r\n", "\n").replace("\r", "\n")
    for line in normalized.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)


def _replace_paragraph_text_preserve_style(paragraph: DocxParagraph, new_text: str) -> None:
    """Заменяет текст абзаца: стиль абзаца сохраняется (clear), оформление знаков — с первого run."""
    r0 = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    new_run = paragraph.add_run(new_text)
    if r0 is not None and r0._r.rPr is not None:
        new_run._r.insert(0, deepcopy(r0._r.rPr))


def _replace_paragraph_text_preserve_style_multiline(paragraph: DocxParagraph, new_text: str) -> None:
    """Как _replace_paragraph_text_preserve_style, но переводы строк — как мягкий перенос в Word."""
    r0 = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    parts = (new_text or "").split("\n")
    for i, part in enumerate(parts):
        if i:
            br_run = paragraph.add_run()
            br_run.add_break(WD_BREAK.LINE)
            if r0 is not None and r0._r.rPr is not None:
                br_run._r.insert(0, deepcopy(r0._r.rPr))
        new_run = paragraph.add_run(part)
        if r0 is not None and r0._r.rPr is not None:
            new_run._r.insert(0, deepcopy(r0._r.rPr))


def _iter_all_paragraphs_in_document(doc: Document):
    """Все абзацы документа, включая ячейки таблиц (объединённые ячейки не дублируются)."""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for p in cell.paragraphs:
                    yield p


def apply_protocol_signature_placeholders_in_template(
    doc: Document,
    *,
    chair_text: str,
    members_text: str,
) -> None:
    """
    В основном шаблоне протокола заменяет плейсхолдеры подписей комиссии (формат И.О. Фамилия — см. БД).
    Плейсхолдеры: {{ПРЕДСЕДАТЕЛЬ}}, {{ЧЛЕНЫ_КОМИССИИ}}, {{CHAIR}}, {{MEMBERS}}.
    """
    repl = (
        ("{{ПРЕДСЕДАТЕЛЬ}}", chair_text or ""),
        ("{{ЧЛЕНЫ_КОМИССИИ}}", members_text or ""),
        ("{{CHAIR}}", chair_text or ""),
        ("{{MEMBERS}}", members_text or ""),
    )
    for para in _iter_all_paragraphs_in_document(doc):
        t = para.text
        if not t:
            continue
        new_t = t
        for key, val in repl:
            if key in new_t:
                new_t = new_t.replace(key, val)
        if new_t == t:
            continue
        if "\n" in new_t:
            _replace_paragraph_text_preserve_style_multiline(para, new_t)
        else:
            _replace_paragraph_text_preserve_style(para, new_t)


def _protocol_results_table_header_match(table: Table) -> bool:
    """Таблица бланка: заголовок с «п/п» и «Фамилия», не меньше 7 колонок."""
    if not table.rows or len(table.rows[0].cells) < 7:
        return False
    joined = " ".join(c.text for c in table.rows[0].cells).lower().replace("\xa0", " ")
    return "п/п" in joined and "фамилия" in joined


def _normalize_fio_marker(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ").replace("\xa0", " ")).strip().upper()


def _format_registry_number_token(token: str) -> str:
    t = token.strip()
    if not t:
        return ""
    if re.match(r"^№\s*", t):
        return t
    return f"№ {t}"


def _format_n_registry_lines(grade: str, registry_no: str, n_registry_rows: int) -> str:
    """
    Оценка и ровно n_registry_rows строк «№ …» (как в блоке «В»); лишние номера из поля не выводятся.
    При нехватке номеров — строка «№» без текста.
    """
    if n_registry_rows <= 0:
        return _format_table_result_grade(grade, registry_no)
    g = (grade or "").strip()
    titled = g[:1].upper() + g[1:] if len(g) > 1 else (g.upper() if g else "")
    regs = [r.strip() for r in re.split(r"[\n,;]+", (registry_no or "").strip()) if r.strip()]
    lines: list[str] = []
    if titled:
        lines.append(f"{titled},")
    for i in range(n_registry_rows):
        r = regs[i] if i < len(regs) else ""
        lines.append(_format_registry_number_token(r) if r else "№")
    return "\n".join(lines)


def _format_table_result_grade(grade: str, registry_no: str) -> str:
    """Оценка + регистрационные номера реестра Минтруда (не номер протокола)."""
    g = (grade or "").strip()
    if not g:
        return ""
    titled = g[:1].upper() + g[1:] if len(g) > 1 else g.upper()
    raw = (registry_no or "").strip()
    if not raw:
        return f"{titled},"
    parts = [p.strip() for p in re.split(r"[\n,;]+", raw) if p.strip()]
    lines = [f"{titled},"]
    lines.extend(_format_registry_number_token(p) for p in parts)
    return "\n".join(lines)


def _find_fio_marker_row_metas(table: Table) -> list[tuple[int, int, int]]:
    """
    Индексы строк таблицы с маркером «ФИО» и номером вида «целое.целое» в 1-й колонке.
    Возвращает список (индекс_строки, major, minor), например (3, 1, 1) для «1.1».
    """
    out: list[tuple[int, int, int]] = []
    for i, row in enumerate(table.rows):
        if len(row.cells) < 7:
            continue
        c0 = row.cells[0].text.strip()
        c1 = row.cells[1].text
        m = re.match(r"^(\d+)\.(\d+)$", c0)
        if not m:
            continue
        if _normalize_fio_marker(c1) != "ФИО":
            continue
        out.append((i, int(m.group(1)), int(m.group(2))))
    return out


def _insert_duplicate_tr_after(template_tr: Any, extra_copies: int) -> None:
    """Вставляет подряд extra_copies копий строки таблицы (XML w:tr) сразу после template_tr."""
    if extra_copies <= 0:
        return
    last = template_tr
    for _ in range(extra_copies):
        new_tr = deepcopy(template_tr)
        last.addnext(new_tr)
        last = new_tr


def _table_tr_elements(tbl: Any) -> list[Any]:
    tag_tr = qn("w:tr")
    return [e for e in tbl if e.tag == tag_tr]


def _fill_static_protocol_result_table(
    doc: Document,
    *,
    persons: list[EmployeeRecord],
    grade: str,
    registry_no: str,
    check_type: str,
) -> tuple[int, int]:
    """
    Старый режим: в шаблоне уже есть блоки 1.1, 2.1, … — только дублируем строки
    под число сотрудников.
    """
    if not persons:
        return 0, 0

    chk = (check_type or "плановая").strip()
    if chk:
        chk = chk[:1].upper() + chk[1:]
    result_text = _format_table_result_grade(grade, registry_no)
    n = len(persons)

    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        markers = _find_fio_marker_row_metas(table)
        if not markers:
            continue
        to_expand = [(idx, major, minor) for idx, major, minor in markers if minor == 1]
        if not to_expand:
            to_expand = list(markers)
        for row_idx, major, _minor in sorted(to_expand, key=lambda x: -x[0]):
            template_tr = table.rows[row_idx]._tr
            _insert_duplicate_tr_after(template_tr, n - 1)
            for j in range(n):
                row = table.rows[row_idx + j]
                row.cells[0].text = f"{major}.{j + 1}"
                p = persons[j]
                row.cells[1].text = p.fio
                row.cells[2].text = p.profession
                row.cells[3].text = p.subdivision
                row.cells[4].text = result_text
                row.cells[5].text = chk
        return n, 0

    return 0, n


def _profession_cell_text_for_table(emp: EmployeeRecord) -> str:
    """Должность в строке таблицы: основная и совмещаемая в одной ячейке."""
    p2 = (emp.profession2 or "").strip()
    if p2:
        return f"{emp.profession}; {p2}" if (emp.profession or "").strip() else p2
    return emp.profession


def _rebuild_protocol_result_table(
    doc: Document,
    *,
    program_titles: list[str],
    program_keys: list[str] | None,
    excel_path: Path | None,
    persons: list[EmployeeRecord],
    persons_v_raw: list[EmployeeRecord] | None,
    grade: str,
    registry_no: str,
    check_type: str,
    persons_b_all_rows: list[EmployeeRecord] | None = None,
) -> tuple[int, int]:
    """
    Удаляет строки таблицы ниже шапки (две первые строки-заголовка) и строит заново:
    для каждой выбранной программы — строка с номером и полным названием из Excel,
    затем строки сотрудников N.1…N.M. Для «В» перед каждой строкой работника — ещё строка-шапка
    с перечнем тем этого работника; для ПП и СИЗ — по объединённому списку ФИО;
    для «Б» — по persons_b_all_rows (все исходные строки, даже с одинаковым ФИО), если задан.
    """
    merged_rows = persons_v_raw if persons_v_raw is not None else persons
    b_rows = persons_b_all_rows if persons_b_all_rows is not None else merged_rows
    if not merged_rows:
        return 0, 0
    if not program_titles:
        return 0, len(merged_rows)

    chk = (check_type or "плановая").strip()
    if chk:
        chk = chk[:1].upper() + chk[1:]
    n_tab = len(merged_rows)
    keys = program_keys if program_keys and len(program_keys) == len(program_titles) else [""] * len(
        program_titles
    )
    v_path = excel_path if excel_path and excel_path.is_file() else None

    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        tbl = table._tbl
        trs = _table_tr_elements(tbl)
        header_rows = 2
        if len(trs) < header_rows + 2:
            return 0, n_tab
        prog_tpl = deepcopy(trs[header_rows])
        emp_tpl = deepcopy(trs[header_rows + 1])
        for tr in trs[header_rows:]:
            tbl.remove(tr)

        for pi, (pkey, title) in enumerate(zip(keys, program_titles), start=1):
            pk = (pkey or "").strip().upper()
            is_v = pk == "V" and v_path is not None
            is_b = pk == "B"
            block_rows = b_rows if is_b else merged_rows
            n_block = len(block_rows)

            if is_v:
                v_fb = _protocol_program_fallback_title("V")
                for j in range(n_block):
                    p = block_rows[j]
                    v_parts = v_program_merged_parts_for_raw_employee(v_path, p)
                    inner_j = ", ".join(v_parts) if v_parts else ""
                    sub_title = format_v_program_table_block_title(inner_j, v_fb)
                    p_tr = deepcopy(prog_tpl)
                    tbl.append(p_tr)
                    row_p = table.rows[-1]
                    row_p.cells[0].text = str(pi) if j == 0 else ""
                    for ci in range(1, min(7, len(row_p.cells))):
                        row_p.cells[ci].text = sub_title
                    e_tr = deepcopy(emp_tpl)
                    tbl.append(e_tr)
                    row = table.rows[-1]
                    row.cells[0].text = f"{pi}.{j + 1}"
                    row.cells[1].text = p.fio
                    row.cells[2].text = _profession_cell_primary_only(p)
                    row.cells[3].text = p.subdivision
                    row.cells[4].text = _format_v_result_cell(grade, registry_no, v_parts)
                    row.cells[5].text = chk
            else:
                p_tr = deepcopy(prog_tpl)
                tbl.append(p_tr)
                row_p = table.rows[-1]
                row_p.cells[0].text = str(pi)
                for ci in range(1, min(7, len(row_p.cells))):
                    row_p.cells[ci].text = title

                e_tr = deepcopy(emp_tpl)
                tbl.append(e_tr)
                emp_base_idx = len(table.rows) - 1
                emp_base_row = table.rows[emp_base_idx]
                _insert_duplicate_tr_after(emp_base_row._tr, max(n_block - 1, 0))
                for j in range(n_block):
                    row = table.rows[emp_base_idx + j]
                    p = block_rows[j]
                    row.cells[0].text = f"{pi}.{j + 1}"
                    row.cells[1].text = p.fio
                    row.cells[2].text = _profession_cell_primary_only(p)
                    row.cells[3].text = p.subdivision
                    n_reg = _rebuild_registry_rows_for_program(pkey or "", p)
                    row.cells[4].text = _format_n_registry_lines(grade, registry_no, n_reg)
                    row.cells[5].text = chk

        return n_tab, 0

    return 0, n_tab


def fill_protocol_result_table(
    doc: Document,
    *,
    program_titles: list[str] | None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons: list[EmployeeRecord],
    persons_v_raw: list[EmployeeRecord] | None = None,
    grade: str,
    registry_no: str,
    check_type: str,
    persons_b_all_rows: list[EmployeeRecord] | None = None,
) -> tuple[int, int]:
    if program_titles:
        return _rebuild_protocol_result_table(
            doc,
            program_titles=program_titles,
            program_keys=program_keys,
            excel_path=excel_path,
            persons=persons,
            persons_v_raw=persons_v_raw,
            grade=grade,
            registry_no=registry_no,
            check_type=check_type,
            persons_b_all_rows=persons_b_all_rows,
        )
    return _fill_static_protocol_result_table(
        doc,
        persons=persons,
        grade=grade,
        registry_no=registry_no,
        check_type=check_type,
    )


def build_filled_protocol_document(
    template_path: Path,
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
    table_persons: list[EmployeeRecord],
    program_titles: list[str] | None = None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons_v_raw: list[EmployeeRecord] | None = None,
    grade: str = "",
    registry_no: str = "",
    check_type: str = "плановая",
) -> tuple[Document, int]:
    """
    Загружает шаблон .docx, подставляет поля, возвращает (документ, «остаток»).
    program_titles — названия: «Б» с листа B, остальное с V_PROF; пересборка таблицы.
    """
    doc = Document(str(template_path))
    emp_source = persons_v_raw if persons_v_raw is not None else table_persons
    emp_all = list(emp_source)
    emp_for_doc = _table_employees_dedupe_by_fio(emp_all)
    paras = _all_document_paragraphs_ordered(doc)
    lines = [p.text for p in paras]
    start, end = _find_form_template_bounds(lines)

    if program_titles:
        paragraph_theme = "; ".join(program_titles).strip()
    else:
        paragraph_theme = (theme or "").strip()

    commission_payload = build_commission_template_payload(_format_date_protocol_line)
    i = start
    while i < end:
        para = paras[i]
        line = lines[i]
        nxt_line = lines[i + 1] if i + 1 < end else None

        if (
            nxt_line is not None
            and _is_protocol_title_line_without_underscores(line)
            and _is_placeholder_only_underscore_line(nxt_line)
        ):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            t1 = line
            t2 = (
                _replace_first_underscore_run_with_protocol_number(nxt_line, pn)
                if pn
                else nxt_line
            )
            t1 = apply_commission_insertions_to_line(
                t1,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            t2 = apply_commission_insertions_to_line(
                t2,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            if t1 != line:
                _replace_paragraph_text_preserve_style(para, t1)
            p2 = paras[i + 1]
            if t2 != nxt_line:
                _replace_paragraph_text_preserve_style(p2, t2)
            i += 2
        elif _line_fills_protocol_number_slot(line):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            t = _apply_protocol_number_after_anchor(line, pn) if pn else line
            t = apply_commission_insertions_to_line(
                t,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            if t != line:
                _replace_paragraph_text_preserve_style(para, t)
            i += 1
        else:
            t = line
            if line.strip().startswith("«__»") and line.strip().endswith("г."):
                t = _format_date_protocol_line(date_str)
            elif _is_program_underscore_line(line) and paragraph_theme:
                t = paragraph_theme + ","
            t = apply_commission_insertions_to_line(
                t,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            if t != line:
                _replace_paragraph_text_preserve_style(para, t)
            i += 1

    if excel_path and excel_path.is_file() and program_keys and emp_for_doc:
        fg_lines = build_fg_lines_for_selected_programs(
            excel_path, program_keys, emp_for_doc
        )
        insert_program_fg_lines_after_anchor(doc, fg_lines)

    _, excess = fill_protocol_result_table(
        doc,
        program_titles=program_titles if program_titles else None,
        program_keys=program_keys,
        excel_path=excel_path,
        persons=emp_for_doc,
        persons_v_raw=emp_for_doc,
        grade=grade,
        registry_no=registry_no,
        check_type=check_type,
        persons_b_all_rows=emp_all,
    )

    lines_after = [p.text for p in _all_document_paragraphs_ordered(doc)]
    start_body, end_body = _find_form_template_bounds(lines_after)
    _apply_protocol_body_font_pt(doc, start_body, end_body, PROTOCOL_BODY_FONT_PT)

    chair_sig, members_sig = build_commission_signature_suffix_payload()
    apply_protocol_signature_placeholders_in_template(
        doc,
        chair_text=chair_sig,
        members_text=members_sig,
    )

    return doc, excess


def save_protocol_docx_from_template(
    template_path: Path,
    output_path: str,
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
    table_persons: list[EmployeeRecord],
    program_titles: list[str] | None = None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons_v_raw: list[EmployeeRecord] | None = None,
    grade: str = "",
    registry_no: str = "",
    check_type: str = "плановая",
) -> None:
    doc, _ = build_filled_protocol_document(
        template_path,
        protocol_no=protocol_no,
        date_str=date_str,
        theme=theme,
        table_persons=table_persons,
        program_titles=program_titles,
        program_keys=program_keys,
        excel_path=excel_path,
        persons_v_raw=persons_v_raw,
        grade=grade,
        registry_no=registry_no,
        check_type=check_type,
    )
    doc.save(output_path)


def _iter_paragraph_runs(paragraph: DocxParagraph):
    for item in paragraph.iter_inner_content():
        if isinstance(item, DocxRun):
            yield item
        elif isinstance(item, Hyperlink):
            for r in item.runs:
                yield r


def _apply_font_pt_to_paragraph_runs(para: DocxParagraph, pt: int) -> None:
    for run in _iter_paragraph_runs(para):
        run.font.size = Pt(pt)


def _is_v_program_table_header_row(row) -> bool:
    """Строка-шапка блока «В» в таблице (до строки сотрудника)."""
    cells = row.cells
    if len(cells) < 2:
        return False
    return (cells[1].text or "").lstrip().startswith("Программа (В)")


def _apply_protocol_body_font_pt(doc: Document, form_start: int, form_end: int, pt: int) -> None:
    """
    Абзацы бланка [form_start, form_end) — без организационной шапки над «ПРОТОКОЛ №»;
    в таблице «Результат проверки» — все строки ниже двух верхних строк шапки.
    """
    paras = _all_document_paragraphs_ordered(doc)
    last = min(form_end, len(paras))
    for i in range(form_start, last):
        _apply_font_pt_to_paragraph_runs(paras[i], pt)

    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        header_rows = 2
        for ri, row in enumerate(table.rows):
            if ri < header_rows:
                continue
            row_pt = (
                PROTOCOL_V_TABLE_HEADER_FONT_PT
                if _is_v_program_table_header_row(row)
                else pt
            )
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for para in cell.paragraphs:
                    _apply_font_pt_to_paragraph_runs(para, row_pt)


def document_to_plain_text(doc: Document) -> str:
    lines = [p.text for p in doc.paragraphs]
    if doc.tables:
        lines.append("")
        lines.append("── Таблица (текст ячеек) ──")
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                if not any(cells):
                    continue
                lines.append("  |  ".join(cells))
    return "\n".join(lines)


def render_document_to_text_widget(widget: tk.Text, doc: Document) -> None:
    """Выводит документ Word в tk.Text с базовым сохранением жирного/курсива по run."""
    widget.configure(state=tk.NORMAL)
    widget.delete("1.0", tk.END)
    for tag in (
        "pv_bold",
        "pv_italic",
        "pv_bi",
        "pv_body",
        "pv_body_bold",
        "pv_body_italic",
        "pv_body_bi",
    ):
        try:
            widget.tag_delete(tag)
        except tk.TclError:
            pass
    widget.tag_configure("pv_bold", font=(*PROTOCOL_PREVIEW_HEADER_FONT, "bold"))
    widget.tag_configure("pv_italic", font=(*PROTOCOL_PREVIEW_HEADER_FONT, "italic"))
    widget.tag_configure("pv_bi", font=(*PROTOCOL_PREVIEW_HEADER_FONT, "bold italic"))
    widget.tag_configure("pv_body", font=PROTOCOL_PREVIEW_BODY_FONT)
    widget.tag_configure("pv_body_bold", font=(*PROTOCOL_PREVIEW_BODY_FONT, "bold"))
    widget.tag_configure("pv_body_italic", font=(*PROTOCOL_PREVIEW_BODY_FONT, "italic"))
    widget.tag_configure("pv_body_bi", font=(*PROTOCOL_PREVIEW_BODY_FONT, "bold italic"))

    ordered = _all_document_paragraphs_ordered(doc)
    plines = [p.text for p in ordered]
    try:
        body_start, body_end = _find_form_template_bounds(plines)
    except ValueError:
        body_start, body_end = 0, len(plines)

    for pi, para in enumerate(ordered):
        in_body = body_start <= pi < body_end
        for run in _iter_paragraph_runs(para):
            chunk = run.text
            if not chunk:
                continue
            b = bool(run.bold)
            i = bool(run.italic)
            if in_body:
                if b and i:
                    tags: tuple[str, ...] = ("pv_body_bi",)
                elif b:
                    tags = ("pv_body_bold",)
                elif i:
                    tags = ("pv_body_italic",)
                else:
                    tags = ("pv_body",)
            else:
                if b and i:
                    tags = ("pv_bi",)
                elif b:
                    tags = ("pv_bold",)
                elif i:
                    tags = ("pv_italic",)
                else:
                    tags = ()
            widget.insert(tk.END, chunk, tags)
        widget.insert(tk.END, "\n")

    if doc.tables:
        widget.insert(tk.END, "\n── Таблица (текст ячеек) ──\n", ("pv_body",))
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                if not any(cells):
                    continue
                widget.insert(tk.END, "  |  ".join(cells) + "\n", ("pv_body",))


class ProtocolApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.template_path: Path | None = None
        self.employees_excel_path: Path | None = None
        self._employee_records: list[EmployeeRecord] = []
        self._employee_filter_indices: list[int] = []
        self.var_emp_search = tk.StringVar(value="")
        self._prog_vars: dict[str, tk.BooleanVar] = {
            key: tk.BooleanVar(value=False) for key, _, _ in PROTOCOL_PROGRAM_DEFS
        }
        self._admin_win: tk.Toplevel | None = None
        self._commission_state = CommissionState()
        self._commission_win: tk.Toplevel | None = None
        self._commission_panel: CommissionAdminPanel | None = None
        self._apply_icon()
        self.title("Протокол проверки знаний")
        self.minsize(480, 560)
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        self._build_menu()
        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_app_quit)
        self._setup_admin_window()
        self._try_autoload_employees()
        self._refresh_employees_file_label()

    def _persist_protocol_no_field(self) -> None:
        save_last_protocol_no(self.entry_protocol_no.get().strip())

    def _on_app_quit(self) -> None:
        self._persist_protocol_no_field()
        self.destroy()

    def _employees_file_resolved(self) -> Path:
        return self.employees_excel_path or employees_excel_default_path()

    def _apply_icon(self) -> None:
        cwd_ico = Path("icon.ico")
        beside_script = Path(__file__).resolve().parent / "icon.ico"
        if cwd_ico.is_file():
            path, use_literal = cwd_ico, True
        elif beside_script.is_file():
            path, use_literal = beside_script, False
        else:
            return
        try:
            self.iconbitmap("icon.ico" if use_literal else str(path.resolve()))
        except (tk.TclError, OSError):
            pass

    def _build_menu(self) -> None:
        mbar = tk.Menu(self)
        self.config(menu=mbar)
        adm = tk.Menu(mbar, tearoff=0)
        mbar.add_cascade(label="Администрирование", menu=adm)
        adm.add_command(label="Настройки и данные…", command=self._open_admin_window)
        adm.add_command(label="Приказ и комиссия…", command=self._open_commission_window)

    def _build_ui(self) -> None:
        g = {"padx": 5, "pady": 5}

        main = ttk.Frame(self, padding=5)
        main.grid(row=0, column=0, sticky=tk.NSEW)
        main.columnconfigure(0, weight=1)
        main.rowconfigure(3, weight=1)

        lf = ttk.Labelframe(main, text="Формирование протокола", padding=5)
        lf.grid(row=0, column=0, sticky=tk.EW, **g)
        lf.columnconfigure(1, weight=1)

        ttk.Label(
            lf,
            text=(
                "Сотрудники (Ctrl/Shift — несколько). Список из Excel; файл, шаблон и комиссия — "
                "меню «Администрирование»."
            ),
            wraplength=460,
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, **g)

        ttk.Label(lf, text="Поиск:").grid(row=1, column=0, sticky=tk.W, **g)
        self.entry_emp_search = ttk.Entry(
            lf,
            textvariable=self.var_emp_search,
            width=50,
        )
        self.entry_emp_search.grid(row=1, column=1, sticky=tk.EW, **g)

        emp_box = ttk.Frame(lf)
        emp_box.grid(row=2, column=0, columnspan=2, sticky=tk.NSEW, **g)
        emp_box.columnconfigure(0, weight=1)
        sb_emp = ttk.Scrollbar(emp_box)
        sb_emp.grid(row=0, column=1, sticky=tk.NS)
        self.list_employees = tk.Listbox(
            emp_box,
            height=7,
            selectmode=tk.EXTENDED,
            exportselection=False,
            font=("Segoe UI", 10),
            yscrollcommand=sb_emp.set,
        )
        self.list_employees.grid(row=0, column=0, sticky=tk.NSEW, **g)
        sb_emp.configure(command=self.list_employees.yview)
        self.list_employees.bind("<<ListboxSelect>>", self._on_employee_list_select)
        self.var_emp_search.trace_add("write", lambda *_: self._refilter_employee_list())

        ttk.Label(lf, text="ФИО вручную (если не из списка):").grid(row=3, column=0, sticky=tk.W, **g)
        self.entry_fio = ttk.Entry(lf, width=50)
        self.entry_fio.grid(row=3, column=1, sticky=tk.EW, **g)

        ttk.Label(lf, text="Должность (при ручном вводе):").grid(row=4, column=0, sticky=tk.W, **g)
        self.entry_position = ttk.Entry(lf, width=50)
        self.entry_position.grid(row=4, column=1, sticky=tk.EW, **g)

        ttk.Label(lf, text="Подразделение (при ручном вводе):").grid(row=5, column=0, sticky=tk.W, **g)
        self.entry_subdivision = ttk.Entry(lf, width=50)
        self.entry_subdivision.grid(row=5, column=1, sticky=tk.EW, **g)

        prog_lf = ttk.Labelframe(lf, text="Программы обучения", padding=4)
        prog_lf.grid(row=6, column=0, columnspan=2, sticky=tk.EW, **g)
        for pi, (key, _sheet, _fb) in enumerate(PROTOCOL_PROGRAM_DEFS):
            ttk.Checkbutton(
                prog_lf,
                text=PROTOCOL_PROGRAM_CHECKBOX_SHORT.get(
                    key, PROTOCOL_PROGRAM_UI_LABELS.get(key, key)
                ),
                variable=self._prog_vars[key],
            ).grid(row=pi // 2, column=pi % 2, sticky=tk.W, padx=4, pady=2)

        ttk.Label(lf, text="Дата:").grid(row=7, column=0, sticky=tk.W, **g)
        self.entry_date = ttk.Entry(lf, width=50)
        self.entry_date.grid(row=7, column=1, sticky=tk.EW, **g)
        self.entry_date.insert(0, date.today().strftime("%d.%m.%Y"))

        ttk.Label(
            lf,
            text="№ протокола (в бланк: номер-месяц-год по полю «Дата»):",
            wraplength=280,
        ).grid(row=8, column=0, sticky=tk.NW, **g)
        self.entry_protocol_no = ttk.Entry(lf, width=50)
        self.entry_protocol_no.grid(row=8, column=1, sticky=tk.EW, **g)
        _saved_protocol_no = load_last_protocol_no()
        if _saved_protocol_no:
            self.entry_protocol_no.insert(0, _saved_protocol_no)

        ttk.Label(lf, text="Оценка:").grid(row=9, column=0, sticky=tk.W, **g)
        self.combo_grade = ttk.Combobox(
            lf,
            values=GRADE_OPTIONS,
            state="readonly",
            width=47,
        )
        self.combo_grade.grid(row=9, column=1, sticky=tk.EW, **g)
        self.combo_grade.current(0)

        ttk.Label(lf, text="Проверка знаний:").grid(row=10, column=0, sticky=tk.W, **g)
        self.combo_check_type = ttk.Combobox(
            lf,
            values=CHECK_TYPE_OPTIONS,
            state="readonly",
            width=47,
        )
        self.combo_check_type.grid(row=10, column=1, sticky=tk.EW, **g)
        self.combo_check_type.current(0)

        gen_row = ttk.Frame(main)
        gen_row.grid(row=1, column=0, sticky=tk.EW, **g)
        gen_row.columnconfigure(0, weight=1)
        ttk.Button(gen_row, text="Сформировать протокол", command=self.generate_protocol).grid(
            row=0, column=0, sticky=tk.EW
        )

        ttk.Label(main, text="Предварительный просмотр:").grid(
            row=2, column=0, sticky=tk.W, **g
        )

        text_frame = ttk.Frame(main)
        text_frame.grid(row=3, column=0, sticky=tk.NSEW, **g)
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)

        self.text_preview = tk.Text(
            text_frame,
            height=16,
            wrap=tk.WORD,
            font=PROTOCOL_PREVIEW_HEADER_FONT,
        )
        self.text_preview.grid(row=0, column=0, sticky=tk.NSEW, **g)
        scroll = ttk.Scrollbar(text_frame, command=self.text_preview.yview)
        scroll.grid(row=0, column=1, sticky=tk.NS, **g)
        self.text_preview.configure(yscrollcommand=scroll.set)

        btn_row = ttk.Frame(main)
        btn_row.grid(row=4, column=0, sticky=tk.EW, **g)
        btn_row.columnconfigure(0, weight=1)
        btn_row.columnconfigure(1, weight=1)

        self.btn_save = ttk.Button(
            btn_row,
            text="Сохранить в DOCX",
            command=self.save_to_docx,
            state="disabled",
        )
        self.btn_save.grid(row=0, column=0, sticky=tk.EW, padx=(0, 3))

        self.btn_save_pdf = ttk.Button(
            btn_row,
            text="Сохранить в PDF",
            command=self.save_to_pdf,
            state="disabled",
        )
        self.btn_save_pdf.grid(row=0, column=1, sticky=tk.EW, padx=(3, 0))

    def _setup_admin_window(self) -> None:
        win = tk.Toplevel(self)
        win.title("Администрирование — настройки и данные")
        win.minsize(520, 420)
        win.withdraw()
        win.transient(self)
        self._admin_win = win
        win.protocol("WM_DELETE_WINDOW", self._close_admin_window)
        self._build_admin_window_content(win)

    def _open_admin_window(self) -> None:
        if self._admin_win is None:
            return
        self._admin_win.deiconify()
        self._admin_win.lift()
        self._admin_win.focus_force()

    def _close_admin_window(self) -> None:
        if self._admin_win is not None:
            self._admin_win.withdraw()

    def _open_commission_window(self) -> None:
        if self._commission_win is not None and self._commission_win.winfo_exists():
            self._commission_win.deiconify()
            self._commission_win.lift()
            self._commission_win.focus_force()
            if self._commission_panel is not None:
                refresh_commission_pool_from_excel(
                    self._commission_state,
                    self._employees_file_resolved(),
                    show_errors=False,
                    parent=self._commission_win,
                )
                self._commission_panel.refresh_pool_display()
                self._commission_panel.load_from_db_into_ui()
            return

        win = tk.Toplevel(self)
        win.title("Приказ и комиссия по проверке знаний")
        win.minsize(520, 520)
        win.transient(self)
        self._commission_win = win

        def _on_close_commission() -> None:
            self._commission_panel = None
            win.withdraw()

        win.protocol("WM_DELETE_WINDOW", _on_close_commission)

        outer = ttk.Frame(win, padding=8)
        outer.pack(fill=tk.BOTH, expand=True)
        refresh_commission_pool_from_excel(
            self._commission_state,
            self._employees_file_resolved(),
            show_errors=False,
            parent=win,
        )
        panel = CommissionAdminPanel(
            outer,
            state=self._commission_state,
            get_excel_path=self._employees_file_resolved,
            dialog_parent=win,
        )
        panel.pack(fill=tk.BOTH, expand=True)
        self._commission_panel = panel

        bf = ttk.Frame(outer)
        bf.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(bf, text="Закрыть", command=_on_close_commission).pack(side=tk.LEFT)

        win.deiconify()
        win.lift()
        win.focus_force()

    def _build_admin_window_content(self, win: tk.Toplevel) -> None:
        g = {"padx": 5, "pady": 5}
        outer = ttk.Frame(win, padding=8)
        outer.pack(fill=tk.BOTH, expand=True)
        outer.columnconfigure(0, weight=1)

        lf_ex = ttk.Labelframe(outer, text="Файл данных (сотрудники)", padding=6)
        lf_ex.grid(row=0, column=0, sticky=tk.EW, **g)
        ttk.Label(
            lf_ex,
            text=(
                "Файл Data_base.xlsx (или выбранный вами). Лист с сотрудниками: «rabotnik» "
                "или «работники», «сотрудники» и т. п. (без учёта регистра). "
                "Лист «komission»: с 3-й строки A+должность в B, D+должность в E; "
                "настройка приказа и состава — «Администрирование» → «Приказ и комиссия…». "
                "Лист сотрудников — как раньше (заголовки в 1-й строке)."
            ),
            wraplength=500,
        ).grid(row=0, column=0, columnspan=3, sticky=tk.W, **g)
        eb = ttk.Frame(lf_ex)
        eb.grid(row=1, column=0, columnspan=3, sticky=tk.W, pady=(6, 0))
        ttk.Button(eb, text="Загрузить из Excel", command=self.reload_employees).grid(
            row=0, column=0, sticky=tk.W, padx=(0, 6)
        )
        ttk.Button(eb, text="Файл сотрудников…", command=self.pick_employees_excel).grid(
            row=0, column=1, sticky=tk.W, padx=(0, 8)
        )
        self.lbl_employees_file = ttk.Label(lf_ex, text="", wraplength=500)
        self.lbl_employees_file.grid(row=2, column=0, columnspan=3, sticky=tk.W, pady=(6, 0))

        lf_tpl = ttk.Labelframe(outer, text="Шаблон протокола", padding=6)
        lf_tpl.grid(row=1, column=0, sticky=tk.EW, **g)
        tr = ttk.Frame(lf_tpl)
        tr.grid(row=0, column=0, sticky=tk.W)
        ttk.Button(tr, text="Выбрать шаблон", command=self.pick_template).grid(
            row=0, column=0, sticky=tk.W, padx=(0, 6)
        )
        ttk.Button(
            tr,
            text="Переменные шаблона",
            command=self.show_template_variables_help,
        ).grid(row=0, column=1, sticky=tk.W, padx=(0, 8))
        self.lbl_template = ttk.Label(lf_tpl, text=self._template_status_text(), wraplength=500)
        self.lbl_template.grid(row=1, column=0, sticky=tk.W, pady=(8, 0))

        lf_more = ttk.Labelframe(outer, text="Дополнительно при заполнении", padding=6)
        lf_more.grid(row=2, column=0, sticky=tk.EW, **g)
        lf_more.columnconfigure(1, weight=1)
        ttk.Label(lf_more, text="Регистрационный номер (реестр Минтруда):").grid(
            row=0, column=0, sticky=tk.NW, **g
        )
        self.entry_registry_no = ttk.Entry(lf_more, width=55)
        self.entry_registry_no.grid(row=0, column=1, sticky=tk.EW, **g)
        ttk.Label(lf_more, text="Доп. тема / строка для .txt:").grid(
            row=1, column=0, sticky=tk.W, **g
        )
        self.entry_theme = ttk.Entry(lf_more, width=55)
        self.entry_theme.grid(row=1, column=1, sticky=tk.EW, **g)

        jf = ttk.Frame(outer)
        jf.grid(row=3, column=0, sticky=tk.W, **g)
        ttk.Button(jf, text="Журнал протоколов в базе", command=self.show_protocol_journal).grid(
            row=0, column=0, sticky=tk.W, padx=(0, 8)
        )
        ttk.Button(jf, text="Очистить журнал…", command=self.clear_protocol_database).grid(
            row=0, column=1, sticky=tk.W
        )

        ttk.Label(
            outer,
            text="Справка по источникам полей программ — в подсказках ниже (разверните окно).",
            font=("Segoe UI", 8),
            foreground="#444",
        ).grid(row=4, column=0, sticky=tk.W, pady=(10, 2))
        help_fr = ttk.Frame(outer)
        help_fr.grid(row=5, column=0, sticky=tk.EW, **g)
        for pi, (key, _sheet, _fb) in enumerate(PROTOCOL_PROGRAM_DEFS):
            ttk.Label(
                help_fr,
                text=f"{key}: {PROTOCOL_PROGRAM_UI_LABELS.get(key, key)}",
                wraplength=500,
            ).grid(row=pi, column=0, sticky=tk.W, pady=1)

        ttk.Label(
            outer,
            text="Приказ о комиссии и состав комиссии — пункт меню «Администрирование» → «Приказ и комиссия…».",
            font=("Segoe UI", 8),
            foreground="#444",
            wraplength=500,
        ).grid(row=6, column=0, sticky=tk.W, pady=(6, 0))

    def _template_status_text(self) -> str:
        if self.template_path is None:
            return f"Файл шаблона: {PROTOCOL_TEMPLATE_FILENAME} (по умолчанию, папка с программой)"
        return f"Файл шаблона: {self.template_path}"

    def show_template_variables_help(self) -> None:
        win = tk.Toplevel(self)
        win.title("Переменные и маркеры шаблона протокола")
        win.minsize(560, 420)
        frm = ttk.Frame(win, padding=8)
        frm.pack(fill=tk.BOTH, expand=True)
        frm.rowconfigure(0, weight=1)
        frm.columnconfigure(0, weight=1)
        box = tk.Text(frm, wrap=tk.WORD, font=("Segoe UI", 10))
        sb = ttk.Scrollbar(frm, command=box.yview)
        box.configure(yscrollcommand=sb.set)
        box.grid(row=0, column=0, sticky=tk.NSEW)
        sb.grid(row=0, column=1, sticky=tk.NS)
        box.insert("1.0", PROTOCOL_TEMPLATE_VARIABLES_DOC)
        box.configure(state=tk.DISABLED)
        ttk.Button(frm, text="Закрыть", command=win.destroy).grid(
            row=1, column=0, columnspan=2, pady=(10, 0)
        )

    def clear_protocol_database(self) -> None:
        """Удаление всех записей журнала protocols (с подтверждением)."""
        if not messagebox.askyesno(
            "Очистка журнала",
            "Удалить все записи журнала протоколов из базы?\n\nДействие нельзя отменить.",
        ):
            return
        try:
            n = clear_protocol_journal()
        except sqlite3.Error as e:
            messagebox.showerror("База данных", str(e))
            return
        messagebox.showinfo("Журнал", f"Удалено записей: {n}.")

    def show_protocol_journal(self) -> None:
        """Окно со списком сохранённых протоколов из SQLite и текстом записи."""
        try:
            rows = get_all_protocols()
        except sqlite3.Error as e:
            messagebox.showerror("База данных", str(e))
            return

        win = tk.Toplevel(self)
        win.title(f"Журнал протоколов — {DATABASE_FILENAME}")
        win.minsize(720, 520)
        win.geometry("900x600")

        outer = ttk.Frame(win, padding=8)
        outer.pack(fill=tk.BOTH, expand=True)
        outer.rowconfigure(2, weight=1)
        outer.columnconfigure(0, weight=1)

        dbp = database_path()
        ttk.Label(
            outer,
            text=f"Файл базы: {dbp}",
            font=("Segoe UI", 9),
        ).grid(row=0, column=0, sticky=tk.W, pady=(0, 4))

        btn_bar = ttk.Frame(outer)
        btn_bar.grid(row=1, column=0, sticky=tk.W, pady=(0, 6))

        pw = ttk.Panedwindow(outer, orient=tk.VERTICAL)
        pw.grid(row=2, column=0, sticky=tk.NSEW)
        outer.rowconfigure(2, weight=1)

        top_fr = ttk.Frame(pw, padding=2)
        bot_fr = ttk.Frame(pw, padding=2)
        pw.add(top_fr, weight=1)
        pw.add(bot_fr, weight=2)

        top_fr.rowconfigure(0, weight=1)
        top_fr.columnconfigure(0, weight=1)
        sb_list = ttk.Scrollbar(top_fr)
        lb = tk.Listbox(
            top_fr,
            height=8,
            font=("Consolas", 10),
            yscrollcommand=sb_list.set,
            exportselection=False,
        )
        lb.grid(row=0, column=0, sticky=tk.NSEW)
        sb_list.grid(row=0, column=1, sticky=tk.NS)
        sb_list.configure(command=lb.yview)

        bot_fr.rowconfigure(1, weight=1)
        bot_fr.columnconfigure(0, weight=1)
        ttk.Label(bot_fr, text="Текст записи (content):").grid(row=0, column=0, sticky=tk.W)
        sb_txt = ttk.Scrollbar(bot_fr)
        txt = tk.Text(bot_fr, wrap=tk.WORD, font=("Segoe UI", 10), height=14, state=tk.DISABLED)
        txt.grid(row=1, column=0, sticky=tk.NSEW)
        sb_txt.grid(row=1, column=1, sticky=tk.NS)
        sb_txt.configure(command=txt.yview)
        txt.configure(yscrollcommand=sb_txt.set)

        def journal_line(r: dict[str, Any]) -> str:
            tid = r.get("id", "")
            dt = (r.get("date") or "").strip()
            gr = (r.get("grade") or "").strip()
            fio = (r.get("fio") or "").strip()
            if len(fio) > 52:
                fio = fio[:49] + "…"
            top = (r.get("topic") or "").strip()
            if len(top) > 36:
                top = top[:33] + "…"
            ca = (r.get("created_at") or "").strip()
            return f"#{tid}  {dt}  {gr}  |  {fio}  |  {top}  |  {ca}"

        def refresh_list() -> None:
            nonlocal rows
            try:
                rows = get_all_protocols()
            except sqlite3.Error as e:
                messagebox.showerror("База данных", str(e))
                return
            lb.delete(0, tk.END)
            for r in rows:
                lb.insert(tk.END, journal_line(r))
            txt.configure(state=tk.NORMAL)
            txt.delete("1.0", tk.END)
            txt.configure(state=tk.DISABLED)

        def on_select(_evt: object | None = None) -> None:
            sel = lb.curselection()
            if not sel:
                return
            r = rows[int(sel[0])]
            body = r.get("content") or ""
            txt.configure(state=tk.NORMAL)
            txt.delete("1.0", tk.END)
            txt.insert("1.0", body)
            txt.configure(state=tk.DISABLED)

        def copy_to_main_preview() -> None:
            sel = lb.curselection()
            if not sel:
                messagebox.showinfo("Журнал", "Выберите запись в списке.")
                return
            r = rows[int(sel[0])]
            body = (r.get("content") or "").strip()
            if not body:
                messagebox.showinfo("Журнал", "В записи нет текста.")
                return
            self.text_preview.configure(state=tk.NORMAL)
            self.text_preview.delete("1.0", tk.END)
            self.text_preview.insert("1.0", body)
            self.btn_save.state(["!disabled"])
            self.btn_save_pdf.state(["!disabled"])
            messagebox.showinfo("Журнал", "Текст загружен в «Предварительный просмотр».\nМожно сохранить в DOCX или PDF.")

        def on_clear_journal() -> None:
            if not messagebox.askyesno(
                "Очистка журнала",
                "Удалить все записи журнала протоколов из базы?\n\nДействие нельзя отменить.",
                parent=win,
            ):
                return
            try:
                n = clear_protocol_journal()
            except sqlite3.Error as e:
                messagebox.showerror("База данных", str(e), parent=win)
                return
            refresh_list()
            messagebox.showinfo("Журнал", f"Удалено записей: {n}.", parent=win)

        ttk.Button(btn_bar, text="Обновить список", command=refresh_list).grid(
            row=0, column=0, padx=(0, 8)
        )
        ttk.Button(btn_bar, text="В предпросмотр", command=copy_to_main_preview).grid(
            row=0, column=1, padx=(0, 8)
        )
        ttk.Button(btn_bar, text="Очистить журнал…", command=on_clear_journal).grid(
            row=0, column=2, padx=(0, 8)
        )
        ttk.Button(btn_bar, text="Закрыть", command=win.destroy).grid(row=0, column=3)

        lb.bind("<<ListboxSelect>>", on_select)
        for r in rows:
            lb.insert(tk.END, journal_line(r))
        if rows:
            lb.selection_set(0)
            on_select()

    def pick_template(self) -> None:
        path = filedialog.askopenfilename(
            title="Выберите шаблон протокола",
            filetypes=[
                ("Текст UTF-8", "*.txt"),
                ("Документ Word", "*.docx"),
                ("Все файлы", "*.*"),
            ],
        )
        if not path:
            return
        self.template_path = Path(path).expanduser().resolve()
        self.lbl_template.configure(text=self._template_status_text())

    def _refresh_employees_file_label(self) -> None:
        p = self._employees_file_resolved()
        src = "выбранный файл" if self.employees_excel_path else "по умолчанию рядом с программой"
        extra = ""
        if not p.is_file():
            extra = " — файл не найден (положите Data_base.xlsx в папку с main.py или «Файл сотрудников…»)"
        elif self._employee_records:
            extra = f" — загружено: {len(self._employee_records)}"
        else:
            extra = (
                " — список пуст или файл не прочитан (кнопка «Загрузить из Excel» покажет ошибку; "
                "проверьте лист и первую строку заголовков)"
            )
        self.lbl_employees_file.configure(text=f"{src}: {p.name}{extra}")

    def pick_employees_excel(self) -> None:
        path = filedialog.askopenfilename(
            title="Файл с базой сотрудников",
            filetypes=[
                ("Excel", "*.xlsx"),
                ("Excel с макросами", "*.xlsm"),
                ("Все файлы", "*.*"),
            ],
        )
        if not path:
            return
        self.employees_excel_path = Path(path).expanduser().resolve()
        self._refresh_employees_file_label()
        self.reload_employees(show_errors=True)

    def reload_employees(self, *, show_errors: bool = False) -> None:
        path = self._employees_file_resolved()
        try:
            self._employee_records = load_employees_from_excel(path)
        except EmployeeExcelError as e:
            self._employee_records = []
            self._employee_filter_indices = []
            self.list_employees.delete(0, tk.END)
            self._refresh_employees_file_label()
            if show_errors:
                messagebox.showerror("Сотрудники Excel", str(e))
            return
        self._refilter_employee_list()
        self._refresh_employees_file_label()
        refresh_commission_pool_from_excel(
            self._commission_state,
            self._employees_file_resolved(),
            show_errors=False,
            parent=self,
        )
        p = self._commission_panel
        if p is not None:
            try:
                if p.winfo_exists():
                    p.refresh_pool_display()
            except tk.TclError:
                self._commission_panel = None

    def _refilter_employee_list(self) -> None:
        if not hasattr(self, "list_employees"):
            return
        records = self._employee_records
        q = self.var_emp_search.get().strip().lower()
        prev_sel_global: list[int] = []
        for li in self.list_employees.curselection():
            li = int(li)
            if 0 <= li < len(self._employee_filter_indices):
                prev_sel_global.append(self._employee_filter_indices[li])
        if not q:
            self._employee_filter_indices = list(range(len(records)))
        else:
            self._employee_filter_indices = []
            for i, rec in enumerate(records):
                blob = f"{rec.fio} {rec.profession} {rec.subdivision}".lower()
                if q in blob:
                    self._employee_filter_indices.append(i)
        self.list_employees.delete(0, tk.END)
        for gi in self._employee_filter_indices:
            self.list_employees.insert(tk.END, listbox_label_for_employee(records[gi]))
        for gi in prev_sel_global:
            if gi in self._employee_filter_indices:
                pos = self._employee_filter_indices.index(gi)
                self.list_employees.selection_set(pos)

    def _try_autoload_employees(self) -> None:
        if self._employees_file_resolved().is_file():
            self.reload_employees(show_errors=False)

    def _on_employee_list_select(self, _event: object | None = None) -> None:
        sel = self.list_employees.curselection()
        if len(sel) != 1:
            return
        li = int(sel[0])
        if li < 0 or li >= len(self._employee_filter_indices):
            return
        rec = self._employee_records[self._employee_filter_indices[li]]
        self.entry_position.delete(0, tk.END)
        self.entry_position.insert(0, rec.profession)
        self.entry_subdivision.delete(0, tk.END)
        self.entry_subdivision.insert(0, rec.subdivision)

    def _collect_program_keys_and_titles(
        self, persons_raw: list[EmployeeRecord]
    ) -> tuple[list[str], list[str]]:
        """Порядок: Б → PP → СИЗ → В; «Б» — название с листа B; PP/СИЗ — V_PROF; «В» — блок по V_PROF."""
        path = self._employees_file_resolved()
        keys: list[str] = []
        titles: list[str] = []
        for key, _, fallback in PROTOCOL_PROGRAM_DEFS:
            if not self._prog_vars[key].get():
                continue
            keys.append(key)
            if key == "V":
                inner = resolve_v_program_inner_text_global(path, persons_raw, fallback)
                titles.append(format_v_program_table_block_title(inner, fallback))
            elif key == "B":
                t = load_excel_first_nonempty_in_column(
                    path, B_PROGRAM_SHEET_NAME, B_PROGRAM_TITLE_COL
                ).strip()
                titles.append(t if t else fallback)
            else:
                col = {
                    "PP": V_PROF_TITLE_COL_PP,
                    "SIZ": V_PROF_TITLE_COL_SIZ,
                }.get(key)
                if col is None:
                    titles.append(fallback)
                    continue
                t = load_v_prof_first_nonempty_in_column(path, col).strip()
                titles.append(t if t else fallback)
        return keys, titles

    def _collect_table_persons(self) -> list[EmployeeRecord]:
        sel = self.list_employees.curselection()
        if sel:
            order = sorted(int(i) for i in sel)
            return [
                self._employee_records[self._employee_filter_indices[i]]
                for i in order
                if 0 <= i < len(self._employee_filter_indices)
            ]
        fio = self.entry_fio.get().strip()
        if fio:
            return [
                EmployeeRecord(
                    fio=fio,
                    profession=self.entry_position.get().strip(),
                    subdivision=self.entry_subdivision.get().strip(),
                )
            ]
        return []

    def generate_protocol(self) -> None:
        theme = self.entry_theme.get().strip()
        date_str = self.entry_date.get().strip()
        protocol_no = self.entry_protocol_no.get().strip()
        registry_no = self.entry_registry_no.get().strip()
        grade = self.combo_grade.get().strip()
        check_type = self.combo_check_type.get().strip() or "плановая"

        persons_raw = self._collect_table_persons()
        if not persons_raw:
            messagebox.showwarning(
                "Проверка",
                "Выберите одного или нескольких сотрудников в списке "
                "или введите ФИО вручную (без выбора в списке).",
            )
            return

        excel_path = self._employees_file_resolved()
        program_keys, program_titles = self._collect_program_keys_and_titles(persons_raw)
        tpl = self.template_path or protocol_template_path()
        is_docx = _is_word_protocol_template(tpl)
        if is_docx and not program_titles:
            messagebox.showwarning(
                "Программы",
                "Для шаблона Word отметьте хотя бы одну программу обучения "
                f"(данные — листы {B_PROGRAM_SHEET_NAME} и {V_PROF_SHEET_NAME}).",
            )
            return

        if not date_str:
            date_str = date.today().strftime("%d.%m.%Y")
            self.entry_date.delete(0, tk.END)
            self.entry_date.insert(0, date_str)

        try:
            if is_docx:
                doc, table_excess = build_filled_protocol_document(
                    tpl,
                    protocol_no=protocol_no,
                    date_str=date_str,
                    theme=theme,
                    table_persons=persons_raw,
                    program_titles=program_titles,
                    program_keys=program_keys,
                    excel_path=excel_path,
                    persons_v_raw=persons_raw,
                    grade=grade,
                    registry_no=registry_no,
                    check_type=check_type,
                )
                if table_excess > 0:
                    messagebox.showwarning(
                        "Таблица протокола",
                        "В документе не найдена таблица с маркерными строками «1.1», «2.1», … "
                        "и текстом «ФИО» во второй колонке — список сотрудников в таблицу не подставлен. "
                        "Проверьте шаблон .docx.",
                    )
                render_document_to_text_widget(self.text_preview, doc)
                text = document_to_plain_text(doc)
            else:
                txt_theme = "; ".join(program_titles) if program_titles else theme
                if not txt_theme.strip():
                    messagebox.showwarning(
                        "Тема",
                        "Отметьте программы обучения или введите доп. тему для текстового шаблона.",
                    )
                    return
                text = build_protocol_text(
                    txt_theme,
                    date_str,
                    protocol_no=protocol_no,
                    template_path=self.template_path,
                )
                self.text_preview.delete("1.0", tk.END)
                self.text_preview.insert("1.0", text)
        except ProtocolTemplateError as e:
            messagebox.showerror("Шаблон протокола", str(e))
            return
        except ValueError as e:
            messagebox.showerror("Шаблон протокола", str(e))
            return

        self.btn_save.state(["!disabled"])
        self.btn_save_pdf.state(["!disabled"])
        self._persist_protocol_no_field()

        try:
            fio_summary = ", ".join(p.fio for p in persons_raw)
            topic_db = "; ".join(program_titles) if program_titles else theme
            save_protocol(fio_summary, topic_db, date_str, grade, text)
        except sqlite3.Error as e:
            messagebox.showwarning(
                "База данных",
                f"Протокол сформирован, но не удалось сохранить в базу:\n{e}",
            )

    def save_to_pdf(self) -> None:
        content = self.text_preview.get("1.0", tk.END).rstrip()
        if not content:
            messagebox.showwarning("Сохранение", "Нет текста протокола для сохранения.")
            return

        path = filedialog.asksaveasfilename(
            title="Сохранить протокол в PDF",
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf"), ("Все файлы", "*.*")],
        )
        if not path:
            return

        tpl = self.template_path or protocol_template_path()
        theme = self.entry_theme.get().strip()
        date_str = self.entry_date.get().strip()
        protocol_no = self.entry_protocol_no.get().strip()
        registry_no = self.entry_registry_no.get().strip()
        grade = self.combo_grade.get().strip()
        check_type = self.combo_check_type.get().strip() or "плановая"
        persons_raw = self._collect_table_persons()
        if not persons_raw:
            messagebox.showwarning(
                "Сохранение",
                "Нет данных сотрудников: сформируйте протокол или выберите сотрудников / введите ФИО.",
            )
            return
        excel_path = self._employees_file_resolved()
        program_keys, program_titles = self._collect_program_keys_and_titles(persons_raw)

        try:
            if _is_word_protocol_template(tpl):
                if not program_titles:
                    messagebox.showwarning(
                        "Сохранение",
                        "Для DOCX отметьте хотя бы одну программу обучения (как при формировании).",
                    )
                    return
                write_protocol_pdf_from_docx_template(
                    tpl,
                    path,
                    protocol_no=protocol_no,
                    date_str=date_str,
                    theme=theme,
                    table_persons=persons_raw,
                    program_titles=program_titles,
                    program_keys=program_keys,
                    excel_path=excel_path,
                    persons_v_raw=persons_raw,
                    grade=grade,
                    registry_no=registry_no,
                    check_type=check_type,
                )
            else:
                write_protocol_pdf(path, content)
        except ValueError as e:
            messagebox.showerror("Шаблон", str(e))
            return
        except OSError as e:
            messagebox.showerror("Ошибка", f"Не удалось записать PDF-файл:\n{e}")
            return
        except RuntimeError as e:
            messagebox.showerror("Ошибка PDF", str(e))
            return
        except Exception as e:
            messagebox.showerror(
                "Ошибка PDF",
                f"Не удалось сформировать PDF (нужны Word и docx2pdf):\n{e}",
            )
            return

        messagebox.showinfo("Готово", f"Протокол сохранён в PDF:\n{path}")

    def save_to_docx(self) -> None:
        content = self.text_preview.get("1.0", tk.END).rstrip()
        if not content:
            messagebox.showwarning("Сохранение", "Нет текста протокола для сохранения.")
            return

        path = filedialog.asksaveasfilename(
            title="Сохранить протокол в Word",
            defaultextension=".docx",
            filetypes=[("Документ Word", "*.docx"), ("Все файлы", "*.*")],
        )
        if not path:
            return

        tpl = self.template_path or protocol_template_path()
        theme = self.entry_theme.get().strip()
        date_str = self.entry_date.get().strip()
        protocol_no = self.entry_protocol_no.get().strip()
        registry_no = self.entry_registry_no.get().strip()
        grade = self.combo_grade.get().strip()
        check_type = self.combo_check_type.get().strip() or "плановая"
        persons_raw = self._collect_table_persons()
        if not persons_raw:
            messagebox.showwarning(
                "Сохранение",
                "Нет данных сотрудников: сформируйте протокол или выберите сотрудников / введите ФИО.",
            )
            return
        excel_path = self._employees_file_resolved()
        program_keys, program_titles = self._collect_program_keys_and_titles(persons_raw)

        try:
            if _is_word_protocol_template(tpl):
                if not program_titles:
                    messagebox.showwarning(
                        "Сохранение",
                        "Для DOCX отметьте хотя бы одну программу обучения (как при формировании).",
                    )
                    return
                save_protocol_docx_from_template(
                    tpl,
                    path,
                    protocol_no=protocol_no,
                    date_str=date_str,
                    theme=theme,
                    table_persons=persons_raw,
                    program_titles=program_titles,
                    program_keys=program_keys,
                    excel_path=excel_path,
                    persons_v_raw=persons_raw,
                    grade=grade,
                    registry_no=registry_no,
                    check_type=check_type,
                )
            else:
                write_protocol_docx(path, content)
        except ValueError as e:
            messagebox.showerror("Шаблон", str(e))
            return
        except OSError as e:
            messagebox.showerror("Ошибка", f"Не удалось записать DOCX:\n{e}")
            return
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сформировать DOCX:\n{e}")
            return

        messagebox.showinfo("Готово", f"Протокол сохранён в DOCX:\n{path}")


def main() -> None:
    init_db()
    app = ProtocolApp()
    app.mainloop()


if __name__ == "__main__":
    main()
