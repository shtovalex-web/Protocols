# -*- coding: utf-8 -*-
"""Сборка и правка документа протокола (.docx): шаблон, таблица результатов, кэш фрагментов V_PROF."""

from __future__ import annotations

import json
import re
import sqlite3
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.enum.text import WD_BREAK, WD_TAB_ALIGNMENT, WD_UNDERLINE
from docx.shared import Inches, Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run as DocxRun

from app_paths import application_bundle_dir
from commission_admin import (
    COMMISSION_KIND_OT,
    COMMISSION_KIND_TECH,
    COMMISSION_ORDER_APPROVER_PLACEHOLDER,
    COMMISSION_VENUE_PLACEHOLDER,
    apply_commission_insertions_to_line,
    build_commission_signature_suffix_payload,
    build_commission_template_payload,
    commission_chair_anchor_prefix,
    commission_members_anchor_prefix,
    load_commission_protocol_context_from_db,
    load_commission_state_from_db,
    parse_commission_members_two_column_text,
)
from employees_io import EmployeeRecord, format_fio_filename_surname_initials
from excel_data_cache import (
    format_training_hours_ru,
    get_cached_v_prof_column,
    get_cached_v_registry_rows,
    get_training_hours_for_program_key,
)
from mintrud_trained_registry import (
    TrainedRegistryIndex,
    filter_candidates_for_program_block,
    load_trained_registry_index,
    merge_registry_tokens,
)
from protocol_paths import database_path, load_last_protocol_no
from programs_v_prof import (
    VProfLayout,
    v_prof_layout_for_path,
    v_prof_row_program_header_titles,
    v_prof_row_program_table_fragments,
)
from v_program_registry_match import (
    fg_line_comparison_key as _fg_line_comparison_key,
    match_v_registry_fragment,
    norm_profession_key as _norm_profession_key,
)

# Табуляция в шаблоне default_protocol.docx (w:tab/@w:pos) — первая зона ФИО в шапке.
_COMMISSION_HEADER_TAB1_TWIPS = 284
# Вторая колонка членов комиссии (~8,4 см от левого поля).
_COMMISSION_HEADER_TAB2 = Inches(3.3)

# Текст бланка протокола без организационной шапки (абзацы от «ПРОТОКОЛ №» и таблица ниже двух строк шапки).
PROTOCOL_BODY_FONT_PT = 11
# Наименования программ в шапке (после комиссии) и в заголовках блоков таблицы результатов.
PROTOCOL_PROGRAM_TITLE_FONT_PT = 10
PROTOCOL_TEMPLATE_FILENAME = "default_protocol.docx"
PROTOCOL_TEMPLATE_TECH_FILENAME = "default_protocol_tehnicheskiy.docx"
# Таблица «Председатель комиссии» в .docx: только ФИО (без должности и блока подписи).
COMMISSION_CHAIR_FIO_PLACEHOLDER = "{{ПРЕДСЕДАТЕЛЬ_ФИО}}"
# Таблица «Члены комиссии»: строка-шаблон дублируется; во 2-й и 3-й колонках (индексы 1 и 2) — по одному ФИО в строке.
COMMISSION_MEMBER_FIO_PLACEHOLDER = "{{ЧЛЕНЫ_КОМИССИИ_ФИО}}"


def is_word_protocol_template(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() in (".docx", ".docm")


def _sanitize_windows_filename_part(s: str) -> str:
    """Символы, недопустимые в имени файла Windows, заменяются на «_»."""
    if not s:
        return ""
    bad = '\\/:*?"<>|\n\r\t'
    t = "".join("_" if c in bad else c for c in s.strip())
    return t.rstrip(" .") or ""


def default_protocol_save_filename(
    protocol_no: str,
    date_str: str,
    extension: str,
    *,
    person_suffix: str = "",
) -> str:
    """
    Предлагаемое имя: «Протокол № <номер> от <дата>» [Фамилия И.О.]. extension — '.pdf' или '.docx'.
    """
    ext = extension if extension.startswith(".") else f".{extension}"
    no = _sanitize_windows_filename_part((protocol_no or "").strip()) or "без номера"
    dt = _sanitize_windows_filename_part((date_str or "").strip())
    if not dt:
        dt = date.today().strftime("%d.%m.%Y")
    base = f"Протокол № {no} от {dt}"
    ps = _sanitize_windows_filename_part((person_suffix or "").strip())
    if ps:
        base = f"{base} {ps}"
    if len(base) > 180:
        base = base[:180].rstrip()
    return f"{base}{ext}"


def protocol_sequence_start_int(raw: str) -> int:
    """Целый номер протокола для нумерации подряд (из поля «№» или вида N-М-ГГ)."""
    s = (raw or "").strip()
    if not s:
        fall = load_last_protocol_no().strip()
        if fall:
            if "-" in fall:
                left = fall.split("-", 1)[0].strip()
                if left.isdigit():
                    return int(left)
            elif fall.isdigit():
                return int(fall)
        return 1
    if "-" in s:
        left = s.split("-", 1)[0].strip()
        if left.isdigit():
            return int(left)
    m = re.match(r"^(\d+)", s)
    if m:
        return int(m.group(1))
    return 1
V_PROF_SHEET_NAME = "V_PROF"
# Лист B: столбец 2 — полное название программы «Б» в шапке и в таблице протокола.
B_PROGRAM_SHEET_NAME = "B"
B_PROGRAM_TITLE_COL = 2
# Листы PP и SIZ: столбец 2 — наименование блока в таблице протокола (как на листе B).
PP_TABLE_SHEET_NAME = "PP"
SIZ_TABLE_SHEET_NAME = "SIZ"
# Лист V_PROF: A — должность; 2 — текст для якорных абзацев после «…по программе(ам):» для «Б»;
# 3–4 — PP / СИЗ; с 5 — фрагменты программы «В».
# Значения по умолчанию (старый формат); фактические столбцы — из programs_v_prof.v_prof_layout_for_path.
V_PROF_TITLE_COL_B = 2
V_PROF_TITLE_COL_PP = 3
V_PROF_TITLE_COL_SIZ = 4
V_PROF_PARTS_FIRST_COL = 5
V_PROF_PARTS_LAST_COL = 22
# Версия схемы кэша фрагментов программы «В»; 4 — таблица: столбец B листа V, шапка: C.
V_PROF_PARTS_CACHE_SCHEMA = 4

# Лист V: столбец B — сопоставление с фрагментом «В» с V_PROF; столбец C — текст в списке после
# «…провела проверку… по программе(ам):» (таблица и шапка темы «В» — без подстановки с листа V).
V_PROGRAM_REGISTRY_SHEET = "V"
V_PROGRAM_SHEET_NAME_ALIASES: tuple[str, ...] = ("v", "в")
# Лист V: A — ID программы в Гос. реестре; B — сопоставление с V_PROF; C — наименование в списке после проверки.
V_PROGRAM_GOS_REGISTRY_ID_COL_A = 1
V_PROGRAM_MATCH_COL_B = 2
V_PROGRAM_TITLE_COL_C = 3
V_PROGRAM_SHEET_MAX_ROWS = 2000

# Ключ, лист-источник названия/якорей, текст по умолчанию если ячейка пуста.
PROTOCOL_PROGRAM_DEFS: tuple[tuple[str, str, str], ...] = (
    ("B", B_PROGRAM_SHEET_NAME, "Программа обучения «Б»"),
    (
        "PP",
        PP_TABLE_SHEET_NAME,
        "Программа обучения по оказанию первой помощи пострадавшим",
    ),
    (
        "SIZ",
        SIZ_TABLE_SHEET_NAME,
        "Программа обучения использование (применение) средств индивидуальной защиты",
    ),
    ("V", V_PROF_SHEET_NAME, "Программа обучения «В»"),
)
PROTOCOL_PROGRAM_UI_LABELS: dict[str, str] = {
    "B": (
        f"Программа «Б» (таблица — {B_PROGRAM_SHEET_NAME}, ст. {B_PROGRAM_TITLE_COL}; "
        f"абзац после проверки — {V_PROF_SHEET_NAME}, ст. {V_PROF_TITLE_COL_B})"
    ),
    "PP": (
        f"ПП: в таблице — лист {PP_TABLE_SHEET_NAME}, ст. {B_PROGRAM_TITLE_COL}; "
        f"абзац после проверки — {V_PROF_SHEET_NAME}, ст. {V_PROF_TITLE_COL_PP}"
    ),
    "SIZ": (
        f"СИЗ: в таблице — лист {SIZ_TABLE_SHEET_NAME}, ст. {B_PROGRAM_TITLE_COL}; "
        f"абзац после проверки — {V_PROF_SHEET_NAME}, ст. {V_PROF_TITLE_COL_SIZ}"
    ),
    "V": (
        f"Программа «В» (лист {V_PROF_SHEET_NAME}: матрица «Да»; шапка — лист V ст. C; "
        "таблица — лист V ст. B)"
    ),
}
# Краткие подписи чекбоксов на главном экране (подробности — в окне администрирования).
PROTOCOL_PROGRAM_CHECKBOX_SHORT: dict[str, str] = {
    "B": "«Б»",
    "PP": "ПП",
    "SIZ": "СИЗ",
    "V": "«В»",
}

# Полный перечень маркеров шаблона (см. кнопку «Переменные шаблона» в окне программы).
PROTOCOL_TEMPLATE_VARIABLES_DOC = """\
ПЕРЕМЕННЫЕ И МАРКЕРЫ ШАБЛОНА ПРОТОКОЛА
========================================

В шаблоне (.docx или .txt) не используются имена вроде {{ФИО}}. Программа ищет
описанные ниже фрагменты текста и подставляет данные из полей формы.

────────────────────────────────────────
1) НОМЕР ПРОТОКОЛА  →  поле «№ протокола»
────────────────────────────────────────
В шаблоне должна быть строка с фрагментом  ПРОТОКОЛ №  (любой регистр; допускаются # или знак №;
можно после названия организации). Подчёркивания _____ после знака не обязательны.

Если после «№» идут подчёркивания (от трёх символов «_» подряд) — они заменяются на номер;
если подчёркиваний нет — номер вставляется сразу после знака.

Если в форме указан номер — подставляется строка вида
  <номер>-<месяц>-<год>
где <номер> — то, что введено в поле «№ протокола», <месяц> — число 1–12 из поля «Дата»
(без ведущего нуля), <год> — две последние цифры года из «Даты» (например, при номере 12
и дате 25.03.2026 в бланк попадёт 12-3-26; строка шаблона остаётся «ПРОТОКОЛ № …»).
Если дата не в формате ДД.ММ.ГГГГ (или ГГ) — подставляется только введённый номер, без суффикса.
Если номер не введён — строка с «ПРОТОКОЛ №» не меняется (подчёркивания остаются).

────────────────────────────────────────
2) ДАТА  →  поле «Дата» (формат ДД.ММ.ГГГГ)
────────────────────────────────────────
Отдельная строка, у которой после удаления пробелов по краям:
  • начинается с  «__
  • заканчивается на  г.

Типичный вид (как в бланке):
  «__» ___________ 20__ г.

Всю такую строку программа заменяет на дату словами, например:
  «25» марта 2026 г.

Важно: строка «В соответствии с приказом … от «__» … 20__ г.» не подходит
под это правило (она не начинается с «__»).

────────────────────────────────────────
2а) ПРИКАЗ О КОМИССИИ  →  меню «Администрирование» → «Приказ и комиссия…» (база protocols.db)
────────────────────────────────────────
В шаблоне .docx (не удалять при правке бланка):
  • {{ПОДРАЗДЕЛЕНИЕ_ПРОВЕРКИ}} — подразделение / место проверки (многострочный текст из формы);
  • {{УТВЕРДИЛ_ПРИКАЗ}} — кем утверждён приказ (фрагмент после «В соответствии с приказом »,
    например: руководителя филиала Урайское УМН АО «Транснефть – Сибирь»).
Отдельные значения для вкладок «Охрана труда» и «Технич. вопросы». Программа читает шаблон с диска
в память — исходный файл при формировании не перезаписывается.

В абзаце, где есть слова о комиссии / приказе (например: комисс, председателя, членов, приказ и « от»),
при сохранённых в базе данных № и дате приказа, председателе и членах:
  • сразу после подстроки « от» (пробел и «от») вставляется дата приказа словами, как в п. 2
    (поле «Дата приказа» в формате ДД.ММ.ГГГГ);
  • после первого в абзаце « №», не входящего в сочетание «ПРОТОКОЛ №», подставляется номер приказа;
  • после слова «председателя» — с новой строки ФИО и должность в родительном падеже (с подчёркиванием, как у членов);
  • после фразы «членов комиссии» (если есть) или слова «членов» — с новой строки в два столбца
    (табуляция, выравнивание по позициям табуляции шаблона; ФИО и должность в родительном падеже,
    с подчёркиванием в .docx).
Родительный падеж: при установленных пакетах pymorphy2 и pymorphy2-dicts-ru; иначе текст без склонения.
Рекомендуется: pip install pymorphy2 pymorphy2-dicts-ru

────────────────────────────────────────
3) ПРОГРАММА / ТЕМА  →  чекбоксы программ и строка подчёркиваний
────────────────────────────────────────
Для .docx: лист B — название программы «Б» в шапке и в таблице из первой непустой
ячейки столбца 2 (со 2-й строки). Названия блоков PP и СИЗ в таблице — с листов PP и SIZ
(так же: первая непустая ячейка столбца 2 со 2-й строки). После абзаца «…по программе(ам):» для «Б»
подставляется текст из листа V_PROF, столбец 2, по совпадению должности (столбец A). На V_PROF для
абзацев после проверки: «PP» — столбец 3, «СИЗ» — столбец 4. Программа «В»: столбец A — должность;
фрагменты — непустые ячейки столбцов 5–19.
Лист V: A — ID программы в Гос. реестре (для программ «В»); B — сопоставление с фрагментом V_PROF;
C — наименование в маркированном списке после «…провела проверку… по программе(ам):».
Соответствие A↔B↔C загружается в память при чтении листа (для дальнейшего использования ID реестра).
Шапка протокола, заголовок блока «В» в таблице и строки сотрудников — по тексту фрагментов с V_PROF (подстановка C только в списке после проверки). Совпадения (сцепка в заголовке блока и в таблице).
Кэш V_PROF — в protocols.db
(таблица v_prof_cache; поле parts_schema — сброс кэша фрагментов при смене логики столбцов),
сброс при изменении времени файла Excel.
Для ПП, СИЗ, «В» при совмещаемой второй профессии: в таблице — одна строка на работника после объединения
одинакового ФИО (в графе должности — только основная). Блок «Б» — все выбранные строки списка,
даже если ФИО повторяется. Фрагменты программы «В»
по двум должностям объединяются без повторов; строк «№ …» в ячейке результата столько же,
сколько фрагментов. Для «Б» при совмещении — две строки «№ …» (две записи в реестре);
для PP и СИЗ — одна строка «№ …» (обучение одно). Для «В» в таблице перед каждым работником
добавляется строка-шапка: «Программа (В)» и в скобках нумерованный перечень тем этого работника (1. … 2. …);
в ячейке результата — строки «1. № …», «2. № …» в том же порядке для сопоставления с шапкой блока;
в первой колонке номер блока программы — только у шапки первого работника в блоке, у следующих шапок пусто; в колонке «ФИО» — только ФИО.
В шапке протокола названия программ через «; ».

Отдельная строка с подчёркиваниями должна:
  • в конце (после пробелов) иметь запятую:  ,
  • состоять в основном из символов подчёркивания _ (длина > 20),
  • кроме подчёркиваний, пробелов и запятой не содержать другого текста.

Для .txt или если программы не отмечены: используется поле «Доп. тема».

После абзаца, в котором есть слова: провела, проверку, программе, охраны труда
(типично: «…провела проверку знания требований охраны труда работников по программе(ам):»),
при отмеченных программах B / PP / SIZ / V подставляются абзацы с префиксом «- »:
лист V_PROF, столбец A — должность; для «Б» — ячейка столбца 2, «PP» — 3, «СИЗ» — 4;
для «В» — по одному абзацу на каждую непустую ячейку столбцов 5–19. Если совпадения
по должности или ячейки пусты — соответствующая строка не добавляется.

────────────────────────────────────────
4) ДОПОЛНИТЕЛЬНО ДЛЯ ФАЙЛА .docx (и для .txt со срезом)
────────────────────────────────────────
Чтобы из документа выделялся фрагмент бланка:

  • Нужна строка с маркером начала: фрагмент  ПРОТОКОЛ №  (или #); в этой строке не должно быть
    слова  ТЕХНИЧЕСКИХ  (так отличается пустой бланк от заполненного примера в одном файле).

  • Конец фрагмента: строка, начинающаяся с  Приложение
    (всё после неё в шаблон не попадает).

Если в .txt нет строки с  ПРОТОКОЛ № , обрабатывается весь
файл построчно — правила п. 1–3 всё равно применяются, где найдутся подходящие строки.

  • В самом шаблоне протокола (.docx) в конце бланка можно поставить плейсхолдеры подписей комиссии
    (лучше одним фрагментом в run Word): {{ПРЕДСЕДАТЕЛЬ}} (также {{CHAIR}}) — двухстрочный блок с И.О. Фамилия
    и строкой «должность, подпись» / «И.О. Фамилия, подпись».
    Таблица «Председатель комиссии»: во второй ячейке строки — {{ПРЕДСЕДАТЕЛЬ_ФИО}}; подставляется только ФИО.
    Таблица «Члены комиссии»: отдельная строка шаблона с маркерами {{ЧЛЕНЫ_КОМИССИИ_ФИО}} во второй и третьей
    ячейках (два члена в одной строке). Строк шаблона добавляется ceil(N/2). Если маркер только в одной ячейке —
    режим как раньше: одна строка на одного члена. Старое {{ЧЛЕНЫ_КОМИССИИ}} / {{MEMBERS}} в абзацах даёт
    многострочный блок подписей (если маркер не только в таблице).
    Данные — «Приказ и комиссия»; отдельный файл подписей не используется.

────────────────────────────────────────
ИТОГ: какие «переменные» задавать в шаблоне явно
────────────────────────────────────────
  • ____  в строке «ПРОТОКОЛ № ____»     — номер в виде N-М-ГГ (см. п. 1)
  • строка вида «__» … 20__ г.           — дата
  • {{ПОДРАЗДЕЛЕНИЕ_ПРОВЕРКИ}}, {{УТВЕРДИЛ_ПРИКАЗ}} — см. п. 2а

Защита стандартных шаблонов (админка): режим «только чтение» в Word для default_protocol*.docx
в папке программы; сформированные протоколы не защищаются. Снятие — «Снять защиту (для правки)…».
  • абзац о комиссии по приказу          — см. п. 2а ( от, №, председателя, членов)
  • длинная строка _____ … ,             — тема в шапке (см. п. 3)
  • в таблице .docx: см. п. 5

────────────────────────────────────────
5) ТАБЛИЦА «Результат проверки знаний»  →  поля формы
────────────────────────────────────────
В шаблоне .docx — таблица: две верхние строки-шапка (п/п, Фамилия, … и 1…7),
затем пример блока «строка программы» + «строка 1.1 с маркером ФИО».

При формировании все строки ниже двух шапок удаляются и строятся заново:
число блоков = числу отмеченных программ (B, PP, SIZ, V). В каждом блоке
первая строка — номер программы (1, 2, …) и название: «Б» — лист B (столб. 2), PP — лист PP (столб. 2),
СИЗ — лист SIZ (столб. 2), «В» — по V_PROF; абзацы после проверки для PP/СИЗ — по-прежнему V_PROF (3–4);
далее строки сотрудников N.1…N.M. Для ПП, СИЗ, «В»: M — после объединения записей с одинаковым ФИО
(одна строка на человека; совмещение учитывается внутри записи). Для программы «Б»: M — по одной строке
на каждую должность: все выбранные строки списка (несколько строк с одним ФИО и разными должностями —
отдельные строки таблицы), плюс разворот полей «должность» и «совмещаемая» в отдельные строки.
Список работников — лист rabotnik.
  • результат проверки: оценка и регистрационные номера реестра Минтруда
    (поле в окне администрирования) — не путать с номером протокола.
    Несколько номеров — через запятую или с новой строки.
    Для «В»: по одной строке «№ …» на каждый непустой фрагмент (после объединения должностей);
    перед строкой каждого работника вставляется подзаголовок «Программа (В)» и в скобках — только его темы
    (как в шаблоне 123 прогВ.docx); в колонке «ФИО» — только ФИО.
    Для «Б»: по одной строке «№ …» на каждую строку таблицы (каждую должность); для PP и СИЗ — всегда одна строка «№ …».
  • тип проверки: плановая / внеплановая
  • объём обучения (ч): на листах B, PP и SIZ — столбец 3 в строке с названием (для справки в Excel);
    в таблице протокола к названию блоков ПП и СИЗ добавляется «(в объеме N ч.)», если в Excel указано число;
    для программы «Б» часы в таблицу не выводятся.
    Для «В» — столбец 4 (D) листа V; к шапке блока «В» для каждого работника добавляется «(в объеме N ч.)»
    по сумме часов по его фрагментам, сопоставленным со строками листа V.

────────────────────────────────────────
6) ПРОТОКОЛ ПО ТЕХНИЧЕСКИМ ВОПРОСАМ (лист Tech_V, отдельная комиссия)
────────────────────────────────────────
• В форме — флажок «Протокол по техническим вопросам» и список «Программа по Tech_V» (все строки листа
  с непустым наименованием); в файле программ — лист Tech_V: первая строка может быть заголовком
  (утвердил, программа/наименование, дата); иначе столбцы A, B, C — утвердил, программа, дата.
  Строка «…технических вопросов… по программе:» (в т.ч. с продолжением про подразделение) без маркеров
  дополняется выбранным наименованием программы с новой строки под этим абзацем;
  плейсхолдеры {{ТЕХ_ПРОГРАММА}} и др. подставляются как раньше.
• Шаблоны Word кладутся в папку с программой (рядом с main.py / .exe): default_protocol.docx и
  default_protocol_tehnicheskiy.docx. Редактируйте эти файлы — при формировании подхватывается актуальная
  копия с диска (отдельной «служебной» копии в базе нет). При включённом тех. протоколе без своего файла
  в настройках используется default_protocol_tehnicheskiy.docx, если он есть.
• В шаблоне .docx строка с «ПРОТОКОЛ №» может содержать «техническ…» (для обычного протокола такие строки пропускались).
• Реестр Минтруда для этого режима не используется; в графе результата — оценка и при необходимости ручной «Регистрационный номер».
• Приказ и состав комиссии — вкладка «Технич. вопросы» в окне «Приказ и комиссия…» (отдельные поля в базе).
• Плейсхолдеры: {{ТЕХ_ПРОГРАММА}}, {{ТЕХ_УТВЕРДИЛ}}, {{ТЕХ_ДАТА_УТВЕРЖДЕНИЯ}} и варианты {{TECH_PROGRAM}}, {{TECH_APPROVER}}, {{TECH_APPROVAL_DATE}}.
• Таблица «Результат проверки»: строка с наименованием программы в таблице не вставляется; нумерация не
  «1.1», а 1, 2, 3 … в графе п/п; работники — по алфавиту ФИО. Под двумя строками шапки должен быть образец
  строки с «ФИО» во второй ячейке (или первая строка данных копируется как шаблон).

Если маркеров «ФИО» в таблице нет, таблица не меняется.

Остальной текст шаблона — произвольный, его можно менять.
"""

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Подчёркивания-заполнители после «ПРОТОКОЛ №»: ASCII _, полноширинная ＿, double low line ‗
_US_PLACEHOLDER_RUN = re.compile(r"(?:_|＿|‗){3,}")

# Якорь бланка: слово «протокол» и знак номера (№ / # / U+2116), без требования подчёркиваний.
_PROTOCOL_NUMBER_ANCHOR_RE = re.compile(r"(?i)протокол\s*(?:№|#|\u2116)")


def _normalize_docx_plain_text(s: str) -> str:
    """Неразрывный пробел и невидимые символы — для надёжного поиска маркеров."""
    if not s:
        return s
    return (
        s.replace("\xa0", " ")
        .replace("\u200b", "")
        .replace("\ufeff", "")
        .replace("\u00ad", "")
    )


def _reject_protocol_line_as_false_positive(s: str) -> bool:
    """«по протоколу …» и т.п. — не бланк."""
    return bool(re.search(r"(?i)по\s+протокол", s))


def _line_has_protocol_form_marker(
    line: str, *, allow_technical_heading: bool = False
) -> bool:
    """Строка содержит якорь «ПРОТОКОЛ №» (регистр и пробелы гибкие; допускаются # и знак U+2116)."""
    s = _normalize_docx_plain_text(line)
    if _reject_protocol_line_as_false_positive(s):
        return False
    if not allow_technical_heading and "техническ" in s.lower():
        return False
    return bool(_PROTOCOL_NUMBER_ANCHOR_RE.search(s))


def _line_fills_protocol_number_slot(
    line: str, *, allow_technical_heading: bool = False
) -> bool:
    """Абзац, в котором подставляем номер протокола (есть якорь «ПРОТОКОЛ №»)."""
    return _line_has_protocol_form_marker(line, allow_technical_heading=allow_technical_heading)


def _is_protocol_title_line_without_underscores(
    line: str, *, allow_technical_heading: bool = False
) -> bool:
    """«… ПРОТОКОЛ №» в абзаце без символов подчёркивания (номер на следующей строке)."""
    s = _normalize_docx_plain_text(line)
    if _reject_protocol_line_as_false_positive(s):
        return False
    if not allow_technical_heading and "техническ" in s.lower():
        return False
    if _US_PLACEHOLDER_RUN.search(s) is not None:
        return False
    return bool(_PROTOCOL_NUMBER_ANCHOR_RE.search(s))


def _is_placeholder_only_underscore_line(line: str) -> bool:
    """Абзац только из группы подчёркиваний (и пробелов) — подстановка номера сюда."""
    s = _normalize_docx_plain_text(line).strip()
    m = _US_PLACEHOLDER_RUN.search(s)
    if not m:
        return False
    return not (s[: m.start()].strip() or s[m.end() :].strip())


def _replace_first_underscore_run_with_protocol_number(line: str, pn: str) -> str:
    s = _normalize_docx_plain_text(line)
    s2, n = _US_PLACEHOLDER_RUN.subn(pn, s, count=1)
    return s2 if n else line


def _apply_protocol_number_after_anchor(line: str, pn: str) -> str:
    """
    После «ПРОТОКОЛ №» (/#/U+2116): если дальше идут подчёркивания — заменить их на номер;
    иначе вставить номер сразу после знака (если его ещё нет).
    """
    if not pn:
        return line
    s = _normalize_docx_plain_text(line)
    m = _PROTOCOL_NUMBER_ANCHOR_RE.search(s)
    if not m:
        return line
    pos = m.end()
    tail = s[pos:]
    um = re.match(r"^(\s*)(?:_|＿|‗){3,}", tail)
    if um:
        return s[:pos] + um.group(1) + pn + tail[um.end() :]
    rest = tail.lstrip()
    if rest.startswith(pn):
        return s
    if not tail.strip():
        return s[:pos] + " " + pn
    return s[:pos] + " " + pn + tail


def _ooxml_sym_char(elem: ET.Element) -> str | None:
    for k, v in elem.attrib.items():
        if k.endswith("}char") and v:
            try:
                return chr(int(v, 16))
            except ValueError:
                return None
    return None


def _w_p_element_plain_text(p_el: ET.Element) -> str:
    """Текст абзаца как в Word: w:t, символы w:sym (например №), табуляция, перенос строки."""
    parts: list[str] = []
    W = _W_NS
    for node in p_el.iter():
        tag = node.tag
        if tag == f"{W}t":
            parts.append(node.text or "")
        elif tag == f"{W}sym":
            ch = _ooxml_sym_char(node)
            if ch:
                parts.append(ch)
        elif tag == f"{W}tab":
            parts.append("\t")
        elif tag == f"{W}br":
            parts.append("\n")
    return "".join(parts)


def _sorted_word_hf_xml_names(zf: zipfile.ZipFile, kind: str) -> list[str]:
    prefix = f"word/{kind}"
    return sorted(n for n in zf.namelist() if n.startswith(prefix) and n.endswith(".xml"))


def _paragraph_texts_from_hf_or_fragment_xml(xml_bytes: bytes) -> list[str]:
    # Локальный фрагмент .docx (zip), не сетевой поток — stdlib ElementTree (bandit B314).
    root = ET.fromstring(xml_bytes)  # nosec B314
    return [_w_p_element_plain_text(p) for p in root.iter(f"{_W_NS}p")]


def _iter_document_paragraphs_in_order(doc: Document):
    """Колонтитулы, затем все w:p тела (таблицы, надписи и т.д.) — в порядке XML."""
    seen_hf: set[int] = set()
    for section in doc.sections:
        hel = section.header._element
        hid = id(hel)
        if hid not in seen_hf:
            seen_hf.add(hid)
            for p_el in hel.iter(qn("w:p")):
                yield DocxParagraph(p_el, doc)
    for p_el in doc.element.body.iter(qn("w:p")):
        yield DocxParagraph(p_el, doc)
    seen_hf.clear()
    for section in doc.sections:
        fel = section.footer._element
        fid = id(fel)
        if fid not in seen_hf:
            seen_hf.add(fid)
            for p_el in fel.iter(qn("w:p")):
                yield DocxParagraph(p_el, doc)


def _all_document_paragraphs_ordered(doc: Document) -> list[DocxParagraph]:
    return list(_iter_document_paragraphs_in_order(doc))


_MONTHS_GEN = (
    "",
    "января",
    "февраля",
    "марта",
    "апреля",
    "мая",
    "июня",
    "июля",
    "августа",
    "сентября",
    "октября",
    "ноября",
    "декабря",
)
class ProtocolTemplateError(Exception):
    """Файл шаблона протокола отсутствует или не читается."""


def protocol_template_path() -> Path:
    """Шаблон протокола по охране труда: файл в папке с программой (корень проекта при запуске из исходников)."""
    return application_bundle_dir() / PROTOCOL_TEMPLATE_FILENAME


def protocol_technical_template_path() -> Path:
    """Шаблон протокола по техническим вопросам — там же, рядом с default_protocol.docx."""
    return application_bundle_dir() / PROTOCOL_TEMPLATE_TECH_FILENAME


def resolve_protocol_template_path(
    *,
    technical_protocol: bool,
    user_override: Path | None,
    technical_user_override: Path | None = None,
) -> Path:
    """
    Тех. протокол: сначала свой шаблон тех. (настройки), затем default_protocol_tehnicheskiy.docx, иначе ОТ-шаблон.
    Обычный режим: выбранный шаблон или default_protocol.docx.
    """
    if technical_protocol:
        if technical_user_override is not None:
            p = technical_user_override.expanduser().resolve()
            if p.is_file():
                return p
        tech = protocol_technical_template_path()
        if tech.is_file():
            return tech
        return protocol_template_path()
    if user_override is not None:
        return user_override.expanduser().resolve()
    return protocol_template_path()

def load_excel_first_nonempty_in_column(
    path: Path, sheet_name: str, column_one_based: int
) -> str:
    """Первая непустая ячейка в столбце листа, начиная со 2-й строки."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    if not path.is_file():
        return ""
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        names = {n.lower(): n for n in wb.sheetnames}
        sk = sheet_name.lower()
        if sk not in names:
            return ""
        ws = wb[names[sk]]
        max_r = min(ws.max_row or 200, 500)
        for r in range(2, max_r + 1):
            v = ws.cell(row=r, column=column_one_based).value
            if v is not None and str(v).strip():
                return str(v).strip()
        return ""
    finally:
        wb.close()


def load_v_prof_first_nonempty_in_column(path: Path, column_one_based: int) -> str:
    """Лист V_PROF: первая непустая ячейка в столбце, начиная со 2-й строки (кэш каталога программ)."""
    return get_cached_v_prof_column(path, column_one_based)


def _load_v_program_registry_rows(path: Path) -> list[tuple[str, str, str, float | None]]:
    """
    Лист V со 2-й строки: (ID гос. реестра — столбец A, B, C, часы D).
    Строки с пустым B пропускаются. Пустой A допускается (тогда ID — пустая строка).
    Данные из SQLite-кэша каталога программ (пересборка при смене файла).
    """
    return get_cached_v_registry_rows(path)


def _match_v_program_registry_row(
    fragment: str, rows: list[tuple[str, str, str, float | None]]
) -> tuple[str, str, float | None] | None:
    """
    Сопоставление фрагмента V_PROF со столбцом B листа V.
    Возвращает (наименование из C или B, ID из A, часы из D) или None.
    """
    return match_v_registry_fragment(fragment, rows)


def _best_v_program_title_for_fragment(
    fragment: str, rows: list[tuple[str, str, str, float | None]],
) -> str | None:
    """Наименование из листа V (столбец C) для фрагмента; сопоставление по столбцу B."""
    m = _match_v_program_registry_row(fragment, rows)
    return m[0] if m else None


def v_program_gos_registry_id_for_fragment(
    fragment: str, rows: list[tuple[str, str, str, float | None]],
) -> str | None:
    """ID программы в Гос. реестре (столбец A листа V) для совпавшего с B фрагмента; иначе None."""
    m = _match_v_program_registry_row(fragment, rows)
    if not m:
        return None
    gid = (m[1] or "").strip()
    return gid if gid else None


def _collect_unique_professions_ordered(persons: list[EmployeeRecord]) -> list[str]:
    """Уникальные должности из выбранных записей: основная и совмещаемая, порядок первого появления."""
    seen: set[str] = set()
    out: list[str] = []
    for p in persons:
        for pr in (p.profession, p.profession2):
            t = (pr or "").strip()
            if not t:
                continue
            k = _norm_profession_key(t)
            if k in seen:
                continue
            seen.add(k)
            out.append(t)
    return out


def _select_best_row_by_profession_col_a(
    path: Path,
    profession: str,
    sheet_name: str,
    max_col: int,
    tie_break_cols_1based: tuple[int, int] | None,
) -> tuple[Any, ...] | None:
    """
    Строка Excel: лучшее совпадение должности в столбце A (1-based).
    tie_break_cols_1based — (первая, последняя) колонки для суммы длин при равном совпадении;
    None — без доп. критерия.
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        return None
    if not path.is_file() or not profession.strip():
        return None
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        names = {n.lower(): n for n in wb.sheetnames}
        sk = sheet_name.lower()
        if sk not in names:
            return None
        ws = wb[names[sk]]
        target = _norm_profession_key(profession)
        skip_headers = {
            "профессия",
            "должность",
            "специальность",
            "фио",
            "п/п",
            "наименование",
        }
        best: tuple[int, int, tuple[Any, ...]] | None = None
        for row in ws.iter_rows(
            min_row=1,
            max_row=min(ws.max_row or 500, 2000),
            min_col=1,
            max_col=max_col,
            values_only=True,
        ):
            if not row or row[0] is None:
                continue
            c0 = _norm_profession_key(str(row[0]))
            if not c0 or c0 in skip_headers:
                continue
            match_score = 0
            if c0 == target:
                match_score = 3
            elif target in c0 or c0 in target:
                match_score = 2
            if match_score == 0:
                continue
            row_tuple = tuple(row)
            merged_len = 0
            if tie_break_cols_1based is not None:
                t_lo, t_hi = tie_break_cols_1based
                lo_idx = t_lo - 1
                hi_excl = min(len(row_tuple), t_hi)
                for idx in range(lo_idx, hi_excl):
                    cell = row_tuple[idx]
                    if cell is not None and str(cell).strip():
                        merged_len += len(str(cell).strip())
            cand = (match_score, merged_len, row_tuple)
            if best is None or cand[0] > best[0] or (
                cand[0] == best[0] and cand[1] > best[1]
            ):
                best = cand
        return best[2] if best else None
    finally:
        wb.close()


def _v_prof_select_best_row(path: Path, profession: str) -> tuple[Any, ...] | None:
    """
    Строка листа V_PROF (ячейки 1–19), лучшее совпадение должности в столбце A.
    """
    layout = v_prof_layout_for_path(path)
    tie: tuple[int, int] | None
    if layout.format == "matrix":
        lo = min((c for c, _ in layout.v_marker_columns), default=4)
        hi = layout.last_col + 1
        tie = (lo + 1, hi)
    else:
        tie = (V_PROF_PARTS_FIRST_COL, V_PROF_PARTS_LAST_COL)
    return _select_best_row_by_profession_col_a(
        path,
        profession,
        V_PROF_SHEET_NAME,
        layout.last_col + 1,
        tie,
    )


def _v_prof_anchor_line_from_row(row: tuple[Any, ...], column_one_based: int) -> str | None:
    """Одна ячейка заданного столбца строки V_PROF с префиксом «- »."""
    i = column_one_based - 1
    if i < 0 or i >= len(row):
        return None
    v = row[i]
    if v is None:
        return None
    s = str(v).strip()
    if not s:
        return None
    return "- " + s


def _v_prof_anchor_lines_program_v(
    row: tuple[Any, ...],
    layout: VProfLayout,
    v_reg_rows: list[tuple[str, str, str, float | None]],
) -> list[str]:
    """Программы «В» после проверки: шапка — столбец C листа V (матрица «Да»)."""
    titles = v_prof_row_program_header_titles(row, layout, v_reg_rows)
    return ["- " + t for t in titles if t]


def build_fg_lines_for_selected_programs(
    path: Path,
    program_keys: list[str],
    persons_raw: list[EmployeeRecord],
) -> list[str]:
    """
    V_PROF, столбец A — должность; для «Б»/PP/СИЗ — ячейки столбцов 2–4 (якорь после проверки).
    «В» — по одной строке на фрагмент; в абзацах после проверки — столбец C листа V.
    Таблица результатов — столбец B листа V (см. v_prof_row_program_table_fragments).
    persons_raw — список записей с теми же должностями, что должны попасть в шапку/якоря; для «Б» это
    развёрнутые строки блока «Б» (expand_persons_block_b_rows), а не объединение по ФИО.
    """
    profs = _collect_unique_professions_ordered(persons_raw)
    if not profs or not program_keys:
        return []
    layout = v_prof_layout_for_path(path)
    anchor_col_for_key: dict[str, int] = {
        "B": layout.col_b_one_based,
        "PP": layout.col_pp_one_based,
        "SIZ": layout.col_siz_one_based,
    }
    row_cache: dict[str, tuple[Any, ...] | None] = {}

    def row_for(pr: str) -> tuple[Any, ...] | None:
        if pr not in row_cache:
            row_cache[pr] = _v_prof_select_best_row(path, pr)
        return row_cache[pr]

    v_reg_rows: list[tuple[str, str, str, float | None]] = []
    if "V" in program_keys and path.is_file():
        v_reg_rows = get_cached_v_registry_rows(path)

    lines: list[str] = []
    for pkey in program_keys:
        if (pkey or "").strip().upper() == "TECH":
            continue
        for pr in profs:
            row = row_for(pr)
            if not row:
                continue
            if pkey == "V":
                lines.extend(_v_prof_anchor_lines_program_v(row, layout, v_reg_rows))
            else:
                col = anchor_col_for_key.get(pkey)
                if col is None:
                    continue
                s = _v_prof_anchor_line_from_row(row, col)
                if s:
                    lines.append(s)
    return lines


def _program_bullet_paragraph_indices_after_anchor(
    doc: Document, anchor: DocxParagraph
) -> tuple[int, int]:
    """
    Абзацы сразу после anchor, идущие подряд и начинающиеся с маркера списка (-, –, —).
    Возвращает (lo, hi) — полуинтервал индексов в doc.paragraphs; если пусто, lo == hi.
    """
    paras = list(doc.paragraphs)
    try:
        i = paras.index(anchor)
    except ValueError:
        return (0, 0)
    lo = i + 1
    hi = lo
    while hi < len(paras):
        t = paras[hi].text.replace("\xa0", " ").strip()
        if not t:
            break
        if re.match(r"^[-–—]", t):
            hi += 1
        else:
            break
    return (lo, hi)


def _is_anchor_paragraph_program_table(text: str) -> bool:
    """Абзац после которого вставляются строки из V_PROF (в т. ч. столб. 2 для «Б»)."""
    t = text.replace("\xa0", " ").lower()
    return (
        "провела" in t
        and "проверку" in t
        and "программе" in t
        and "охраны труда" in t
    )


def _insert_paragraph_after(
    paragraph: DocxParagraph, text: str = ""
) -> DocxParagraph:
    new_el = OxmlElement("w:p")
    paragraph._p.addnext(new_el)
    new_para = DocxParagraph(new_el, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_program_fg_lines_after_anchor(
    doc: Document,
    lines: list[str],
) -> int:
    """
    Вставляет абзацы после блока про проверку по программе(ам).
    Уже имеющиеся в шаблоне пункты (подряд после якоря, с «-»/«–»/«—») не дублируются;
    из списка на вставку убираются повторы. Новые строки добавляются после последнего
    такого пункта (или сразу после якоря, если пунктов ещё нет).
    """
    if not lines:
        return 0
    anchor: DocxParagraph | None = None
    for para in doc.paragraphs:
        if _is_anchor_paragraph_program_table(para.text):
            anchor = para
            break
    if anchor is None:
        return 0
    lo, hi = _program_bullet_paragraph_indices_after_anchor(doc, anchor)
    paras = list(doc.paragraphs)
    ref: DocxParagraph = paras[hi - 1] if hi > lo else anchor
    existing_keys: set[str] = set()
    for j in range(lo, hi):
        existing_keys.add(_fg_line_comparison_key(paras[j].text))
    to_insert: list[str] = []
    batch_seen: set[str] = set()
    for line in lines:
        k = _fg_line_comparison_key(line)
        if k in existing_keys or k in batch_seen:
            continue
        batch_seen.add(k)
        to_insert.append(line)
    if not to_insert:
        return 0
    for line in to_insert:
        ref = _insert_paragraph_after(ref, line)
    return len(to_insert)


def _table_employees_dedupe_by_fio(records: list[EmployeeRecord]) -> list[EmployeeRecord]:
    """
    Одна строка таблицы на одно ФИО: записи с одинаковым ФИО (несколько строк Excel и т. п.)
    объединяются; уникальные должности идут в profession и profession2 (остальные отбрасываются).
    """
    if not records:
        return []
    buckets: dict[str, list[EmployeeRecord]] = {}
    order_keys: list[str] = []
    for i, p in enumerate(records):
        key = _norm_profession_key(p.fio or "")
        if not key:
            key = f"__noid_{i}__"
        if key not in buckets:
            buckets[key] = []
            order_keys.append(key)
        buckets[key].append(p)
    out: list[EmployeeRecord] = []
    for key in order_keys:
        grp = buckets[key]
        base = grp[0]
        profs_ordered: list[str] = []

        def add_prof(s: str) -> None:
            t = (s or "").strip()
            if not t:
                return
            nk = _norm_profession_key(t)
            if not any(_norm_profession_key(x) == nk for x in profs_ordered):
                profs_ordered.append(t)

        add_prof(base.profession)
        add_prof(base.profession2)
        for extra in grp[1:]:
            add_prof(extra.profession)
            add_prof(extra.profession2)
        main_p = profs_ordered[0] if profs_ordered else (base.profession or "").strip()
        p2 = profs_ordered[1] if len(profs_ordered) > 1 else ""
        sub = (base.subdivision or "").strip()
        for extra in grp[1:]:
            if not sub:
                sub = (extra.subdivision or "").strip()
        snils = (base.snils or "").strip()
        if not snils:
            for extra in grp[1:]:
                t = (extra.snils or "").strip()
                if t:
                    snils = t
                    break
        out.append(
            EmployeeRecord(
                fio=base.fio,
                profession=main_p,
                subdivision=sub,
                profession2=p2,
                snils=snils,
            )
        )
    return out


def ensure_v_prof_cache_table(conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS v_prof_cache (
            file_path TEXT NOT NULL,
            profession_norm TEXT NOT NULL,
            file_mtime REAL NOT NULL,
            row_text TEXT NOT NULL,
            parts_json TEXT,
            PRIMARY KEY (file_path, profession_norm)
        )
        """
    )
    try:
        conn.execute("ALTER TABLE v_prof_cache ADD COLUMN parts_json TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        conn.execute("ALTER TABLE v_prof_cache ADD COLUMN parts_schema INTEGER DEFAULT 1")
    except sqlite3.OperationalError:
        pass


def _v_cache_invalidate_if_stale(conn: sqlite3.Connection, path_resolved: str, mtime: float) -> None:
    row = conn.execute(
        "SELECT file_mtime FROM v_prof_cache WHERE file_path = ? LIMIT 1",
        (path_resolved,),
    ).fetchone()
    if row is not None and row[0] != mtime:
        conn.execute("DELETE FROM v_prof_cache WHERE file_path = ?", (path_resolved,))


def v_cache_get_row_text(path: Path, profession: str) -> str | None:
    """Кэш строки V_PROF по должности (обновляется при смене времени файла Excel)."""
    if not path.is_file() or not profession.strip():
        return None
    p = str(path.resolve())
    mtime = path.stat().st_mtime
    key = _norm_profession_key(profession)
    with sqlite3.connect(database_path()) as conn:
        ensure_v_prof_cache_table(conn)
        _v_cache_invalidate_if_stale(conn, p, mtime)
        conn.commit()
        row = conn.execute(
            """
            SELECT row_text FROM v_prof_cache
            WHERE file_path = ? AND profession_norm = ? AND file_mtime = ?
            """,
            (p, key, mtime),
        ).fetchone()
        if row:
            return row[0]
    return None


def v_cache_put_row_text(path: Path, profession: str, text: str, parts: list[str] | None = None) -> None:
    p = str(path.resolve())
    mtime = path.stat().st_mtime
    key = _norm_profession_key(profession)
    pj = json.dumps(parts, ensure_ascii=False) if parts is not None else None
    with sqlite3.connect(database_path()) as conn:
        ensure_v_prof_cache_table(conn)
        conn.execute(
            """
            INSERT OR REPLACE INTO v_prof_cache
            (file_path, profession_norm, file_mtime, row_text, parts_json, parts_schema)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (p, key, mtime, text, pj, V_PROF_PARTS_CACHE_SCHEMA),
        )
        conn.commit()


def v_cache_get_parts_json(path: Path, profession: str) -> list[str] | None:
    if not path.is_file() or not profession.strip():
        return None
    p = str(path.resolve())
    mtime = path.stat().st_mtime
    key = _norm_profession_key(profession)
    with sqlite3.connect(database_path()) as conn:
        ensure_v_prof_cache_table(conn)
        _v_cache_invalidate_if_stale(conn, p, mtime)
        conn.commit()
        row = conn.execute(
            """
            SELECT parts_json, COALESCE(parts_schema, 0) FROM v_prof_cache
            WHERE file_path = ? AND profession_norm = ? AND file_mtime = ?
            """,
            (p, key, mtime),
        ).fetchone()
        if row and row[0] and int(row[1]) == V_PROF_PARTS_CACHE_SCHEMA:
            try:
                data = json.loads(row[0])
                if isinstance(data, list):
                    return [str(x) for x in data]
            except json.JSONDecodeError:
                return None
    return None


def _v_prof_find_best_row_parts(path: Path, profession: str) -> list[str]:
    """Фрагменты «В» для таблицы протокола: столбец B листа V (матрица «Да» в V_PROF)."""
    row = _v_prof_select_best_row(path, profession)
    if not row:
        return []
    layout = v_prof_layout_for_path(path)
    v_rows = get_cached_v_registry_rows(path) if path.is_file() else []
    return v_prof_row_program_table_fragments(row, layout, v_rows)


def read_v_prof_row_parts_list(path: Path, profession: str) -> list[str]:
    """Упорядоченный список непустых ячеек столбцов 5–19 для должности (A); с кэшем."""
    cached = v_cache_get_parts_json(path, profession)
    if cached is not None:
        return cached
    parts = _v_prof_find_best_row_parts(path, profession)
    joined = ", ".join(parts)
    v_cache_put_row_text(path, profession, joined, parts)
    return parts


def read_v_prof_row_joined(path: Path, profession: str) -> str:
    """Сцепка ячеек V_PROF через запятую (для совместимости)."""
    return ", ".join(read_v_prof_row_parts_list(path, profession))


def expand_persons_block_b_rows(persons: list[EmployeeRecord]) -> list[EmployeeRecord]:
    """
    Строки блока «Б» в таблице протокола: по одной на каждую должность.
    Для каждой записи из списка (строка Excel / выбор) поля profession и profession2
    дают до двух разных строк; несколько строк с одним ФИО и разными должностями
    сохраняются все (не схлопываются).
    """
    out: list[EmployeeRecord] = []
    for emp in persons:
        profs: list[str] = []
        seen: set[str] = set()
        for pr in (emp.profession, emp.profession2):
            t = (pr or "").strip()
            if not t:
                continue
            nk = _norm_profession_key(t)
            if nk in seen:
                continue
            seen.add(nk)
            profs.append(t)
        sn = (emp.snils or "").strip()
        sub = (emp.subdivision or "").strip()
        fio = emp.fio or ""
        if not profs:
            out.append(
                EmployeeRecord(
                    fio=fio,
                    profession="",
                    subdivision=sub,
                    profession2="",
                    snils=sn,
                )
            )
            continue
        for pr in profs:
            out.append(
                EmployeeRecord(
                    fio=fio,
                    profession=pr,
                    subdivision=sub,
                    profession2="",
                    snils=sn,
                )
            )
    return out


def expand_persons_for_separate_profession_rows(persons: list[EmployeeRecord]) -> list[EmployeeRecord]:
    """Совместимость: то же, что expand_persons_block_b_rows."""
    return expand_persons_block_b_rows(persons)


def raw_employee_rows_same_fio_as(
    raw_selection: list[EmployeeRecord],
    emp: EmployeeRecord,
) -> list[EmployeeRecord]:
    """Все выбранные строки с тем же ФИО, что у emp (ключ как при объединении по ФИО)."""
    key = _norm_profession_key(emp.fio or "")
    if not key:
        return [emp]
    out = [r for r in raw_selection if _norm_profession_key(r.fio or "") == key]
    return out if out else [emp]


def v_program_merged_parts_for_raw_employee(path: Path, emp: EmployeeRecord) -> list[str]:
    """
    Все фрагменты V_PROF для сотрудника из исходной записи Excel: основная + совмещаемая
    должность без повторов по тексту (порядок: сначала основная строка V_PROF, затем уникальное из второй).
    """
    primary = read_v_prof_row_parts_list(path, emp.profession)
    seen = {_norm_profession_key(x) for x in primary}
    out = list(primary)
    p2 = (emp.profession2 or "").strip()
    if not p2:
        return out
    for x in read_v_prof_row_parts_list(path, p2):
        nk = _norm_profession_key(x)
        if nk not in seen:
            seen.add(nk)
            out.append(x)
    return out


def v_program_ordered_unique_parts_global(
    path: Path, persons_raw: list[EmployeeRecord],
) -> list[str]:
    """Уникальные фрагменты программы «В» по всем выбранным сотрудникам (порядок сохраняется)."""
    seen: set[str] = set()
    ordered: list[str] = []
    for emp in persons_raw:
        for part in v_program_merged_parts_for_raw_employee(path, emp):
            nk = _norm_profession_key(part)
            if nk not in seen:
                seen.add(nk)
                ordered.append(part)
    return ordered


def resolve_v_program_inner_text_global(
    path: Path, persons_raw: list[EmployeeRecord], fallback: str
) -> str:
    """Один объединённый текст фрагментов В по всем выбранным сотрудникам (без повторов)."""
    parts = v_program_ordered_unique_parts_global(path, persons_raw)
    return ", ".join(parts) if parts else fallback


def _v_line_number_and_body(fragment: str) -> tuple[int | None, str]:
    """
    Ведущий номер пункта в строке V_PROF («2. …», «3) …») и текст без него.
    Если номера нет или после номера пусто — (None, исходная строка).
    """
    s = (fragment or "").replace("\xa0", " ").strip()
    if not s:
        return None, ""
    m = re.match(r"^(\d+)\s*[\.\)]\s*(.+)$", s, re.DOTALL)
    if not m:
        return None, s
    try:
        n = int(m.group(1))
    except ValueError:
        return None, s
    body = (m.group(2) or "").strip()
    if not body:
        return None, s
    return n, body


def v_fragment_labels_and_hints(parts: list[str]) -> tuple[list[int], list[str]]:
    """
    Для каждого фрагмента «В»: метка в шапке/в ячейке результата (номер из текста или 1,2,… по порядку)
    и строка для сопоставления с реестром (без ведущего «n.»).
    """
    labels: list[int] = []
    hints: list[str] = []
    for i, raw in enumerate(parts):
        p = (raw or "").strip()
        if not p:
            labels.append(i + 1)
            hints.append("")
            continue
        n, body = _v_line_number_and_body(p)
        labels.append(n if n is not None else i + 1)
        hints.append(body if n is not None else p)
    return labels, hints


def format_v_program_table_block_title(parts: list[str], fallback: str) -> str:
    """Заголовок блока «В»: перечень с номерами как в V_PROF (2., 3., …) или по порядку 1., 2., …"""
    clean = [p.strip() for p in parts if (p or "").strip()]
    if not clean:
        fb = (fallback or "").strip() or fallback
        return f"Программа (В)\n({fb})"
    lines_out: list[str] = []
    for i, p in enumerate(clean):
        n, body = _v_line_number_and_body(p)
        lbl = n if n is not None else i + 1
        text = body if n is not None else p
        lines_out.append(f"{lbl}. {text}")
    numbered = "\n".join(lines_out)
    return f"Программа (В)\n({numbered})"


def _protocol_block_title_with_hours(
    title: str, hours: float | None, *, multiline: bool
) -> str:
    """Добавляет объём в ч. в скобках «(в объеме N ч.)»; multiline — с новой строки перед скобками."""
    if hours is None or hours != hours:
        return title
    lab = format_training_hours_ru(hours)
    suffix = f"(в объеме {lab} ч.)"
    if multiline:
        return f"{title.rstrip()}\n{suffix}"
    return f"{title.rstrip()} {suffix}"


def _v_person_training_hours_sum(
    v_parts: list[str],
    v_rows: list[tuple[str, str, str, float | None]],
) -> float | None:
    """Сумма часов из столбца D листа V по фрагментам «В» этого работника (сопоставление по столбцу B)."""
    tot = 0.0
    any_h = False
    for raw in v_parts:
        raw = (raw or "").strip()
        if not raw:
            continue
        m = match_v_registry_fragment(raw, v_rows)
        if m and m[2] is not None:
            tot += float(m[2])
            any_h = True
    return tot if any_h else None


def _format_n_registry_lines_with_v_labels(
    grade: str,
    registry_no: str,
    line_labels: list[int],
    file_line_tokens: list[str] | None = None,
) -> str:
    """Строки «n. № …», где n — номер из шапки блока «В» (2., 3., …), не всегда с 1."""
    n = len(line_labels)
    if n <= 0:
        return _format_table_result_grade(grade, registry_no)
    g = (grade or "").strip()
    titled = g[:1].upper() + g[1:] if len(g) > 1 else (g.upper() if g else "")
    regs = merge_registry_tokens(registry_no, n, file_line_tokens)
    lines: list[str] = []
    if titled:
        lines.append(f"{titled},")
    for i in range(n):
        r = regs[i] if i < len(regs) else ""
        token = _format_registry_number_token(r) if r else "№"
        lbl = line_labels[i] if i < len(line_labels) else i + 1
        lines.append(f"{lbl}. {token}")
    return "\n".join(lines)


def _format_v_result_cell(
    grade: str,
    registry_no: str,
    v_parts: list[str],
    v_file_tokens: list[str] | None = None,
) -> str:
    """
    Результат для блока В: оценка, затем строки «n. № …» с тем же n, что в шапке блока
    (если в V_PROF указано «2. …», то «2. № …», а не «1. № …»).
    """
    if not v_parts:
        return _format_table_result_grade(grade, registry_no)
    labels, _ = v_fragment_labels_and_hints(v_parts)
    if len(labels) != len(v_parts):
        labels = [i + 1 for i in range(len(v_parts))]
    return _format_n_registry_lines_with_v_labels(
        grade, registry_no, labels, file_line_tokens=v_file_tokens
    )


def _protocol_program_fallback_title(program_key: str) -> str:
    for k, _, fb in PROTOCOL_PROGRAM_DEFS:
        if k == program_key:
            return fb
    return ""


def build_protocol_header_theme_text(
    path: Path | None,
    program_keys: list[str],
    program_titles: list[str],
    persons_for_v_prof: list[EmployeeRecord],
) -> str:
    """
    Текст шапки протокола (абзац с «__» под номером):
    «Б» / ПП / СИЗ — ячейки V_PROF по должностям (в матричном формате: D / B / C);
    «В» — наименования с листа V (столбец C) по матрице «Да» и номерам в шапке V_PROF.
    Таблица «В» — столбец B листа V; «Б»/ПП/СИЗ в таблице — листы B, PP, SIZ.
    """
    clean_titles = [(t or "").strip() for t in program_titles]
    if not program_keys:
        return "; ".join(t for t in clean_titles if t)
    if path is None or not path.is_file():
        return "; ".join(t for t in clean_titles if t)

    layout = v_prof_layout_for_path(path)
    profs = _collect_unique_professions_ordered(persons_for_v_prof)
    if not profs:
        return "; ".join(t for t in clean_titles if t)

    row_cache: dict[str, tuple[Any, ...] | None] = {}

    def row_for(pr: str) -> tuple[Any, ...] | None:
        if pr not in row_cache:
            row_cache[pr] = _v_prof_select_best_row(path, pr)
        return row_cache[pr]

    v_reg_rows: list[tuple[str, str, str, float | None]] = []
    if any((k or "").strip().upper() == "V" for k in program_keys):
        v_reg_rows = get_cached_v_registry_rows(path)

    anchor_col_for_key: dict[str, int] = {
        "B": layout.col_b_one_based,
        "PP": layout.col_pp_one_based,
        "SIZ": layout.col_siz_one_based,
    }
    keys_titles = list(zip(program_keys, program_titles))
    theme_parts: list[str] = []

    for key, fallback_title in keys_titles:
        pk = (key or "").strip().upper()
        if pk == "V":
            seen_v: set[str] = set()
            for pr in profs:
                row = row_for(pr)
                if not row:
                    continue
                for t in v_prof_row_program_header_titles(row, layout, v_reg_rows):
                    t = (t or "").strip()
                    if not t:
                        continue
                    nk = _norm_profession_key(t)
                    if nk in seen_v:
                        continue
                    seen_v.add(nk)
                    theme_parts.append(t)
            if not theme_parts and (fallback_title or "").strip():
                theme_parts.append((fallback_title or "").strip())
        elif pk in anchor_col_for_key:
            seen_a: set[str] = set()
            added = False
            for pr in profs:
                row = row_for(pr)
                if not row:
                    continue
                line = _v_prof_anchor_line_from_row(row, anchor_col_for_key[pk])
                if not line:
                    continue
                core = line.lstrip("-").strip()
                if not core:
                    continue
                nk = _norm_profession_key(core)
                if nk in seen_a:
                    continue
                seen_a.add(nk)
                theme_parts.append(core)
                added = True
            if not added and (fallback_title or "").strip():
                theme_parts.append((fallback_title or "").strip())
        else:
            ts = (fallback_title or "").strip()
            if ts:
                theme_parts.append(ts)

    return "; ".join(theme_parts).strip()


def _profession_cell_primary_only(emp: EmployeeRecord) -> str:
    """Только основная должность (совмещение в таблице не дублируется второй строкой)."""
    return (emp.profession or "").strip()


def _rebuild_registry_rows_for_program(pkey: str, emp: EmployeeRecord) -> int:
    """Сколько строк «№» в ячейке результата для блока программы."""
    pk = (pkey or "").strip().upper()
    has2 = bool((emp.profession2 or "").strip())
    if pk == "B":
        # Блок «Б»: после expand_persons_block_b_rows — одна должность на строку, одна строка «№»;
        # запасной вариант для записи с двумя должностями в одной строке.
        return 2 if has2 else 1
    if pk in ("PP", "SIZ"):
        return 1
    return 1

def existing_per_employee_docx_in_folder(
    out_dir: Path,
    persons: list[EmployeeRecord],
    base_no: int,
    date_str: str,
) -> list[Path]:
    """Пути к уже существующим DOCX с теми же именами, что получит партия «по одному»."""
    o = Path(out_dir)
    found: list[Path] = []
    for i, emp in enumerate(persons):
        n = base_no + i
        fname = default_protocol_save_filename(
            str(n),
            date_str,
            ".docx",
            person_suffix=format_fio_filename_surname_initials(emp.fio),
        )
        p = o / fname
        if p.is_file():
            found.append(p)
    return found
def _load_form_lines_from_txt(path: Path) -> list[str]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as e:
        raise ProtocolTemplateError(
            f"Текстовый шаблон должен быть в кодировке UTF-8:\n{path}"
        ) from e
    lines = text.splitlines()
    try:
        return _slice_form_template(lines)
    except ValueError:
        return lines


def load_protocol_form_lines(path: Path | None = None) -> list[str]:
    p = protocol_template_path() if path is None else Path(path).expanduser().resolve()
    if not p.is_file():
        hint = (
            f"Положите файл «{PROTOCOL_TEMPLATE_FILENAME}» в папку с программой или выберите шаблон."
            if path is None
            else ""
        )
        raise ProtocolTemplateError(
            f"Файл шаблона не найден:\n{p}" + (f"\n\n{hint}" if hint else "")
        )
    suf = p.suffix.lower()
    if suf == ".txt":
        try:
            return _load_form_lines_from_txt(p)
        except OSError as e:
            raise ProtocolTemplateError(f"Не удалось прочитать шаблон:\n{e}") from e
    if suf in (".docx", ".docm"):
        try:
            return _load_form_lines_from_docx(p)
        except zipfile.BadZipFile as e:
            raise ProtocolTemplateError(
                f"Файл не является корректным документом Word (.docx):\n{p}"
            ) from e
        except ET.ParseError as e:
            raise ProtocolTemplateError(f"Не удалось разобрать XML в шаблоне:\n{p}") from e
        except ValueError as e:
            raise ProtocolTemplateError(str(e)) from e
        except OSError as e:
            raise ProtocolTemplateError(f"Не удалось прочитать шаблон:\n{e}") from e
    raise ProtocolTemplateError(
        f"Поддерживаются шаблоны .docx / .docm и .txt; указан файл: {p.suffix or '(без расширения)'}"
    )


def _paragraphs_from_docx(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        lines: list[str] = []
        for name in _sorted_word_hf_xml_names(zf, "header"):
            lines.extend(_paragraph_texts_from_hf_or_fragment_xml(zf.read(name)))
        root = ET.fromstring(zf.read("word/document.xml"))  # nosec B314
        body = root.find(f".//{_W_NS}body")
        if body is None:
            raise ValueError("В документе Word не найден раздел body.")
        for p in body.iter(f"{_W_NS}p"):
            lines.append(_w_p_element_plain_text(p))
        for name in _sorted_word_hf_xml_names(zf, "footer"):
            lines.extend(_paragraph_texts_from_hf_or_fragment_xml(zf.read(name)))
    return lines


def _document_line_has_protocol_number_anchor(line: str) -> bool:
    """Якорь «ПРОТОКОЛ №» для границ бланка (включая заголовки с «техническ…»)."""
    s = _normalize_docx_plain_text(line)
    if _reject_protocol_line_as_false_positive(s):
        return False
    return bool(_PROTOCOL_NUMBER_ANCHOR_RE.search(s))


def _find_form_template_bounds(lines: list[str]) -> tuple[int, int]:
    start = None
    for i, line in enumerate(lines):
        if _document_line_has_protocol_number_anchor(line):
            start = i
            break
    if start is None:
        for i in range(len(lines) - 1):
            for glue in ("", " ", "\n"):
                combo = glue.join((lines[i], lines[i + 1]))
                if _document_line_has_protocol_number_anchor(combo):
                    start = i
                    break
            if start is not None:
                break
    if start is None:
        raise ValueError(
            "В файле шаблона не найден бланк: в тексте, таблице или колонтитуле должна быть строка "
            "с фрагментом «ПРОТОКОЛ №» (допускаются латинский # или знак № U+2116; любой регистр). "
            "Номер подставляется после знака или вместо подчёркиваний сразу после него. "
            "Пример шаблона: default_protocol.docx в папке программы."
        )
    end = len(lines)
    for j in range(start, len(lines)):
        if lines[j].startswith("Приложение"):
            end = j
            break
    return start, end


def _slice_form_template(lines: list[str]) -> list[str]:
    start, end = _find_form_template_bounds(lines)
    return lines[start:end]


def _load_form_lines_from_docx(path: Path) -> list[str]:
    raw = _paragraphs_from_docx(path)
    try:
        return _slice_form_template(raw)
    except ValueError as e:
        raise ValueError(f"{e}\n\nПроверяемый файл:\n{path.resolve()}") from e


def _split_dmy_triplet(date_str: str) -> tuple[str, str, str] | None:
    """День, месяц, год как строки; разделитель . / или - (порядок ДД.ММ.ГГГГ)."""
    s = date_str.strip()
    for sep in (".", "/", "-"):
        if sep not in s:
            continue
        parts = [p.strip() for p in s.split(sep)]
        if len(parts) == 3:
            return (parts[0], parts[1], parts[2])
    return None


def _format_date_protocol_line(date_str: str) -> str:
    sp = _split_dmy_triplet(date_str)
    if sp is None:
        return date_str
    d_s, m_s, y = sp
    try:
        d, m = int(d_s), int(m_s)
        if len(y) == 2:
            y = f"20{y}"
        if not (1 <= m <= 12):
            return date_str
        return f"«{d}» {_MONTHS_GEN[m]} {y} г."
    except ValueError:
        return date_str


def _format_tech_approval_date_for_template(raw: str) -> str:
    """Дата утверждения программы (Tech_V): ДД.ММ.ГГГГ → как дата протокола словами; иначе как в Excel."""
    t = (raw or "").strip()
    if not t:
        return ""
    if re.match(r"^\d{1,2}\.\d{1,2}\.\d{2,4}$", t):
        return _format_date_protocol_line(t)
    if re.match(r"^\d{4}-\d{2}-\d{2}", t):
        y, m, d = t[:10].split("-")
        return _format_date_protocol_line(f"{d}.{m}.{y}")
    return t


def _parse_dmy_month_year(date_str: str) -> tuple[int, int] | None:
    """ДД.ММ.ГГГГ или ДД.ММ.ГГ (разделитель . / -) → (месяц 1–12, полный год). Некорректная дата — None."""
    sp = _split_dmy_triplet(date_str)
    if sp is None:
        return None
    d_s, m_s, y = sp
    try:
        d, m = int(d_s), int(m_s)
        if len(y) == 2:
            y_full = int(f"20{y}")
        else:
            y_full = int(y)
        if not (1 <= m <= 12 and 1 <= d <= 31):
            return None
        return m, y_full
    except ValueError:
        return None


def format_protocol_number_for_template(protocol_no: str, date_str: str) -> str:
    """
    Подстановка вместо ____ в строке «ПРОТОКОЛ №»: <номер>-<месяц>-<год_2_цифры>.
    Месяц без ведущего нуля. Пустой номер — пустая строка (не подставлять).
    """
    raw = (protocol_no or "").strip()
    if not raw:
        return ""
    my = _parse_dmy_month_year(date_str)
    if my is None:
        return raw
    month, year_full = my
    yy = year_full % 100
    return f"{raw}-{month}-{yy:02d}"


def _is_program_underscore_line(line: str) -> bool:
    s = line.strip()
    if not s.endswith(","):
        return False
    core = s[:-1].replace("_", "").strip()
    return len(s) > 20 and not core


def _is_technical_program_intro_paragraph(line: str, *, technical_protocol: bool) -> bool:
    lo = line.lower().replace("\xa0", " ")
    if not (
        "провела" in lo
        and "проверку" in lo
        and "программе" in lo
    ):
        return False
    if "техническ" in lo:
        return True
    if technical_protocol and "охраны труда" in lo:
        return True
    return False


def _line_has_tech_v_placeholder_marker(line: str) -> bool:
    u = line.upper().replace("\xa0", " ")
    return any(
        m in u
        for m in (
            "{{ТЕХ_ПРОГРАММА}}",
            "{{TECH_PROGRAM}}",
            "{{ТЕХ_УТВЕРДИЛ}}",
            "{{TECH_APPROVER}}",
            "{{ТЕХ_ДАТА_УТВЕРЖДЕНИЯ}}",
            "{{TECH_APPROVAL_DATE}}",
        )
    )


def _inline_tech_program_in_intro_line(line: str, program_name: str) -> str:
    """
    Без маркеров {{ТЕХ_ПРОГРАММА}}: наименование программы — с новой строки под фразой
    «…провела проверку… технических вопросов… по программе:» (в т.ч. если далее текст
    про подразделение, например «… для работников отдела главного энергетика»).
    """
    pn = (program_name or "").strip()
    if not pn or _line_has_tech_v_placeholder_marker(line):
        return line
    if pn in line:
        return line
    s = line.rstrip()
    if "\n" in s:
        last = s.split("\n")[-1].strip()
        if last == pn:
            return line
    return s + "\n" + pn


def _fill_protocol_form(
    form: list[str],
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
) -> list[str]:
    """Подстановка полей по маркерам шаблона. Список маркеров — PROTOCOL_TEMPLATE_VARIABLES_DOC."""
    payload = build_commission_template_payload(_format_date_protocol_line)
    result: list[str] = []
    idx = 0
    while idx < len(form):
        line = form[idx]
        nxt = form[idx + 1] if idx + 1 < len(form) else None
        chunks: list[str]

        if (
            nxt is not None
            and _is_protocol_title_line_without_underscores(line)
            and _is_placeholder_only_underscore_line(nxt)
        ):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            chunks = [
                line,
                _replace_first_underscore_run_with_protocol_number(nxt, pn) if pn else nxt,
            ]
            idx += 2
        elif _line_fills_protocol_number_slot(line):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            chunks = [_apply_protocol_number_after_anchor(line, pn) if pn else line]
            idx += 1
        elif line.strip().startswith("«__»") and line.strip().endswith("г."):
            chunks = [_format_date_protocol_line(date_str)]
            idx += 1
        elif _is_program_underscore_line(line):
            chunks = [(theme.strip() + ",") if theme.strip() else line]
            idx += 1
        else:
            chunks = [line]
            idx += 1

        for out_line in chunks:
            result.append(
                apply_commission_insertions_to_line(
                    out_line,
                    date_words=payload["date_words"],
                    order_no=payload["order_no"],
                    chair_gen=payload["chair"],
                    members_gen=payload["members"],
                )
            )
    return result
def _replace_paragraph_text_preserve_style(paragraph: DocxParagraph, new_text: str) -> None:
    """Заменяет текст абзаца: стиль абзаца сохраняется (clear), оформление знаков — с первого run."""
    r0 = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    new_run = paragraph.add_run(new_text)
    if r0 is not None and r0._r.rPr is not None:
        new_run._r.insert(0, deepcopy(r0._r.rPr))


def _replace_paragraph_text_preserve_style_multiline(paragraph: DocxParagraph, new_text: str) -> None:
    """Как _replace_paragraph_text_preserve_style, но переводы строк — как мягкий перенос в Word."""
    r0 = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    parts = (new_text or "").split("\n")
    for i, part in enumerate(parts):
        if i:
            br_run = paragraph.add_run()
            br_run.add_break(WD_BREAK.LINE)
            if r0 is not None and r0._r.rPr is not None:
                br_run._r.insert(0, deepcopy(r0._r.rPr))
        new_run = paragraph.add_run(part)
        if r0 is not None and r0._r.rPr is not None:
            new_run._r.insert(0, deepcopy(r0._r.rPr))


def _apply_run_pr_from_template(template_run: DocxRun | None, run: DocxRun) -> None:
    if template_run is not None and template_run._r.rPr is not None:
        run._r.insert(0, deepcopy(template_run._r.rPr))


def _ensure_commission_header_tab_stops(paragraph: DocxParagraph, *, two_columns: bool) -> None:
    """Позиции табуляции как в бланке; для членов — вторая колонка."""
    tabs = paragraph.paragraph_format.tab_stops
    existing = {int(t.position) for t in tabs}
    if _COMMISSION_HEADER_TAB1_TWIPS not in existing:
        tabs.add_tab_stop(Twips(_COMMISSION_HEADER_TAB1_TWIPS), WD_TAB_ALIGNMENT.LEFT)
    if two_columns:
        pos2 = int(_COMMISSION_HEADER_TAB2)
        if pos2 not in existing:
            tabs.add_tab_stop(_COMMISSION_HEADER_TAB2, WD_TAB_ALIGNMENT.LEFT)


def _line_inserts_commission_members(orig_line: str, new_text: str) -> bool:
    if new_text == orig_line or "член" not in (orig_line or "").lower():
        return False
    return new_text.count("\n") > (orig_line or "").count("\n")


def _line_inserts_commission_chair(
    orig_line: str, new_text: str, chair_gen: str
) -> bool:
    if not (chair_gen or "").strip() or new_text == orig_line:
        return False
    if "председател" not in (orig_line or "").lower():
        return False
    if chair_gen.strip() in (new_text or ""):
        return True
    return new_text.count("\n") > (orig_line or "").count("\n")


def _add_commission_member_run(
    paragraph: DocxParagraph,
    text: str,
    *,
    template_run: DocxRun | None,
    underline: bool,
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    _apply_run_pr_from_template(template_run, run)
    if underline:
        run.font.underline = WD_UNDERLINE.SINGLE


def _replace_paragraph_commission_members(
    paragraph: DocxParagraph,
    anchor_line: str,
    members_text: str,
) -> None:
    """
    Абзац «членов …»: префикс без подчёркивания; ФИО и должности — с подчёркиванием,
    в два столбца с табуляцией по позициям шаблона.
    """
    r0 = paragraph.runs[0] if paragraph.runs else None
    prefix = commission_members_anchor_prefix(anchor_line)
    rows = parse_commission_members_two_column_text(members_text)
    two_columns = any(r for _, r in rows)
    paragraph.clear()
    _ensure_commission_header_tab_stops(paragraph, two_columns=two_columns)
    if prefix:
        _add_commission_member_run(
            paragraph, prefix, template_run=r0, underline=False
        )
    for i, (left, right) in enumerate(rows):
        if i or prefix:
            br_run = paragraph.add_run()
            br_run.add_break(WD_BREAK.LINE)
            _apply_run_pr_from_template(r0, br_run)
        if left and right:
            _add_commission_member_run(
                paragraph, left, template_run=r0, underline=True
            )
            tab_run = paragraph.add_run("\t")
            _apply_run_pr_from_template(r0, tab_run)
            _add_commission_member_run(
                paragraph, right, template_run=r0, underline=True
            )
        elif left:
            _add_commission_member_run(
                paragraph, left, template_run=r0, underline=True
            )
        elif right:
            _add_commission_member_run(
                paragraph, right, template_run=r0, underline=True
            )


def _replace_paragraph_commission_chair(
    paragraph: DocxParagraph,
    anchor_line: str,
    chair_text: str,
) -> None:
    """
    Абзац «председателя …»: слово «председателя» без подчёркивания;
    ФИО и должность (род. п.) — с подчёркиванием, как у членов комиссии.
    """
    r0 = paragraph.runs[0] if paragraph.runs else None
    prefix = commission_chair_anchor_prefix(anchor_line)
    body = (chair_text or "").strip()
    paragraph.clear()
    _ensure_commission_header_tab_stops(paragraph, two_columns=False)
    if prefix:
        _add_commission_member_run(
            paragraph, prefix, template_run=r0, underline=False
        )
    if not body:
        return
    for i, part in enumerate(body.split("\n")):
        chunk = part.strip()
        if not chunk:
            continue
        if i or prefix:
            br_run = paragraph.add_run()
            br_run.add_break(WD_BREAK.LINE)
            _apply_run_pr_from_template(r0, br_run)
        if "," in chunk:
            left, _, right = chunk.partition(",")
            left = left.strip()
            right = right.strip()
            if left:
                _add_commission_member_run(
                    paragraph, left, template_run=r0, underline=True
                )
            if left and right:
                _add_commission_member_run(
                    paragraph, ", ", template_run=r0, underline=False
                )
            if right:
                _add_commission_member_run(
                    paragraph, right, template_run=r0, underline=True
                )
        else:
            _add_commission_member_run(
                paragraph, chunk, template_run=r0, underline=True
            )


def _apply_commission_paragraph_replacement(
    paragraph: DocxParagraph,
    orig_line: str,
    new_text: str,
    *,
    members_gen: str,
    chair_gen: str = "",
) -> None:
    if members_gen and _line_inserts_commission_members(orig_line, new_text):
        _replace_paragraph_commission_members(paragraph, orig_line, members_gen)
        return
    if chair_gen and _line_inserts_commission_chair(orig_line, new_text, chair_gen):
        _replace_paragraph_commission_chair(paragraph, orig_line, chair_gen)
        return
    if "\n" in new_text:
        _replace_paragraph_text_preserve_style_multiline(paragraph, new_text)
    else:
        _replace_paragraph_text_preserve_style(paragraph, new_text)


def _iter_all_paragraphs_in_document(doc: Document):
    """Все абзацы документа, включая ячейки таблиц (объединённые ячейки не дублируются)."""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for p in cell.paragraphs:
                    yield p


def fill_commission_chair_fio_table(
    doc: Document, *, kind: str = COMMISSION_KIND_OT
) -> None:
    """Ячейки с {{ПРЕДСЕДАТЕЛЬ_ФИО}} — только ФИО председателя из БД (без должности)."""
    _, _, chair, _ = load_commission_state_from_db(kind)
    fio = (chair.fio or "").strip() if chair else ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if COMMISSION_CHAIR_FIO_PLACEHOLDER in cell.text:
                    cell.text = cell.text.replace(COMMISSION_CHAIR_FIO_PLACEHOLDER, fio)


def fill_commission_members_fio_table(
    doc: Document, *, kind: str = COMMISSION_KIND_OT
) -> None:
    """
    Таблица «Члены комиссии»: строка-шаблон с COMMISSION_MEMBER_FIO_PLACEHOLDER.
    • Если маркер во 2-й и 3-й колонках (индексы 1 и 2) — в строке два ФИО; число строк = ceil(N/2).
    • Если маркер только в одной колонке — по одному ФИО на строку (как раньше).
    При нуле членов строка-шаблон удаляется.
    """
    _, _, _chair, members = load_commission_state_from_db(kind)
    fio_list = [(m.fio or "").strip() for m in members if (m.fio or "").strip()]

    col_left, col_right = 1, 2

    for table in doc.tables:
        tpl_row_idx: int | None = None
        pair_mode = False
        fio_single_col = col_left

        for ri, row in enumerate(table.rows):
            if len(row.cells) <= col_right:
                continue
            has_l = COMMISSION_MEMBER_FIO_PLACEHOLDER in row.cells[col_left].text
            has_r = COMMISSION_MEMBER_FIO_PLACEHOLDER in row.cells[col_right].text
            if has_l and has_r:
                tpl_row_idx = ri
                pair_mode = True
                break
            if has_l:
                tpl_row_idx = ri
                pair_mode = False
                fio_single_col = col_left
                break
            if has_r:
                tpl_row_idx = ri
                pair_mode = False
                fio_single_col = col_right
                break

        if tpl_row_idx is None:
            continue

        template_tr = table.rows[tpl_row_idx]._tr
        n = len(fio_list)
        if n == 0:
            template_tr.getparent().remove(template_tr)
            return

        if pair_mode:
            num_rows = (n + 1) // 2
            _insert_duplicate_tr_after(template_tr, num_rows - 1)
            for r in range(num_rows):
                row = table.rows[tpl_row_idx + r]
                li, rio = 2 * r, 2 * r + 1
                row.cells[col_left].text = fio_list[li] if li < n else ""
                row.cells[col_right].text = fio_list[rio] if rio < n else ""
                for ci, cell in enumerate(row.cells):
                    if ci not in (col_left, col_right) and COMMISSION_MEMBER_FIO_PLACEHOLDER in cell.text:
                        cell.text = ""
        else:
            _insert_duplicate_tr_after(template_tr, n - 1)
            for j in range(n):
                row = table.rows[tpl_row_idx + j]
                for ci, cell in enumerate(row.cells):
                    if ci == fio_single_col:
                        cell.text = fio_list[j]
                    elif COMMISSION_MEMBER_FIO_PLACEHOLDER in cell.text:
                        cell.text = ""
        return


def apply_protocol_signature_placeholders_in_template(
    doc: Document,
    *,
    chair_text: str,
    members_text: str,
) -> None:
    """
    В основном шаблоне протокола заменяет плейсхолдеры подписей комиссии (блоки с подчёркиваниями и строкой
    «…, подпись» — см. build_commission_signature_suffix_payload в commission_admin).
    Плейсхолдеры: {{ПРЕДСЕДАТЕЛЬ}}, {{ЧЛЕНЫ_КОМИССИИ}}, {{CHAIR}}, {{MEMBERS}}.
    """
    repl = (
        ("{{ПРЕДСЕДАТЕЛЬ}}", chair_text or ""),
        ("{{ЧЛЕНЫ_КОМИССИИ}}", members_text or ""),
        ("{{CHAIR}}", chair_text or ""),
        ("{{MEMBERS}}", members_text or ""),
    )
    for para in _iter_all_paragraphs_in_document(doc):
        t = para.text
        if not t:
            continue
        new_t = t
        for key, val in repl:
            if key in new_t:
                new_t = new_t.replace(key, val)
        if new_t == t:
            continue
        if "\n" in new_t:
            _replace_paragraph_text_preserve_style_multiline(para, new_t)
        else:
            _replace_paragraph_text_preserve_style(para, new_t)


def apply_document_marker_replacements(
    doc: Document, replacements: Sequence[tuple[str, str]]
) -> None:
    """Замена подстрок во всех абзацах и ячейках таблиц (плейсхолдеры шаблона)."""

    def subst(text: str) -> str:
        new_t = text
        for a, b in replacements:
            if a in new_t:
                new_t = new_t.replace(a, b)
        return new_t

    for para in _iter_all_paragraphs_in_document(doc):
        t = para.text
        new_t = subst(t)
        if new_t == t:
            continue
        if "\n" in new_t:
            _replace_paragraph_text_preserve_style_multiline(para, new_t)
        else:
            _replace_paragraph_text_preserve_style(para, new_t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                new_t = subst(t)
                if new_t != t:
                    cell.text = new_t


def apply_commission_context_markers_in_document(
    doc: Document,
    *,
    kind: str = COMMISSION_KIND_OT,
) -> None:
    """Подразделение и «кем утверждён приказ» из «Приказ и комиссия» → плейсхолдеры в шаблоне."""
    venue, approver = load_commission_protocol_context_from_db(kind)
    repl = (
        (COMMISSION_VENUE_PLACEHOLDER, venue),
        ("{{VENUE_SUBDIVISION}}", venue),
        (COMMISSION_ORDER_APPROVER_PLACEHOLDER, approver),
        ("{{ORDER_APPROVER}}", approver),
    )
    apply_document_marker_replacements(doc, repl)


def apply_technical_program_markers_in_document(
    doc: Document,
    *,
    approver: str,
    program_name: str,
    approval_date_display: str,
) -> None:
    """Плейсхолдеры из листа Tech_V (шаблон default_protocol_tehnicheskiy.docx и аналоги)."""
    repl = (
        ("{{TECH_APPROVER}}", approver),
        ("{{TECH_PROGRAM}}", program_name),
        ("{{TECH_APPROVAL_DATE}}", approval_date_display),
        ("{{ТЕХ_УТВЕРДИЛ}}", approver),
        ("{{ТЕХ_ПРОГРАММА}}", program_name),
        ("{{ТЕХ_ДАТА_УТВЕРЖДЕНИЯ}}", approval_date_display),
    )
    apply_document_marker_replacements(doc, repl)


def _protocol_results_table_header_match(table: Table) -> bool:
    """Таблица бланка: заголовок с «п/п» и «Фамилия», не меньше 7 колонок."""
    if not table.rows or len(table.rows[0].cells) < 7:
        return False
    joined = " ".join(c.text for c in table.rows[0].cells).lower().replace("\xa0", " ")
    return "п/п" in joined and "фамилия" in joined


def _normalize_fio_marker(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ").replace("\xa0", " ")).strip().upper()


def _format_registry_number_token(token: str) -> str:
    t = token.strip()
    if not t:
        return ""
    if re.match(r"^№\s*", t):
        return t
    return f"№ {t}"


def _format_n_registry_lines(
    grade: str,
    registry_no: str,
    n_registry_rows: int,
    file_line_tokens: list[str] | None = None,
) -> str:
    """
    Оценка и ровно n_registry_rows строк «№ …» (как в блоке «В»); лишние номера из поля не выводятся.
    При нехватке номеров — строка «№» без текста.
    file_line_tokens — номера из файла реестра; пустые слоты добираются из поля «Регистрационный номер».
    """
    if n_registry_rows <= 0:
        return _format_table_result_grade(grade, registry_no)
    g = (grade or "").strip()
    titled = g[:1].upper() + g[1:] if len(g) > 1 else (g.upper() if g else "")
    regs = merge_registry_tokens(registry_no, n_registry_rows, file_line_tokens)
    lines: list[str] = []
    if titled:
        lines.append(f"{titled},")
    for i in range(n_registry_rows):
        r = regs[i] if i < len(regs) else ""
        lines.append(_format_registry_number_token(r) if r else "№")
    return "\n".join(lines)


def _format_n_registry_lines_numbered(
    grade: str,
    registry_no: str,
    n_registry_rows: int,
    file_line_tokens: list[str] | None = None,
) -> str:
    """
    Как _format_n_registry_lines, но каждая строка реестра с префиксом «n. » (1, 2, …)
    для сопоставления с нумерацией в заголовке блока «В».
    """
    if n_registry_rows <= 0:
        return _format_table_result_grade(grade, registry_no)
    g = (grade or "").strip()
    titled = g[:1].upper() + g[1:] if len(g) > 1 else (g.upper() if g else "")
    regs = merge_registry_tokens(registry_no, n_registry_rows, file_line_tokens)
    lines: list[str] = []
    if titled:
        lines.append(f"{titled},")
    for i in range(n_registry_rows):
        r = regs[i] if i < len(regs) else ""
        token = _format_registry_number_token(r) if r else "№"
        lines.append(f"{i + 1}. {token}")
    return "\n".join(lines)


def _format_table_result_grade(grade: str, registry_no: str) -> str:
    """Оценка + регистрационные номера реестра Минтруда (не номер протокола)."""
    g = (grade or "").strip()
    if not g:
        return ""
    titled = g[:1].upper() + g[1:] if len(g) > 1 else g.upper()
    raw = (registry_no or "").strip()
    if not raw:
        return f"{titled},"
    parts = [p.strip() for p in re.split(r"[\n,;]+", raw) if p.strip()]
    lines = [f"{titled},"]
    lines.extend(_format_registry_number_token(p) for p in parts)
    return "\n".join(lines)


def _find_fio_marker_row_metas(table: Table) -> list[tuple[int, int, int]]:
    """
    Индексы строк таблицы с маркером «ФИО» и номером вида «целое.целое» в 1-й колонке.
    Возвращает список (индекс_строки, major, minor), например (3, 1, 1) для «1.1».
    """
    out: list[tuple[int, int, int]] = []
    for i, row in enumerate(table.rows):
        if len(row.cells) < 7:
            continue
        c0 = row.cells[0].text.strip()
        c1 = row.cells[1].text
        m = re.match(r"^(\d+)\.(\d+)$", c0)
        if not m:
            continue
        if _normalize_fio_marker(c1) != "ФИО":
            continue
        out.append((i, int(m.group(1)), int(m.group(2))))
    return out


def _insert_duplicate_tr_after(template_tr: Any, extra_copies: int) -> None:
    """Вставляет подряд extra_copies копий строки таблицы (XML w:tr) сразу после template_tr."""
    if extra_copies <= 0:
        return
    last = template_tr
    for _ in range(extra_copies):
        new_tr = deepcopy(template_tr)
        last.addnext(new_tr)
        last = new_tr


def _table_tr_elements(tbl: Any) -> list[Any]:
    tag_tr = qn("w:tr")
    return [e for e in tbl if e.tag == tag_tr]


def _fill_static_protocol_result_table(
    doc: Document,
    *,
    persons: list[EmployeeRecord],
    grade: str,
    registry_no: str,
    check_type: str,
) -> tuple[int, int]:
    """
    Старый режим: в шаблоне уже есть блоки 1.1, 2.1, … — только дублируем строки
    под число сотрудников.
    """
    if not persons:
        return 0, 0

    chk = (check_type or "плановая").strip()
    if chk:
        chk = chk[:1].upper() + chk[1:]
    result_text = _format_table_result_grade(grade, registry_no)
    n = len(persons)

    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        markers = _find_fio_marker_row_metas(table)
        if not markers:
            continue
        to_expand = [(idx, major, minor) for idx, major, minor in markers if minor == 1]
        if not to_expand:
            to_expand = list(markers)
        for row_idx, major, _minor in sorted(to_expand, key=lambda x: -x[0]):
            template_tr = table.rows[row_idx]._tr
            _insert_duplicate_tr_after(template_tr, n - 1)
            for j in range(n):
                row = table.rows[row_idx + j]
                row.cells[0].text = f"{major}.{j + 1}"
                p = persons[j]
                row.cells[1].text = p.fio
                row.cells[2].text = p.profession
                row.cells[3].text = p.subdivision
                row.cells[4].text = result_text
                row.cells[5].text = chk
        return n, 0

    return 0, n


def _profession_cell_text_for_table(emp: EmployeeRecord) -> str:
    """Должность в строке таблицы: основная и совмещаемая в одной ячейке."""
    p2 = (emp.profession2 or "").strip()
    if p2:
        return f"{emp.profession}; {p2}" if (emp.profession or "").strip() else p2
    return emp.profession


def _first_employee_template_tr_after_headers(table: Table, header_rows: int = 2) -> Any | None:
    """Первая строка данных с маркером «ФИО» во 2-й ячейке; иначе первая строка под шапкой."""
    trs = _table_tr_elements(table._tbl)
    for i in range(header_rows, len(trs)):
        row = table.rows[i]
        if len(row.cells) < 7:
            continue
        if _normalize_fio_marker(row.cells[1].text) == "ФИО":
            return trs[i]
    return trs[header_rows] if len(trs) > header_rows else None


def _rebuild_technical_protocol_result_table(
    doc: Document,
    *,
    persons: list[EmployeeRecord],
    grade: str,
    registry_no: str,
    check_type: str,
) -> tuple[int, int]:
    """
    Протокол по техническим вопросам: без строки с наименованием программы в таблице,
    без нумерации вида «1.1»; только строки работников по алфавиту ФИО, в графе п/п — 1, 2, 3, …
    """
    persons_sorted = sorted(
        list(persons),
        key=lambda p: (p.fio or "").strip().casefold(),
    )
    if not persons_sorted:
        return 0, 0
    chk = (check_type or "плановая").strip()
    if chk:
        chk = chk[:1].upper() + chk[1:]
    n_tab = len(persons_sorted)

    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        tbl = table._tbl
        trs = _table_tr_elements(tbl)
        header_rows = 2
        if len(trs) <= header_rows:
            return 0, n_tab
        emp_tr_src = _first_employee_template_tr_after_headers(table, header_rows)
        if emp_tr_src is None:
            return 0, n_tab
        for tr in trs[header_rows:]:
            tbl.remove(tr)
        tbl.append(deepcopy(emp_tr_src))
        emp_base_idx = header_rows
        emp_base_row = table.rows[emp_base_idx]
        _insert_duplicate_tr_after(emp_base_row._tr, max(n_tab - 1, 0))
        for j in range(n_tab):
            row = table.rows[emp_base_idx + j]
            p = persons_sorted[j]
            row.cells[0].text = str(j + 1)
            row.cells[1].text = p.fio
            row.cells[2].text = _profession_cell_primary_only(p)
            row.cells[3].text = p.subdivision
            row.cells[4].text = _format_n_registry_lines(
                grade, registry_no, 1, file_line_tokens=None
            )
            row.cells[5].text = chk
        return n_tab, 0

    return 0, n_tab


def _rebuild_protocol_result_table(
    doc: Document,
    *,
    program_titles: list[str],
    program_keys: list[str] | None,
    excel_path: Path | None,
    persons: list[EmployeeRecord],
    persons_v_raw: list[EmployeeRecord] | None,
    grade: str,
    registry_no: str,
    check_type: str,
    persons_b_all_rows: list[EmployeeRecord] | None = None,
    trained_registry_index: TrainedRegistryIndex | None = None,
    protocol_no_for_registry: str = "",
    date_str_for_registry: str = "",
) -> tuple[int, int]:
    """
    Удаляет строки таблицы ниже шапки (две первые строки-заголовка) и строит заново:
    для каждой выбранной программы — строка с номером и полным названием из Excel,
    затем строки сотрудников N.1…N.M. Для «В» перед каждой строкой работника — ещё строка-шапка
    с перечнем тем этого работника; для ПП и СИЗ — по объединённому списку ФИО;
    для «Б» — по persons_b_all_rows (все исходные строки, даже с одинаковым ФИО), если задан.
    """
    merged_rows = persons_v_raw if persons_v_raw is not None else persons
    b_rows = persons_b_all_rows if persons_b_all_rows is not None else merged_rows
    if not merged_rows:
        return 0, 0
    if not program_titles:
        return 0, len(merged_rows)

    chk = (check_type or "плановая").strip()
    if chk:
        chk = chk[:1].upper() + chk[1:]
    n_tab = len(merged_rows)
    keys = program_keys if program_keys and len(program_keys) == len(program_titles) else [""] * len(
        program_titles
    )
    v_path = excel_path if excel_path and excel_path.is_file() else None
    idx = trained_registry_index
    protocol_queries: list[str] = []
    raw_pn = (protocol_no_for_registry or "").strip()
    if raw_pn:
        protocol_queries.append(raw_pn)
    ds_reg = (date_str_for_registry or "").strip()
    if ds_reg:
        fmt_pn = format_protocol_number_for_template(raw_pn, ds_reg)
        if fmt_pn and fmt_pn not in protocol_queries:
            protocol_queries.append(fmt_pn)

    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        tbl = table._tbl
        trs = _table_tr_elements(tbl)
        header_rows = 2
        if len(trs) < header_rows + 2:
            return 0, n_tab
        prog_tpl = deepcopy(trs[header_rows])
        emp_tpl = deepcopy(trs[header_rows + 1])
        for tr in trs[header_rows:]:
            tbl.remove(tr)

        v_reg_rows_cached: list[tuple[str, str, str, float | None]] = (
            get_cached_v_registry_rows(v_path) if v_path is not None else []
        )

        for pi, (pkey, title) in enumerate(zip(keys, program_titles), start=1):
            pk = (pkey or "").strip().upper()
            is_v = pk == "V" and v_path is not None
            is_b = pk == "B"
            block_rows = b_rows if is_b else merged_rows
            n_block = len(block_rows)

            if is_v:
                v_fb = _protocol_program_fallback_title("V")
                for j in range(n_block):
                    p = block_rows[j]
                    v_parts = v_program_merged_parts_for_raw_employee(v_path, p)
                    _, v_hints = v_fragment_labels_and_hints(v_parts)
                    sub_title = format_v_program_table_block_title(v_parts, v_fb)
                    h_person = _v_person_training_hours_sum(v_parts, v_reg_rows_cached)
                    block_title = _protocol_block_title_with_hours(
                        sub_title, h_person, multiline=True
                    )
                    p_tr = deepcopy(prog_tpl)
                    tbl.append(p_tr)
                    row_p = table.rows[-1]
                    row_p.cells[0].text = str(pi) if j == 0 else ""
                    for ci in range(1, min(7, len(row_p.cells))):
                        row_p.cells[ci].text = block_title
                    e_tr = deepcopy(emp_tpl)
                    tbl.append(e_tr)
                    row = table.rows[-1]
                    row.cells[0].text = f"{pi}.{j + 1}"
                    row.cells[1].text = p.fio
                    row.cells[2].text = _profession_cell_primary_only(p)
                    row.cells[3].text = p.subdivision
                    v_file: list[str] | None = None
                    if idx is not None:
                        cand_all = idx.candidates_for_employee(
                            p,
                            protocol_queries,
                            apply_profession_filter_for_registry=False,
                        )
                        cand_v = filter_candidates_for_program_block(cand_all, v_hints)
                        v_file = idx.registry_numbers_for_hints(
                            cand_v,
                            v_hints,
                            emp=p,
                            require_profession_for_registry=False,
                        )
                    row.cells[4].text = _format_v_result_cell(
                        grade, registry_no, v_parts, v_file_tokens=v_file
                    )
                    row.cells[5].text = chk
            else:
                block_h = (
                    None
                    if pk == "B"
                    else get_training_hours_for_program_key(v_path, pk)
                )
                title_with_h = _protocol_block_title_with_hours(
                    title, block_h, multiline=False
                )
                p_tr = deepcopy(prog_tpl)
                tbl.append(p_tr)
                row_p = table.rows[-1]
                row_p.cells[0].text = str(pi)
                for ci in range(1, min(7, len(row_p.cells))):
                    row_p.cells[ci].text = title_with_h

                e_tr = deepcopy(emp_tpl)
                tbl.append(e_tr)
                emp_base_idx = len(table.rows) - 1
                emp_base_row = table.rows[emp_base_idx]
                _insert_duplicate_tr_after(emp_base_row._tr, max(n_block - 1, 0))
                for j in range(n_block):
                    row = table.rows[emp_base_idx + j]
                    p = block_rows[j]
                    row.cells[0].text = f"{pi}.{j + 1}"
                    row.cells[1].text = p.fio
                    row.cells[2].text = _profession_cell_primary_only(p)
                    row.cells[3].text = p.subdivision
                    if pk == "TECH":
                        row.cells[4].text = _format_n_registry_lines(
                            grade, registry_no, 1, file_line_tokens=None
                        )
                    else:
                        n_reg = _rebuild_registry_rows_for_program(pkey or "", p)
                        n_file: list[str] | None = None
                        if idx is not None:
                            registry_strict_prof = (pkey or "").strip().upper() == "B"
                            cand_all = idx.candidates_for_employee(
                                p,
                                protocol_queries,
                                apply_profession_filter_for_registry=registry_strict_prof,
                            )
                            cand_blk = filter_candidates_for_program_block(cand_all, [title])
                            hints = [title] * n_reg
                            n_file = idx.registry_numbers_for_hints(
                                cand_blk,
                                hints,
                                emp=p,
                                require_profession_for_registry=registry_strict_prof,
                            )
                        row.cells[4].text = _format_n_registry_lines(
                            grade, registry_no, n_reg, file_line_tokens=n_file
                        )
                    row.cells[5].text = chk

        return n_tab, 0

    return 0, n_tab


def fill_protocol_result_table(
    doc: Document,
    *,
    program_titles: list[str] | None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons: list[EmployeeRecord],
    persons_v_raw: list[EmployeeRecord] | None = None,
    grade: str,
    registry_no: str,
    check_type: str,
    persons_b_all_rows: list[EmployeeRecord] | None = None,
    trained_registry_index: TrainedRegistryIndex | None = None,
    protocol_no_for_registry: str = "",
    date_str_for_registry: str = "",
    technical_protocol: bool = False,
) -> tuple[int, int]:
    if technical_protocol and program_titles:
        merged = persons_v_raw if persons_v_raw is not None else persons
        return _rebuild_technical_protocol_result_table(
            doc,
            persons=merged,
            grade=grade,
            registry_no=registry_no,
            check_type=check_type,
        )
    if program_titles:
        return _rebuild_protocol_result_table(
            doc,
            program_titles=program_titles,
            program_keys=program_keys,
            excel_path=excel_path,
            persons=persons,
            persons_v_raw=persons_v_raw,
            grade=grade,
            registry_no=registry_no,
            check_type=check_type,
            persons_b_all_rows=persons_b_all_rows,
            trained_registry_index=trained_registry_index,
            protocol_no_for_registry=protocol_no_for_registry,
            date_str_for_registry=date_str_for_registry,
        )
    return _fill_static_protocol_result_table(
        doc,
        persons=persons,
        grade=grade,
        registry_no=registry_no,
        check_type=check_type,
    )


def build_filled_protocol_document(
    template_path: Path,
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
    table_persons: list[EmployeeRecord],
    program_titles: list[str] | None = None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons_v_raw: list[EmployeeRecord] | None = None,
    persons_b_row_source: list[EmployeeRecord] | None = None,
    grade: str = "",
    registry_no: str = "",
    check_type: str = "плановая",
    trained_registry_path: Path | None = None,
    technical_protocol: bool = False,
    tech_approver: str = "",
    tech_program_name: str = "",
    tech_approval_date_raw: str = "",
) -> tuple[Document, int]:
    """
    Загружает шаблон .docx, подставляет поля, возвращает (документ, «остаток»).
    program_titles — названия: «Б» с листа B, остальное с V_PROF; пересборка таблицы.
    trained_registry_path — выгрузка «Реестр обученных…» с портала Минтруда (.xlsx).
    persons_b_row_source — сырой список выбора для блока «Б» (несколько строк Excel с одним ФИО
    и разными должностями); если None — берётся emp_source, после разворота profession/profession2.
    technical_protocol — протокол по техническим вопросам: реестр Минтруда не читается; комиссия
    из вкладки «Технич. вопросы»; данные программы — лист Tech_V и плейсхолдеры {{TECH_*}} / {{ТЕХ_*}}.
    """
    doc = Document(str(template_path))
    commission_kind = COMMISSION_KIND_TECH if technical_protocol else COMMISSION_KIND_OT
    apply_commission_context_markers_in_document(doc, kind=commission_kind)
    trained_idx = load_trained_registry_index(
        None if technical_protocol else trained_registry_path
    )
    emp_source = persons_v_raw if persons_v_raw is not None else table_persons
    emp_all = list(emp_source)
    emp_for_doc = _table_employees_dedupe_by_fio(emp_all)
    b_src = persons_b_row_source if persons_b_row_source is not None else emp_all
    persons_b_all_rows = expand_persons_block_b_rows(b_src)
    paras = _all_document_paragraphs_ordered(doc)
    lines = [p.text for p in paras]
    start, end = _find_form_template_bounds(lines)

    if program_titles and program_keys and len(program_keys) == len(program_titles):
        paragraph_theme = build_protocol_header_theme_text(
            excel_path,
            program_keys,
            program_titles,
            persons_b_all_rows,
        )
    elif program_titles:
        paragraph_theme = "; ".join(program_titles).strip()
    else:
        paragraph_theme = (theme or "").strip()

    commission_payload = build_commission_template_payload(
        _format_date_protocol_line, kind=commission_kind
    )
    program_header_paragraphs: list[DocxParagraph] = []
    i = start
    while i < end:
        para = paras[i]
        line = lines[i]
        nxt_line = lines[i + 1] if i + 1 < end else None

        if (
            nxt_line is not None
            and _is_protocol_title_line_without_underscores(
                line, allow_technical_heading=technical_protocol
            )
            and _is_placeholder_only_underscore_line(nxt_line)
        ):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            t1 = line
            t2 = (
                _replace_first_underscore_run_with_protocol_number(nxt_line, pn)
                if pn
                else nxt_line
            )
            t1 = apply_commission_insertions_to_line(
                t1,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            t2 = apply_commission_insertions_to_line(
                t2,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            if t1 != line:
                _replace_paragraph_text_preserve_style(para, t1)
            p2 = paras[i + 1]
            if t2 != nxt_line:
                _replace_paragraph_text_preserve_style(p2, t2)
            i += 2
        elif _line_fills_protocol_number_slot(
            line, allow_technical_heading=technical_protocol
        ):
            pn = format_protocol_number_for_template(protocol_no, date_str)
            t = _apply_protocol_number_after_anchor(line, pn) if pn else line
            t = apply_commission_insertions_to_line(
                t,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            if t != line:
                _replace_paragraph_text_preserve_style(para, t)
            i += 1
        else:
            t = line
            if line.strip().startswith("«__»") and line.strip().endswith("г."):
                t = _format_date_protocol_line(date_str)
            elif _is_program_underscore_line(line) and paragraph_theme:
                t = paragraph_theme + ","
                program_header_paragraphs.append(para)
            elif technical_protocol and (tech_program_name or "").strip():
                nxt_prog = lines[i + 1] if i + 1 < end else None
                if _is_technical_program_intro_paragraph(
                    line, technical_protocol=technical_protocol
                ) and not _line_has_tech_v_placeholder_marker(line):
                    if nxt_prog is None or not _line_has_tech_v_placeholder_marker(
                        nxt_prog
                    ):
                        t = _inline_tech_program_in_intro_line(line, tech_program_name)
            t = apply_commission_insertions_to_line(
                t,
                date_words=commission_payload["date_words"],
                order_no=commission_payload["order_no"],
                chair_gen=commission_payload["chair"],
                members_gen=commission_payload["members"],
            )
            if t != line:
                _apply_commission_paragraph_replacement(
                    para,
                    line,
                    t,
                    members_gen=commission_payload["members"],
                    chair_gen=commission_payload["chair"],
                )
            i += 1

    if excel_path and excel_path.is_file() and program_keys and persons_b_all_rows:
        fg_lines = build_fg_lines_for_selected_programs(
            excel_path, program_keys, persons_b_all_rows
        )
        insert_program_fg_lines_after_anchor(doc, fg_lines)

    _, excess = fill_protocol_result_table(
        doc,
        program_titles=program_titles if program_titles else None,
        program_keys=program_keys,
        excel_path=excel_path,
        persons=emp_for_doc,
        persons_v_raw=emp_for_doc,
        grade=grade,
        registry_no=registry_no,
        check_type=check_type,
        persons_b_all_rows=persons_b_all_rows,
        trained_registry_index=trained_idx,
        protocol_no_for_registry=protocol_no,
        date_str_for_registry=date_str,
        technical_protocol=technical_protocol,
    )

    lines_after = [p.text for p in _all_document_paragraphs_ordered(doc)]
    start_body, end_body = _find_form_template_bounds(lines_after)
    _apply_protocol_body_font_pt(
        doc,
        start_body,
        end_body,
        PROTOCOL_BODY_FONT_PT,
        paragraph_theme=paragraph_theme,
    )

    fill_commission_chair_fio_table(doc, kind=commission_kind)
    fill_commission_members_fio_table(doc, kind=commission_kind)

    chair_sig, members_sig = build_commission_signature_suffix_payload(kind=commission_kind)
    apply_protocol_signature_placeholders_in_template(
        doc,
        chair_text=chair_sig,
        members_text=members_sig,
    )

    if technical_protocol:
        td = _format_tech_approval_date_for_template(tech_approval_date_raw)
        apply_technical_program_markers_in_document(
            doc,
            approver=(tech_approver or "").strip(),
            program_name=(tech_program_name or paragraph_theme or "").strip(),
            approval_date_display=td,
        )

    _apply_protocol_program_title_header_fonts(
        doc, program_header_paragraphs, paragraph_theme
    )

    return doc, excess


def save_protocol_docx_from_template(
    template_path: Path,
    output_path: str,
    *,
    protocol_no: str,
    date_str: str,
    theme: str,
    table_persons: list[EmployeeRecord],
    program_titles: list[str] | None = None,
    program_keys: list[str] | None = None,
    excel_path: Path | None = None,
    persons_v_raw: list[EmployeeRecord] | None = None,
    persons_b_row_source: list[EmployeeRecord] | None = None,
    grade: str = "",
    registry_no: str = "",
    check_type: str = "плановая",
    trained_registry_path: Path | None = None,
    technical_protocol: bool = False,
    tech_approver: str = "",
    tech_program_name: str = "",
    tech_approval_date_raw: str = "",
) -> None:
    doc, _ = build_filled_protocol_document(
        template_path,
        protocol_no=protocol_no,
        date_str=date_str,
        theme=theme,
        table_persons=table_persons,
        program_titles=program_titles,
        program_keys=program_keys,
        excel_path=excel_path,
        persons_v_raw=persons_v_raw,
        persons_b_row_source=persons_b_row_source,
        grade=grade,
        registry_no=registry_no,
        check_type=check_type,
        trained_registry_path=trained_registry_path,
        technical_protocol=technical_protocol,
        tech_approver=tech_approver,
        tech_program_name=tech_program_name,
        tech_approval_date_raw=tech_approval_date_raw,
    )
    doc.save(output_path)


def _iter_paragraph_runs(paragraph: DocxParagraph):
    for item in paragraph.iter_inner_content():
        if isinstance(item, DocxRun):
            yield item
        elif isinstance(item, Hyperlink):
            for r in item.runs:
                yield r


def _apply_font_pt_to_paragraph_runs(para: DocxParagraph, pt: int) -> None:
    half = str(int(round(pt * 2)))
    for run in _iter_paragraph_runs(para):
        run.font.size = Pt(pt)
    p_pr = para._element.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.insert(0, r_pr)
    for tag in ("w:sz", "w:szCs"):
        el = r_pr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            r_pr.append(el)
        el.set(qn("w:val"), half)


def _paragraph_is_program_titles_header(text: str, paragraph_theme: str) -> bool:
    """Абзац с перечнем программ в шапке (строка подчёркиваний → тема через «; »)."""
    s = (text or "").replace("\xa0", " ").strip()
    theme = (paragraph_theme or "").replace("\xa0", " ").strip()
    if not s:
        return False
    if theme:
        if s in (theme, theme + ","):
            return True
        if len(theme) >= 15 and theme in s:
            return True
    if not s.endswith(","):
        return False
    core = s[:-1].replace("_", "").strip()
    if len(core) < 20:
        return False
    return ";" in s or _is_program_underscore_line(s) or len(core) >= 40


def _apply_protocol_program_title_header_fonts(
    doc: Document,
    tracked_paragraphs: list[DocxParagraph],
    paragraph_theme: str,
) -> None:
    seen: set[int] = set()
    for para in tracked_paragraphs:
        pid = id(para._p)
        if pid in seen:
            continue
        seen.add(pid)
        _apply_font_pt_to_paragraph_runs(para, PROTOCOL_PROGRAM_TITLE_FONT_PT)
    for para in doc.paragraphs:
        pid = id(para._p)
        if pid in seen:
            continue
        if _paragraph_is_program_titles_header(para.text, paragraph_theme):
            seen.add(pid)
            _apply_font_pt_to_paragraph_runs(para, PROTOCOL_PROGRAM_TITLE_FONT_PT)


def _is_v_program_table_header_row(row) -> bool:
    """Строка-шапка блока «В» в таблице (до строки сотрудника)."""
    cells = row.cells
    if len(cells) < 2:
        return False
    return (cells[1].text or "").lstrip().startswith("Программа (В)")


def _is_protocol_program_title_table_row(row) -> bool:
    """Строка с полным наименованием программы в таблице (не строка сотрудника N.M)."""
    if _is_v_program_table_header_row(row):
        return True
    cells = row.cells
    if len(cells) < 2:
        return False
    c0 = (cells[0].text or "").strip()
    if not c0 or "." in c0:
        return False
    if not c0.isdigit():
        return False
    return len((cells[1].text or "").strip()) > 12


def _apply_protocol_body_font_pt(
    doc: Document,
    form_start: int,
    form_end: int,
    pt: int,
    *,
    paragraph_theme: str = "",
) -> None:
    """
    Абзацы бланка [form_start, form_end) — без организационной шапки над «ПРОТОКОЛ №»;
    в таблице «Результат проверки» — все строки ниже двух верхних строк шапки.
    Наименования программ (шапка и заголовки блоков таблицы) — PROTOCOL_PROGRAM_TITLE_FONT_PT.
    """
    paras = _all_document_paragraphs_ordered(doc)
    last = min(form_end, len(paras))
    for i in range(form_start, last):
        ptxt = paras[i].text or ""
        if _paragraph_is_program_titles_header(ptxt, paragraph_theme):
            _apply_font_pt_to_paragraph_runs(paras[i], PROTOCOL_PROGRAM_TITLE_FONT_PT)
        else:
            _apply_font_pt_to_paragraph_runs(paras[i], pt)

    for para in doc.paragraphs:
        if not _is_anchor_paragraph_program_table(para.text):
            continue
        lo, hi = _program_bullet_paragraph_indices_after_anchor(doc, para)
        body_paras = list(doc.paragraphs)
        for j in range(lo, hi):
            if j < len(body_paras):
                _apply_font_pt_to_paragraph_runs(
                    body_paras[j], PROTOCOL_PROGRAM_TITLE_FONT_PT
                )
        break

    prog_pt = PROTOCOL_PROGRAM_TITLE_FONT_PT
    for table in doc.tables:
        if not _protocol_results_table_header_match(table):
            continue
        header_rows = 2
        for ri, row in enumerate(table.rows):
            if ri < header_rows:
                continue
            row_pt = prog_pt if _is_protocol_program_title_table_row(row) else pt
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for para in cell.paragraphs:
                    _apply_font_pt_to_paragraph_runs(para, row_pt)


def document_to_plain_text(doc: Document) -> str:
    lines = [p.text for p in doc.paragraphs]
    if doc.tables:
        lines.append("")
        lines.append("── Таблица (текст ячеек) ──")
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                if not any(cells):
                    continue
                lines.append("  |  ".join(cells))
    return "\n".join(lines)

