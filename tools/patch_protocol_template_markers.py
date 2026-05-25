# -*- coding: utf-8 -*-
"""
Вставляет в шаблоны протоколов плейсхолдеры подразделения и утверждения приказа.
Закройте .docx в Word. Из корня: py -3 tools/patch_protocol_template_markers.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx_template_protection import protect_standard_protocol_templates
VENUE = "{{ПОДРАЗДЕЛЕНИЕ_ПРОВЕРКИ}}"
APPROVER_LINE = (
    "В соответствии с приказом {{УТВЕРДИЛ_ПРИКАЗ}} "
    "от «__» _______ 20__ г. № ___"
)

TARGETS = [
    ROOT / "default_protocol.docx",
    ROOT / "default_protocol_tehnicheskiy.docx",
    ROOT / "bundle" / "default_protocol.docx",
    ROOT / "bundle" / "default_protocol_tehnicheskiy.docx",
]


def _set_para_text(para, text: str) -> None:
    para.clear()
    para.add_run(text)


def patch_ot(doc: Document) -> int:
    n = 0
    if len(doc.paragraphs) > 3:
        p = doc.paragraphs[3]
        if VENUE not in (p.text or ""):
            _set_para_text(p, VENUE)
            n += 1
    if len(doc.paragraphs) > 8:
        p = doc.paragraphs[8]
        if "{{УТВЕРДИЛ_ПРИКАЗ}}" not in (p.text or ""):
            _set_para_text(p, APPROVER_LINE)
            n += 1
    return n


def patch_tech(doc: Document) -> int:
    n = 0
    if len(doc.paragraphs) > 3:
        p = doc.paragraphs[3]
        t = (p.text or "").strip()
        if VENUE not in t and ("___" in t or "организац" in t.lower() or len(t) < 3):
            _set_para_text(p, VENUE)
            n += 1
    if len(doc.paragraphs) > 8:
        p = doc.paragraphs[8]
        if "{{УТВЕРДИЛ_ПРИКАЗ}}" not in (p.text or ""):
            _set_para_text(p, APPROVER_LINE)
            n += 1
    if len(doc.paragraphs) > 9:
        p = doc.paragraphs[9]
        t = (p.text or "").strip().lower()
        if VENUE not in (p.text or "") and "наименование организации" in t:
            _set_para_text(p, VENUE)
            n += 1
    return n


def main() -> int:
    ok = 0
    for path in TARGETS:
        if not path.is_file():
            print(f"Пропуск (нет файла): {path}")
            continue
        try:
            doc = Document(str(path))
        except Exception as e:
            print(f"Ошибка открытия {path}:\n{e}", file=sys.stderr)
            return 1
        tech = "tehnicheskiy" in path.name.lower()
        ch = patch_tech(doc) if tech else patch_ot(doc)
        if ch == 0:
            print(f"Без изменений: {path}")
            continue
        try:
            doc.save(str(path))
        except OSError as e:
            print(f"Не удалось сохранить {path}:\n{e}", file=sys.stderr)
            return 1
        print(f"Обновлено абзацев: {ch} — {path}")
        ok += 1
    if ok == 0:
        print("Ни один файл не изменён (возможно, маркеры уже стоят).")
    prot = protect_standard_protocol_templates(ROOT)
    for line in prot:
        print(f"Защита Word (только чтение): {line}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
