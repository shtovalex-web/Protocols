# -*- coding: utf-8 -*-
"""Конвертация инструкций Markdown → Word (заголовки, списки, таблицы, **жирный**)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

_SEG = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
_NUM = re.compile(r"^(\d+)\.\s+(.*)$")
_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|$")
_IMG = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")


def _iter_md_lines_skip_html_comments(lines: list[str]):
    in_comment = False
    for raw in lines:
        s = raw.strip()
        if in_comment:
            if "-->" in s:
                in_comment = False
            continue
        if s.startswith("<!--"):
            if "-->" in s:
                continue
            in_comment = True
            continue
        yield raw


def _set_body_font(doc: Document, name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)


def _add_runs(paragraph, text: str) -> None:
    pos = 0
    for m in _SEG.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        seg = m.group(0)
        if seg.startswith("**"):
            run = paragraph.add_run(seg[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(seg[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _paragraph(doc: Document, text: str) -> None:
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    _add_runs(p, text)


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().split("|") if c.strip()]


def _add_image(doc: Document, md_path: Path, caption: str, rel_path: str) -> None:
    img_path = (md_path.parent / rel_path).resolve()
    if not img_path.is_file():
        p = doc.add_paragraph()
        p.add_run(f"[Снимок экрана: {caption or rel_path} — файл не найден: {rel_path}]")
        return
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(img_path), width=Inches(6.2))
    cap_text = (caption or "").strip()
    if cap_text:
        cap = doc.add_paragraph(cap_text)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(11)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            _add_runs(cell.paragraphs[0], text)


def build_docx_from_markdown(md_path: Path, out_path: Path) -> None:
    if not md_path.is_file():
        raise FileNotFoundError(md_path)

    doc = Document()
    _set_body_font(doc)
    lines = list(_iter_md_lines_skip_html_comments(md_path.read_text(encoding="utf-8").splitlines()))
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line.startswith("|") and not _TABLE_SEP.match(line):
            table_rows: list[list[str]] = []
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line.startswith("|"):
                    break
                if _TABLE_SEP.match(row_line):
                    i += 1
                    continue
                table_rows.append(_parse_table_row(row_line))
                i += 1
            _add_table(doc, table_rows)
            continue

        if not line:
            i += 1
            continue

        img_m = _IMG.match(line)
        if img_m:
            _add_image(doc, md_path, img_m.group(1), img_m.group(2))
            i += 1
            continue

        if line == "---":
            doc.add_paragraph()
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
        else:
            nm = _NUM.match(line)
            if nm:
                p = doc.add_paragraph(style="List Number")
                _add_runs(p, nm.group(2).strip())
            else:
                _paragraph(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
