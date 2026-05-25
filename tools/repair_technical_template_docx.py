# -*- coding: utf-8 -*-
"""
Восстанавливает в default_protocol_tehnicheskiy.docx полную строку приказа с плейсхолдерами даты и №,
если абзац обрезан на «… от» (частая ошибка после правок в Word).

Закройте файл в Word, затем из корня проекта:
  py -3 tools/repair_technical_template_docx.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
TARGET = ROOT / "bundle" / "default_protocol_tehnicheskiy.docx"

COMMISSION_LINE = (
    "В соответствии с приказом {{УТВЕРДИЛ_ПРИКАЗ}} "
    "от «__» _______ 20__ г. № ___"
)


def main() -> int:
    if not TARGET.is_file():
        print(f"Не найден файл:\n{TARGET}", file=sys.stderr)
        return 1
    try:
        doc = Document(str(TARGET))
    except Exception as e:
        print(f"Не удалось открыть:\n{e}", file=sys.stderr)
        return 1
    if len(doc.paragraphs) < 10:
        print("Неожиданно мало абзацев в шаблоне.", file=sys.stderr)
        return 1
    p8 = doc.paragraphs[8]
    cur = (p8.text or "").strip()
    if "20__" in cur and "№" in cur and "«__»" in cur:
        print("Абзац приказа уже полный — изменений не требуется.")
        return 0
    if "приказом" not in cur.lower() or "руководител" not in cur.lower():
        print(
            f"Абзац 8 не похож на строку приказа, пропуск:\n{cur[:200]}",
            file=sys.stderr,
        )
        return 1
    p8.clear()
    p8.add_run(COMMISSION_LINE)
    try:
        doc.save(str(TARGET))
    except OSError as e:
        print(
            f"Не удалось сохранить (закройте файл в Word и повторите):\n{e}",
            file=sys.stderr,
        )
        return 1
    print(f"Обновлён абзац приказа в:\n{TARGET}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
